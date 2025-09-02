import sqlite_fix  # This must be imported first

# Apply SQLite3 fix before any other imports that might use ChromaDB
import sys
try:
    import pysqlite3
    sys.modules['sqlite3'] = pysqlite3
    print("✅ SQLite3 fix applied in document_processor")
except ImportError:
    print("⚠️ pysqlite3-binary not available in document_processor")

from fastapi import UploadFile, HTTPException
from PIL import Image
import os
from dotenv import load_dotenv
import pdfplumber
from io import BytesIO
from typing import List, Dict, Union, Optional, Any
from docx import Document
from google.cloud import vision
import io
from google.cloud.vision import ImageAnnotatorClient
import pandas as pd
from letter_format import LetterFormatStorage
from rag_storage import RAGLegalStorage
# from rag_storage_local import RAGLegalStorage
from pydantic_ai import Agent
from pydantic_ai.models.google import GoogleModel
from pydantic_ai.providers.google import GoogleProvider
from pydantic_ai.settings import ModelSettings
from pydantic import BaseModel
import asyncio
from agentic_doc.parse import parse
from evaluator_optimizer import EvaluatorOptimizer
# nest_asyncio imported conditionally inside __init__ to avoid uvloop conflicts
# Load environment variables from .env 
load_dotenv()



class DocumentAnalysisRequest(BaseModel):
    """Request model for document analysis"""
    payslip_text: Optional[str] = None
    contract_text: Optional[str] = None
    attendance_text: Optional[str] = None
    analysis_type: str = "report"
    context: Optional[str] = None

class DocumentAnalysisResponse(BaseModel):
    """Response model for document analysis"""
    legal_analysis: str
    # relevant_laws: List[Dict]
    # relevant_judgements: List[Dict]
    status: str
    analysis_type: str
    formatted_laws: Optional[str] = None
    formatted_judgements: Optional[str] = None
    documents: Optional[Dict] = None

def get_error_detail(e):
    try:
        # Try accessing the structured body if it's JSON-like
        if hasattr(e, "body"):
            import json
            body = e.body
            # If it's a string, parse it
            if isinstance(body, str):
                body = json.loads(body)
            return body.get("error", {}).get("message", str(e))

        # Fallback to using .message
        if hasattr(e, "message"):
            return str(e.message)

        return str(e)

    except Exception:
        return repr(e)


class DocumentProcessor:
    async def export_to_excel(self, processed_result: dict) -> bytes:
        """
        Process the processed_result, use AI agent to extract employee name, overtime hours, salary, and all relevant data, and generate an Excel file for download.
        Returns the Excel file as bytes.
        """
        import pandas as pd
        import io
        import asyncio

        # Compose a prompt for the AI agent to extract structured data
        prompt = f"""
        Extract a table of all employees from the following legal document analysis results. For each employee, extract the following fields (if available):
        - שם העובד (Employee Name)
        - חודש עבודה (Work Month)
        - שנה (Year)
        - סה"כ ימי עבודה (Total Work Days)
        - סה"כ שעות עבודה (Total Work Hours)
        - תעריף שעתי (Hourly Rate)
        - משכורת 100 % (100% Salary)
        - ש.נ. 125% (Overtime 125%)
        - ש.נ. 150% (Overtime 150%)
        - ש.נ. 175% (Overtime 175%)
        - ש.נ. 200% (Overtime 200%)
        - שעות שבת (Shabbat Hours)
        - שעות חג (Holiday Hours)
        - חופשה (Vacation)
        - ימי מחלה (Sick Days)
        - הבראה (Convalescence)
        - נסיעות (Travel)
        - הפרשות מעסיק לפנסיה (Employer Pension Contributions)
        - הפרשות מעסיק קרן השתלמות (Employer Study Fund Contributions)
        - ניכוי עובד גמל (Employee Pension Deduction)
        - ניכוי עובד קרן התשלמות (Employee Study Fund Deduction)
        - תוספת וותק (Seniority Bonus)
        - ש.נ. גלובלי (Global Overtime)
        - כמות ש.נ גלובלי (Global Overtime Amount)
        - עמלה (Commission)
        - בונוס (Bonus)
        - עמלה (Commission)

        Return the result as a markdown table with columns for all the above fields (in the order listed). If there are multiple payslips or employees, include all rows.

        Payslip Text:
        {processed_result.get('payslip_text', '')}

        Contract Text:
        {processed_result.get('contract_text', '')}

        Attendance Text:
        {processed_result.get('attendance_text', '')}
        """

        # Simplified agent call
        async def get_table_sync():
            try:
                # For FastAPI/uvicorn - direct call (no nest_asyncio needed)
                if not self._is_streamlit():
                    return await self.agent.run(prompt, model_settings=ModelSettings(temperature=0.0))

                # For Streamlit - use nest_asyncio
                import nest_asyncio
                nest_asyncio.apply()
                loop = asyncio.get_event_loop()
                return loop.run_until_complete(self.agent.run(prompt, model_settings=ModelSettings(temperature=0.0)))
            except Exception as e:
                raise e

        result = await get_table_sync()
        print(result)
        table_md = result.output if hasattr(result, 'output') else str(result)

        # Parse the markdown table to a DataFrame
        import re
        from io import StringIO
        def markdown_table_to_df(md):
            # Find the first markdown table in the string
            lines = md.splitlines()
            table_lines = []
            in_table = False
            for line in lines:
                if '|' in line:
                    in_table = True
                    table_lines.append(line)
                elif in_table and line.strip() == '':
                    break
            if not table_lines:
                raise ValueError("No markdown table found in AI output.")
            # Remove markdown separator lines (e.g., | :--- | ... |)
            cleaned_lines = []
            for i, line in enumerate(table_lines):
                # Remove lines where all cells are dashes or colons (markdown separator)
                cells = [c.strip() for c in line.strip().split('|') if c.strip()]
                if all(re.match(r'^:?[-]+:?$', c) for c in cells):
                    continue
                cleaned_lines.append(line)
            table_str = '\n'.join(cleaned_lines)
            # Use pandas to read the markdown table
            try:
                import pandas as pd
                from io import StringIO
                df = pd.read_csv(StringIO(table_str), sep='|', engine='python')
                # Remove unnamed columns and whitespace
                df = df.loc[:, ~df.columns.str.contains('^Unnamed')]
                df.columns = [c.strip() for c in df.columns]
                # Use DataFrame.map for element-wise string strip (applymap is deprecated)
                for col in df.columns:
                    if df[col].dtype == object:
                        df[col] = df[col].map(lambda x: x.strip() if isinstance(x, str) else x)
                return df
            except Exception as e:
                raise ValueError(f"Failed to parse markdown table: {e}\nTable string:\n{table_str}")

        df = markdown_table_to_df(table_md)

        # Write DataFrame to Excel in memory
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name='EmployeeData')
        output.seek(0)
        return output.read()
    
    def _is_streamlit(self) -> bool:
        """Check if the code is running in a Streamlit environment."""
        try:
            import streamlit as st
            # Check if we're actually in a Streamlit runtime, not just if streamlit is installed
            return hasattr(st, 'runtime') and st.runtime.exists()
        except (ImportError, AttributeError):
            return False

    def qna_sync(self, report: str, questions: str) -> str:
        """Synchronous wrapper for the async qna method."""
        import asyncio
        try:
            return asyncio.run(self.qna(report, questions))
        except RuntimeError as e:
            # If there's already a running event loop (e.g. in Streamlit), use alternative
            if self._is_streamlit():
                try:
                    import nest_asyncio
                    nest_asyncio.apply()
                    loop = asyncio.get_event_loop()
                    return loop.run_until_complete(self.qna(report, questions))
                except Exception as inner_e:
                    raise RuntimeError(f"Failed to run qna async in Streamlit: {inner_e}") from e
            else:
                # In uvicorn/FastAPI, this shouldn't happen since we should use the async method directly
                raise RuntimeError(f"Cannot run async method in sync context. Use the async qna method instead: {e}") from e
            
    def __init__(self):
        # Initialize RAG storage
        self.rag_storage = RAGLegalStorage()
        
        # Initialize backward compatibility storages
        self.letter_format = LetterFormatStorage()
        # self.law_storage = LaborLawStorage()
        # self.judgement_storage = JudgementStorage()
        self.law_storage = self.rag_storage
        self.judgement_storage = self.rag_storage
        
        # Initialize PydanticAI model - use Google Gemini
        gemini_api_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GOOGLE_CLOUD_VISION_API_KEY")
        openai_api_key = os.getenv("OPENAI_API_KEY")
        # Use ModelSettings to set temperature to 0.0 (deterministic output)
        model_settings = ModelSettings(temperature=0.0)
        if openai_api_key:
            try:
                import openai
                openai.api_key = openai_api_key
                self.model_type = "gpt-5"
                self.model = "gpt-5"  # Use process_with_gpt5 for calls
            except ImportError:
                print("⚠️ openai package not installed. GPT-5 support unavailable.")
        # elif gemini_api_key:
        #     # Create Google provider with API key
        #     google_provider = GoogleProvider(api_key=gemini_api_key)
        #     self.model = GoogleModel('gemini-2.5-pro', provider=google_provider, settings=model_settings)
        #     self.model_type = "gemini"
        else:
            raise Exception("GOOGLE_API_KEY or GOOGLE_CLOUD_VISION_API_KEY must be set in environment variables")
        # Initialize PydanticAI Agent
        # Configure Vision API
        vision_api_key = os.getenv("GOOGLE_CLOUD_VISION_API_KEY")
        if not vision_api_key:
            raise Exception("Google Cloud Vision API key not found. Please set GOOGLE_CLOUD_VISION_API_KEY in your .env file")
        
        self.vision_client = ImageAnnotatorClient(client_options={"api_key": vision_api_key})
        self.image_context = {"language_hints": ["he"]} 
        # Only apply nest_asyncio if not running under uvloop (e.g., not under uvicorn)
        _can_patch = True
        try:
            import asyncio
            loop = asyncio.get_event_loop()
            if 'uvloop' in str(type(loop)):
                _can_patch = False
        except Exception:
            pass
        if _can_patch and os.environ.get("USE_NEST_ASYNCIO", "0") == "1":
            import nest_asyncio
            nest_asyncio.apply()
        self.agent = Agent(
            model=self.model,
            output_type=str,
            system_prompt="""You are an expert legal document analyzer specializing in Israeli labor law compliance.

You will be provided with:
1. Relevant labor laws retrieved from a legal database
2. Relevant legal judgements and precedents retrieved from a legal database  
3. Document content to analyze (payslips, contracts, attendance records)
4. Additional context if provided

Your analysis must be based STRICTLY on the provided laws and judgements. Do not use external legal knowledge.

🚫 VERY IMPORTANT:
- Read the document carefully and remember its contents like wage per hour, working hours, etc. and Do not recalculate data like wages per hour, sick days, etc. unless the document provides exact values.
- Do not infer or estimate violations without clear proof in the payslip.
- Use **only** the documents provided (e.g., payslip data, employment contracts data, and attendance records data). **Do not extract or reuse any example values (e.g., 6000 ₪, 186 hours, 14 hours overtime) that appear in the legal texts or examples.**
- Do **not invent** missing data. If the document does not include sufficient detail for a violation (e.g., no overtime hours), **do not report a violation**.
- Do not hallucinate sick days, overtime hours, or absences
- think step by step and analyze the documents carefully. do not rush to conclusions.
- while calculating read the whole law and dont miss anything and explain the calculations step by step and how you arrived at the final amounts.

Always respond in Hebrew and follow the specific formatting requirements for each analysis type.

CRITICAL: When providing analysis, do NOT output template text or placeholders. Always replace ALL placeholders with real data from the analysis.""",
        )

        # Only apply nest_asyncio for question_agent if not running under uvloop
        if _can_patch and os.environ.get("USE_NEST_ASYNCIO", "0") == "1":
            import nest_asyncio
            nest_asyncio.apply()
        self.question_agent = Agent(
            model=self.model,
            output_type=List[str],
            system_prompt="""You are an expert legal document analyzer specializing in Israeli labor law compliance.
                            Your task is to analyze document content and generate specific search queries to find relevant and most applicable laws and judgements.
                            break those questions into sub-questions that can be used to search for laws and judgements in the database.
                            Return a list of 3-5 specific search queries that focus on the key legal concepts present in the documents.
            """
        )

        @self.question_agent.system_prompt
        def dynamic_system_prompt():
            """Dynamic system prompt for question agent"""
            # Get all law summaries from RAG storage
            law_summaries = self.rag_storage.get_all_law_summaries()
            if law_summaries:
                print(f"Found {len(law_summaries)} law summaries in the database.")
                summaries_text = "\n".join([f"- {summary}" for summary in law_summaries])
                return f"""
Here is a list of summaries of laws available in the database that you can search for:

{summaries_text}

Based on the document content provided, generate 3-5 specific search queries that will help find the most relevant laws and judgements. Focus on key legal concepts like:
- Salary and wage-related issues
- Working hours and overtime
- Pension and social benefits
- Employment contracts and termination
- Worker rights violations

Return only the search queries as a list of strings.
"""
            else:
                print("no law summaries found in the database.")
                return """
You are an expert legal document analyzer specializing in Israeli labor law compliance.

No law summaries are currently available in the database. Analyze the document content and generate 3-5 specific search queries based on the general legal concepts present. Focus on key areas like wages, working hours, benefits, contracts, and worker rights.

Return only the search queries as a list of strings.
"""

        # --- Review Agent ---
        if _can_patch and os.environ.get("USE_NEST_ASYNCIO", "0") == "1":
            import nest_asyncio
            nest_asyncio.apply()
        self.review_agent = Agent(
            model=self.model,
            output_type=str,
            system_prompt="""
You are a legal analysis review agent specializing in Israeli labor law.
You are given:
1. A set of relevant labor laws (as text)
2. A set of relevant legal judgements (as text)
3. An analysis of a legal case (as text)
4. A set of relevant employee documents (as text)
5. Additional context or information (as text)
Your job is to:
  - Carefully check if the analysis is correct, complete, and strictly based on the provided laws and judgements.
  - Carefully check the documents and laws and judgements and see if the analysis is missing any important legal points or if it contains any errors or is missing any violation.
  - Carefully check if the analysis is based only on the provided laws and judgements and documents and not assuming any external knowledge or making up facts.
  - You must also carefully check that all calculations (amounts, sums, percentages, totals, etc.) are correct and match the provided laws, judgements, and document data. If you find any calculation errors, you must correct them and explain the correction.
  - If the analysis is correct, return it as-is.
  - If the analysis is incorrect, incomplete, or not strictly based on the provided laws and judgements, or if any calculation is wrong, generate a corrected analysis that is fully compliant and mathematically accurate.
  
  🚫 VERY IMPORTANT:
- Do not recalculate data like wages per hour, overtime hours, sick days, etc. unless the document provides exact values.
- Do not infer or estimate violations without clear proof in the payslip.
- Use **only** the documents provided (e.g., payslip data, employment contracts data, and attendance records data). **Do not extract or reuse any example values (e.g., 6000 ₪, 186 hours, 14 hours overtime) that appear in the legal texts or examples.**
- Do **not invent** missing data. If the document does not include sufficient detail for a violation (e.g., no overtime hours), **do not report a violation**.
- Do not hallucinate sick days, overtime hours, or absences
- think step by step and analyze the documents carefully. do not rush to conclusions.
- while calculating read the whole law and dont miss anything and explain the calculations step by step and how you arrived at the final amounts.

Always respond in Hebrew.
Do not use any external knowledge or make up facts.
Always cite the provided laws and judgements in your corrections.
Always check and correct all calculations.
"""
        )

        # Initialize Evaluator-Optimizer system
        try:
            self.evaluator_optimizer = EvaluatorOptimizer()
            print("✅ Evaluator-Optimizer system initialized successfully")
        except Exception as e:
            print(f"⚠️ Failed to initialize Evaluator-Optimizer: {e}")
            self.evaluator_optimizer = None
            
        # Initialize Rule Engine components
        try:
            from engine.loader import RuleLoader
            from engine.evaluator import RuleEvaluator
            from engine.penalty_calculator import PenaltyCalculator
            
            self.rule_loader = RuleLoader
            self.rule_evaluator = RuleEvaluator
            self.penalty_calculator = PenaltyCalculator
            
            # Load rules from the rules file
            rules_path = os.path.join(os.path.dirname(__file__), 'rules/labor_law_rules.json')
            self.rules_data = self.rule_loader.load_rules(rules_path)
            print("✅ Rule Engine initialized successfully")
        except Exception as e:
            print(f"⚠️ Failed to initialize Rule Engine: {e}")
            self.rule_loader = None
            self.rule_evaluator = None
            self.penalty_calculator = None
            self.rules_data = None
        

    async def review_analysis(self, laws: str, judgements: str, analysis: str, documents: Dict[str, str], context: str) -> str:
        """
        Review the analysis against the provided laws and judgements. If correct, return as-is. If not, return a corrected analysis.
        Also, carefully check all calculations (amounts, sums, percentages, totals, etc.) and correct any errors found.
        The output must never mention that it was revised, never explain what was fixed, and must simply return the corrected analysis as if it was always correct.
        """
        print("Reviewing analysis against provided laws and judgements...")
        prompt = f"""
הנך מקבל את החוקים, פסקי הדין, והניתוח המשפטי הבא:

📄 חוקים:
{laws}

📚 פסקי דין:
{judgements}

📑 ניתוח משפטי:
{analysis}

🧾 מסמכים שסופקו לבדיקה:
{ "\n\n".join([f"{doc_type.upper()} CONTENT:\n{content}" for doc_type, content in documents.items()]) }

🔍 הקשר נוסף:
{context}

🔍 הנחיות לבדיקה:

1. בדוק בקפידה האם הניתוח **נכון, שלם, ומבוסס אך ורק** על:
   - החוקים שסופקו
   - פסקי הדין שסופקו
   - המסמכים שסופקו (תלוש שכר, חוזה עבודה, דוח נוכחות)
2. אין להשתמש בידע כללי או חיצוני, ואין להניח מידע שאינו מופיע במפורש במסמכים.
3. אין להעתיק או לעשות שימוש בערכים מהחוקים כדוגמה (כגון 6000 ₪, 186 שעות, 5 ימי מחלה), אלא אם הם מופיעים במפורש במסמכים.
4. אם לא מצוינים נתונים מפורשים (כמו שעות נוספות, ימי מחלה, סכום שכר מינימום), יש להניח **שאין עבירה**, ואין לדווח על הפרה או לבצע חישובים משוערים.
5. כל החישובים (סכומים, אחוזים, טבלאות) חייבים להיות מדויקים ולבוסס אך ורק על הערכים המופיעים במסמכים שסופקו.
6. יש לוודא שהמסקנות המשפטיות תואמות את המסמכים ואת החוק, ללא שגיאות, וללא חריגות או תוספות לא מבוססות.

🛠 אם הניתוח תקין – החזר אותו כפי שהוא.  
🛠 אם יש בו שגיאות – תקן אותו כך שיהיה מדויק, מבוסס אך ורק על המסמכים, החוקים ופסקי הדין שסופקו.  


- לציין שבוצע תיקון.
- להסביר מה שונה.
- להזכיר שהניתוח תוקן או נערך מחדש.

✅ **יש להחזיר תמיד את הניתוח הסופי בלבד, בעברית מלאה וללא כל הערה.**

"""
        try:
            result = await self.review_agent.run(prompt,model_settings=ModelSettings(temperature=0.0))
            return result.output if hasattr(result, 'output') else str(result) 
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error in review_analysis: {get_error_detail(e)}")

    # async def create_report_with_optimization(self, payslip_text: str = None, contract_text: str = None, 
    #                       attendance_text: str = None, analysis_type: str = "combined", 
    #                       context: str = None, use_optimizer: bool = True, target_score: float = 0.85) -> Dict:
    #     """
    #     Create legal analysis report with evaluator-optimizer workflow for enhanced accuracy
        
    #     Args:
    #         payslip_text: Payslip document text
    #         contract_text: Contract document text  
    #         attendance_text: Attendance document text
    #         analysis_type: Type of analysis to perform
    #         context: Additional context
    #         use_optimizer: Whether to use the evaluator-optimizer workflow
    #         target_score: Target quality score for optimization (0-1)
            
    #     Returns:
    #         Dictionary with optimized analysis and quality metrics
    #     """
        
    #     # Step 1: Generate initial analysis and get all the processed data
    #     print("🔄 Generating initial analysis...")
    #     initial_result = await self.create_report(
    #         payslip_text=payslip_text,
    #         contract_text=contract_text,
    #         attendance_text=attendance_text,
    #         analysis_type=analysis_type,
    #         context=context
    #     )
        
    #     initial_analysis = initial_result.legal_analysis
    #     formatted_laws = initial_result.formatted_laws
    #     formatted_judgements = initial_result.formatted_judgements
    #     documents = initial_result.documents
        
    #     # If optimizer is disabled or not available, return initial result
    #     if not use_optimizer or not self.evaluator_optimizer:
    #         print("⚠️ Optimizer disabled or unavailable, returning initial analysis")
    #         return {
    #             'legal_analysis': initial_analysis,
    #             'status': 'success_no_optimization',
    #             'analysis_type': analysis_type,
    #             'quality_score': None,
    #             'improvements_made': [],
    #             'optimization_used': False
    #         }
        
    #     # Step 2: Prepare source documents for evaluation using data from create_report
    #     source_documents = {}
    #     if documents:
    #         for key, value in documents.items():
    #             if key == 'payslip':
    #                 source_documents['payslip_text'] = value
    #             elif key == 'contract':
    #                 source_documents['contract_text'] = value
    #             elif key == 'attendance report':
    #                 source_documents['attendance_text'] = value
    #             else:
    #                 source_documents[key] = value
    #     else:
    #         # Fallback to creating from parameters
    #         if payslip_text:
    #             source_documents['payslip_text'] = payslip_text
    #         if contract_text:
    #             source_documents['contract_text'] = contract_text
    #         if attendance_text:
    #             source_documents['attendance_text'] = attendance_text
        
    #     # Step 3: Use the formatted laws and judgements from create_report (no need to retrieve again!)
    #     if not formatted_laws or not formatted_judgements:
    #         print("⚠️ No formatted laws/judgements from create_report, using fallback")
    #         formatted_laws = formatted_laws or "No laws retrieved for evaluation"
    #         formatted_judgements = formatted_judgements or "No judgements retrieved for evaluation"
    #     else:
    #         print("✅ Using formatted laws and judgements from create_report")
        
    #     # Step 4: Run evaluator-optimizer workflow
    #     try:
    #         print("🚀 Running evaluator-optimizer workflow...")
            
    #         optimized_analysis, final_evaluation, changes = await self.evaluator_optimizer.evaluate_and_optimize(
    #             analysis=initial_analysis,
    #             source_documents=source_documents,
    #             laws=formatted_laws,
    #             judgements=formatted_judgements,
    #             context=context or "",
    #             target_score=target_score
    #         )
            
    #         print(f"✅ Optimization complete! Final score: {final_evaluation.overall_score:.2f}")
            
    #         return {
    #             'legal_analysis': optimized_analysis,
    #             'status': 'success_optimized',
    #             'analysis_type': analysis_type,
    #             'quality_score': final_evaluation.overall_score,
    #             'detailed_scores': {
    #                 'factual_accuracy': final_evaluation.factual_accuracy,
    #                 'legal_compliance': final_evaluation.legal_compliance,
    #                 'calculation_accuracy': final_evaluation.calculation_accuracy,
    #                 'completeness': final_evaluation.completeness,
    #                 'citation_quality': final_evaluation.citation_quality,
    #                 'clarity': final_evaluation.clarity
    #             },
    #             'improvements_made': changes,
    #             'remaining_suggestions': final_evaluation.improvement_suggestions,
    #             'critical_issues': final_evaluation.critical_issues,
    #             'optimization_used': True
    #         }
            
    #     except Exception as e:
    #         print(f"⚠️ Evaluator-optimizer failed: {e}")
    #         # Fallback to initial analysis if optimization fails
    #         return {
    #             'legal_analysis': initial_analysis,
    #             'status': 'success_optimization_failed',
    #             'analysis_type': analysis_type,
    #             'quality_score': None,
    #             'improvements_made': [],
    #             'optimization_used': False,
    #             'optimization_error': str(e)
    #         } 

    async def process_document(self, files: Union[UploadFile, List[UploadFile]], doc_types: Union[str, List[str]], compress: bool = False) -> Dict[str, str]:
        """Process uploaded documents and extract text concurrently"""
        if not files:
            raise HTTPException(
                status_code=400,
                detail="At least one document must be provided"
            )
        # Handle single file case
        if isinstance(files, UploadFile):
            files = [files]
            doc_types = [doc_types]
        elif not isinstance(doc_types, list):
            doc_types = [doc_types] * len(files)

        async def process_single(file, doc_type, idx):
            print(f"Processing file: {file.filename} as type: {doc_type}")
            content = await file.read()
            extracted_text = await asyncio.to_thread(self._extract_text2, content, file.filename, compress)
            return (doc_type.lower(), idx, extracted_text)

        tasks = [process_single(file, doc_type, idx+1) for idx, (file, doc_type) in enumerate(zip(files, doc_types))]
        results = await asyncio.gather(*tasks)

        payslip_text = None
        contract_text = None
        attendance_text = None
        
        # Initialize separate counters for each document type
        payslip_counter = 0
        contract_counter = 0
        attendance_counter = 0

        for doc_type, _, extracted_text in results:
            if doc_type == "payslip":
                payslip_counter += 1
                payslip_text = f"Payslip {payslip_counter}:\n{extracted_text}" if payslip_text is None else f"{payslip_text}\n\nPayslip {payslip_counter}:\n{extracted_text}"
            elif doc_type == "contract":
                contract_counter += 1
                contract_text = f"Contract {contract_counter}:\n{extracted_text}" if contract_text is None else f"{contract_text}\n\nContract {contract_counter}:\n{extracted_text}"
            elif doc_type == "attendance":
                attendance_counter += 1
                attendance_text = f"Attendance {attendance_counter}:\n{extracted_text}" if attendance_text is None else f"{attendance_text}\n\nAttendance {attendance_counter}:\n{extracted_text}"

        # return {
        #     "payslip_text": payslip_text,
        #     "contract_text": contract_text,
        #     "attendance_text": attendance_text
        # }

        return {
            "payslip_text": """Payslip 1:
1/21  תלוש משכורת <!-- marginalia, from page 0 (l=0.316,t=0.015,r=0.523,b=0.036), with ID 74a17059-921b-4df4-9428-8c85894e5ff3 -->

ת.אילת השקעות וניהול פרויקטים בע"מ  
ירושלים השלמה 35 אילת 88000  
טל: 08-6333399 פקס: <!-- text, from page 0 (l=0.603,t=0.023,r=0.917,b=0.071), with ID 85aaa8d3-8d9f-4a3b-a27d-830afa295d06 -->

95009327800
950093278
514256429

תיק ג ב"ל
תיק נכיינים
תיק במי"ה <!-- text, from page 0 (l=0.028,t=0.032,r=0.196,b=0.070), with ID 8dc08c7a-06c0-44e9-9201-ea39407833b3 -->

<table><tr><td>01/07/2018</td><td>תחילת עבודה:</td><td>בנק/סניף:</td><td>מס' עובד:260</td><td>בנאי אברהם</td></tr><tr><td>01/07/2018</td><td>תאריך ותק:</td><td>חשבון:</td><td>ת.ז: 05498201-2</td><td>312</td></tr><tr><td>2.59</td><td>שנות ותק:</td><td></td><td>מחלקה: 6-חברה ת.אילת</td><td>3</td></tr><tr><td>105.00 ש' עבודה: מתוך: 182.00</td><td></td><td></td><td></td><td>אילת</td></tr><tr><td>22.00 ימי עבודה: מתוך: 26.00</td><td></td><td>דרגה:</td><td></td><td>88000</td></tr><tr><td></td><td></td><td>דרוג:</td><td></td><td></td></tr><tr><td>תע' שנה: 35.00</td><td>היקף משרה: %</td><td></td><td>משרה: לפי שעות</td><td></td></tr></table> <!-- table, from page 0 (l=0.028,t=0.071,r=0.917,b=0.187), with ID b1862545-c7ea-4e2a-8389-93a90e5a96f4 -->

<table><thead><tr><th>תיאור תשלום</th><th>כמות</th><th>תעריף</th><th>שווי למס</th><th>גילום</th><th>סכום לתשלום</th><th>ניכויים</th><th>סכום</th></tr></thead><tbody><tr><td>שכר יסוד</td><td></td><td>35.00</td><td></td><td></td><td>3,675.00</td><td>מס הכנסה שולי 14%</td><td>158.00</td></tr><tr><td>נסיעות</td><td></td><td>116.00</td><td></td><td></td><td>116.00</td><td>ביטוח לאומי</td><td></td></tr><tr><td>שעות נוספות 125%</td><td>15.00</td><td>43.75</td><td></td><td></td><td>656.25</td><td>מס בריאות</td><td>291.00</td></tr><tr><td>שעות נוספות 150%</td><td>72.00</td><td>52.50</td><td></td><td></td><td>3,780.00</td><td></td><td></td></tr><tr><td colspan="8"></td></tr><tr><td colspan="6"></td><td>ניכויי רשות</td><td>סכום</td></tr><tr><td colspan="6"></td><td>קניות</td><td>975.00</td></tr></tbody></table> <!-- table, from page 0 (l=0.040,t=0.186,r=0.913,b=0.515), with ID 56278ee9-521c-4fce-9268-c5a5572689ea -->

הפקדות מעסיק לקופות גמל
<table><thead><tr><th>שונות</th><th>א. כושר</th><th>פיצויים</th><th>תגמולים</th><th>שכר לגמל</th><th>שם הקופה</th></tr></thead><tbody><tr><td></td><td></td><td></td><td></td><td></td><td></td></tr></tbody></table> <!-- table, from page 0 (l=0.557,t=0.514,r=0.908,b=0.619), with ID 2c71ae73-f097-49f1-b41c-5fd3086271f8 -->

הפקדות מעסיק לפיצויים
<table><thead><tr><th>מצטבר</th><th>חודשי</th><th>הפקדות מעסיק לפיצויים</th></tr></thead><tbody><tr><td>0.00</td><td>0.00</td><td>משכורת מבוטחת</td></tr><tr><td>0.00</td><td>0.00</td><td>שווי פיצויים</td></tr><tr><td>0.00</td><td>0.00</td><td>פיצויים ללא שווי</td></tr><tr><td>0.00</td><td>0.00</td><td>הפקדה לקרן ותיקה</td></tr></tbody></table> <!-- table, from page 0 (l=0.287,t=0.514,r=0.552,b=0.618), with ID 28c58631-b7bb-4d69-8306-6902ca6e7727 -->

<table><tr><td>8,227.25</td><td>סה"כ תשלומים</td></tr><tr><td>449.00</td><td>ניכויי חובה וגמל</td></tr><tr><td>7,778.25</td><td>שכר נטו</td></tr><tr><td>975.00</td><td>ס"ה ניכויי רשות</td></tr><tr><td>6,803.25</td><td>נטו לתשלום</td></tr></table> <!-- table, from page 0 (l=0.059,t=0.517,r=0.274,b=0.598), with ID 4fec36c3-121e-4a90-a141-1a6971be9a12 -->

<table><thead><tr><th>מצטבר</th><th>חודשי</th><th>תשלומים שווי למס שווי גמל שווי ק"ה חייב מ"ה פטור מ"ה חייב ל"ל ערך ז"ז זיכוי נוסף זיכוי גמל ניכוי ס' 47 הנחת י"פ זיכוי משמרות פטור מיוחד</th></tr></thead><tbody><tr><td>8,227.00</td><td>8,227.25</td><td></td></tr><tr><td>8,227.00</td><td>8,227.25</td><td></td></tr><tr><td>8,227.25</td><td>8,227.25</td><td></td></tr><tr><td>490.50</td><td>490.50</td><td></td></tr><tr><td></td><td></td><td></td></tr><tr><td>823.00</td><td>822.73</td><td></td></tr></tbody></table> <!-- table, from page 0 (l=0.570,t=0.618,r=0.915,b=0.806), with ID 77e34e92-6ec5-432a-a413-fc7ed728b1cb -->

<table><thead><tr><th>ניכוי"ם</th><th>מצטבר</th></tr></thead><tbody><tr><td>מס הכנסה<br>ביטוח לאומי<br>בריאות<br>גמל 25%<br>גמל 35%<br>ב. מנהלים<br>ק. השתלמות<br>סעיף 47<br>יתרת הלוואה</td><td>158.00<br>291.00</td></tr><tr><td colspan="2">מצב משפחתי גרוש/ה</td></tr><tr><td>נ"ז<br>% מס קבוע<br>% הנחת י"פ<br>תנאום מס</td><td>2.25<br><br>10.00</td></tr></tbody></table> <!-- table, from page 0 (l=0.378,t=0.619,r=0.570,b=0.804), with ID f8703e3f-12cc-4e02-b43f-16dba39e87d6 -->

העדרויות
<table><thead><tr><th>תאור</th><th>קודמת</th><th>תוספת</th><th>ניצול</th><th>יתרה</th></tr></thead><tbody><tr><td>חופשה</td><td>9.05</td><td>1.17</td><td></td><td>10.21</td></tr><tr><td>מחלה</td><td>45.00</td><td>1.50</td><td></td><td>46.50</td></tr><tr><td>מילואים</td><td></td><td></td><td></td><td></td></tr><tr><td>הבראה</td><td>3.00</td><td>0.50</td><td></td><td>3.50</td></tr></tbody></table> <!-- table, from page 0 (l=0.028,t=0.614,r=0.369,b=0.713), with ID ea54699f-1aac-45c9-b457-9ff792940d60 -->

<table>
  <tr>
    <td>חודשי עבודה</td>
    <td>חודשי עבודה</td>
    <td>מספר חודשי עבודה מחודש</td>
    <td>חישוב מצטבר מחודש</td>
  </tr>
  <tr>
    <td>1</td>
    <td>1__________</td>
    <td>1</td>
    <td>0</td>
  </tr>
</table> <!-- table, from page 0 (l=0.031,t=0.733,r=0.369,b=0.805), with ID 48f0fb4d-25cb-4418-9d10-146fc2e8c7cd -->

הודעות:

צורת תשלום:  במזומן/בהמחאה

"למען מיצוי זכויותיך בביטוח הלאומי נא לוודא שפרטיך האישיים הרשומים בתלוש השכר זהים לפרטי תעודת הזהות."

שכר מינימום לחודש: 5,300.00 שכר מינימום לשעה: 29.12

חתימה <!-- text, from page 0 (l=0.031,t=0.804,r=0.913,b=0.935), with ID 551c6ee2-7100-4057-9f75-2c2516956ed0 -->

תאריך הדפסה 06/02/21 <!-- marginalia, from page 0 (l=0.759,t=0.936,r=0.906,b=0.952), with ID ea681fd8-ce92-41d7-aefd-e1174821ff8e -->

הופק ע"י: הכוכב שירותי הבה"ש באמצעות "מערכת טריו" של עוקץ מערכות בע"מ <!-- marginalia, from page 0 (l=0.073,t=0.933,r=0.533,b=0.949), with ID 02bce43d-8ecc-4aee-8d8a-e0ae721e05ab -->""",
            "contract_text": "",
            "attendance_text": """Attendance 1:
יום ראשון  
יום שני  
יום שלישי  
יום רביעי  
יום חמישי  
יום שישי  
יום שבת  
יום ראשון  
יום שני  
יום שלישי  
יום רביעי  
יום חמישי  
יום שישי  
יום שבת  
יום ראשון  
יום שני  
יום שלישי  
יום רביעי  
יום חמישי  
יום שישי  
יום שבת  
יום ראשון  
יום שני  
יום שלישי  
יום רביעי  
יום חמישי  
יום שישי  
יום שבת  
יום ראשון  
יום שני  
יום שלישי  
יום רביעי  
יום חמישי  
יום שישי  
יום שבת  
יום ראשון <!-- text, from page 0 (l=0.016,t=0.047,r=0.222,b=0.881), with ID 4eaefa70-ac65-41fe-ad9c-a9e1b4ba3f91 -->

1/2021 <!-- text, from page 0 (l=0.374,t=0.000,r=0.549,b=0.050), with ID a1cfc528-03bb-4213-bc27-34be5acac584 -->

21 <!-- text, from page 0 (l=0.667,t=0.004,r=0.733,b=0.042), with ID c14343a5-dc09-4766-8a14-9d8998a6f495 -->

7k1 9l <!-- text, from page 0 (l=0.769,t=0.002,r=0.897,b=0.042), with ID 63f94d8f-8e42-41ac-b1dc-8dcf6aa61395 -->

<table>
  <tr>
    <td>1/1</td>
    <td>תפוחי עץ</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>2/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>3/1</td>
    <td>חובב חובב חובב חובב</td>
    <td>חובב</td>
    <td>חובב</td>
    <td>24°°</td>
    <td>08°°</td>
  </tr>
  <tr>
    <td>4/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>5/1</td>
    <td>חובב חובב חובב חובב חובב</td>
    <td>חובב</td>
    <td>חובב</td>
    <td>24°°</td>
    <td>08°°</td>
  </tr>
  <tr>
    <td>6/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>7/1</td>
    <td>שזיף צהוב</td>
    <td>49</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>8/1</td>
    <td>שזיף צהוב</td>
    <td>49</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>9/1</td>
    <td>שזיף צהוב</td>
    <td>49</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>10/1</td>
    <td>חובב חובב חובב חובב</td>
    <td>חובב</td>
    <td>חובב</td>
    <td>24°°</td>
    <td>08°°</td>
  </tr>
  <tr>
    <td>11/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>12/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>13/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>14/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>15/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>16/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>17/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>18/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>19/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>20/1</td>
    <td>חובב חובב חובב חובב</td>
    <td>חובב</td>
    <td>חובב</td>
    <td>24°°</td>
    <td>08°°</td>
  </tr>
  <tr>
    <td>21/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>22/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>23/1</td>
    <td>חובב חובב חובב חובב חובב</td>
    <td>חובב</td>
    <td>חובב</td>
    <td>24°°</td>
    <td>08°°</td>
  </tr>
  <tr>
    <td>24/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>25/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>26/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>27/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>28/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>29/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>30/1</td>
    <td>חובב חובב חובב חובב</td>
    <td>חובב</td>
    <td>חובב</td>
    <td>24°°</td>
    <td>08°°</td>
  </tr>
  <tr>
    <td>31/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>32/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
  <tr>
    <td>33/1</td>
    <td>חובב חובב חובב חובב חובב</td>
    <td>חובב</td>
    <td>חובב</td>
    <td>24°°</td>
    <td>08°°</td>
  </tr>
  <tr>
    <td>34/1</td>
    <td>שזיף צהוב</td>
    <td>48</td>
    <td>08°°</td>
    <td>24°°</td>
  </tr>
</table> <!-- table, from page 0 (l=0.207,t=0.037,r=0.989,b=0.863), with ID 8fd127fd-d9d9-404b-92ac-aeccc59109e3 -->

ס"הכ  שינויים  199  שעות  
עלול  להשתנות  ולעבור  ניסיונות <!-- text, from page 0 (l=0.292,t=0.861,r=0.892,b=0.957), with ID 573d2f67-6344-4b0d-9198-7b23c02b38ef -->"""
        }

    async def _build_summary_prompt_from_combined(self, combined_report: str, analysis_type: str) -> str:
        """
        Build a prompt to convert the reviewed combined report into the requested summary type.
        """
        if analysis_type == "table":
            instructions = self._get_table_instructions()
        elif analysis_type == "violation_count_table":
            instructions = self._get_violation_count_table_instructions()
        elif analysis_type == "violations_list":
            instructions = self._get_violations_list_instructions()
        elif analysis_type == "easy":
            instructions = self._get_easy_instructions()
        elif analysis_type == "claim":
            instructions = self._get_claim_instructions()
        elif analysis_type == "warning_letter":
            instructions = self._get_warning_letter_instructions()
        else:
            raise ValueError(f"Unsupported summary analysis_type: {analysis_type}")

        prompt = f"""
להלן דוח ניתוח משפטי משולב של מסמכי העובד:

{combined_report}

---

{instructions}
"""
        return prompt

    async def qna(self, report: str, questions: str) -> str:
        """Generate answer of queries based on the provided document content."""
        prompt = f"""
להלן דוח ניתוח משפטי משולב של מסמכי העובד:

{report}

---

שאלות:
{questions}

---

אנא ספק תשובות לשאלות לעיל.
"""
        result = await self.agent.run(prompt, model_settings=ModelSettings(temperature=0.0))
        return result.output

    async def fix_ocr_content(self, ocr_content: str) -> str:
        """Fix and rearrange OCR content from payslips, attendance sheets, and contracts to improve readability and structure."""
        prompt = f"""
You are an expert document processing assistant specializing in Israeli employment documents. Your task is to fix and rearrange OCR-extracted content from payslips, attendance sheets, and employment contracts to make them more readable and properly structured.

🚫 CRITICAL RULES - DO NOT VIOLATE:
1. DO NOT add, invent, or hallucinate ANY information that is not present in the original OCR content
2. DO NOT remove or omit ANY information from the original content
3. DO NOT change numbers, dates, names, or any factual data
4. ONLY rearrange and fix formatting issues - preserve ALL original content exactly as it appears
5. If text is unclear or garbled, keep it as-is rather than guessing what it should be

The OCR content is from employment documents and may have issues such as:
- Misaligned text and spacing in tables
- Broken words split across lines
- Incorrect line breaks in salary/payment information
- Mixed up columns in payslip tables
- Garbled characters in Hebrew/English text
- Scattered table data that should be aligned

OCR Content to Fix:
{ocr_content}

Instructions for Rearrangement:
1. Fix obvious OCR spacing and alignment issues ONLY
2. If there are tables (salary details, attendance records), align them properly
3. Group related information together (employee details, salary breakdown, deductions, etc.)
4. Maintain proper Hebrew text direction if present
5. Preserve ALL numbers, amounts, dates, and names exactly as they appear
6. Keep all original information - just make it more readable
7. If content appears to be a payslip, organize sections like: Employee Info, Salary Details, Deductions, Net Pay
8. If content appears to be attendance, organize by dates, hours, etc.
9. If content appears to be a contract, maintain clause structure

🔍 Remember: You are ONLY fixing formatting and alignment - NOT interpreting or changing content.

Return only the cleaned and restructured content with NO explanations, comments, or additions.
"""
        
        try:
            result = await self.agent.run(prompt, model_settings=ModelSettings(temperature=0.0))
            return result.output if hasattr(result, 'output') else str(result)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error fixing OCR content: {get_error_detail(e)}")

    async def create_report(self, payslip_text: str = None, contract_text: str = None, 
                          attendance_text: str = None, analysis_type: str = "report", 
                          context: str = None) -> DocumentAnalysisResponse:
        """Create legal analysis report using RAG and PydanticAI"""

        # Prepare documents for analysis
        documents = {}
        if payslip_text:
            documents['payslip'] = payslip_text
        if contract_text:
            documents['contract'] = contract_text
        if attendance_text:
            documents['attendance report'] = attendance_text

        # Special handling for table, violation_count_table, violations_list, easy, claim, warning_letter
        special_types = ["table", "violation_count_table", "violations_list", "easy", "claim", "warning_letter"]
        if analysis_type in special_types:
            # Step 1: Run combined analysis and review
            combined_report = await self.create_report(
                payslip_text=payslip_text,
                contract_text=contract_text,
                attendance_text=attendance_text,
                analysis_type="combined",
                context=context
            )
            reviewed_combined = combined_report.legal_analysis
            print("Review Combined Report:")
            print(reviewed_combined)
            # Step 2: Build a summary prompt from the reviewed combined report
            prompt = await self._build_summary_prompt_from_combined(reviewed_combined, analysis_type)
            try:
                result = await self.agent.run(prompt, model_settings=ModelSettings(temperature=0.0))
                analysis = result.output
            except Exception as pydantic_error:
                raise pydantic_error
            # No review needed for these summary types
            return DocumentAnalysisResponse(
                legal_analysis=analysis,
                status="success",
                analysis_type=analysis_type,
                formatted_laws=combined_report.formatted_laws,
                formatted_judgements=combined_report.formatted_judgements,
                documents=combined_report.documents
            )

        # Normal flow for all other types
        # Pass the documents to the question agent to generate search queries
        docs = "\n\n".join([f"{doc_type.upper()} CONTENT:\n{content}" for doc_type, content in documents.items()])
        try:
            question_result = await self.question_agent.run(docs)
        except Exception as e:
            print(f"Error generating search queries: {get_error_detail(e)}")
            raise HTTPException(status_code=500, detail=f"Error generating search queries: {get_error_detail(e)}")

        # Extract the search queries from the result object
        search_queries = question_result.output if hasattr(question_result, 'output') else question_result
        
        print(f"Generated search queries: {search_queries}")

        # Query laws and judgements using the generated search queries
        all_relevant_laws = []
        all_relevant_judgements = []

        for query in search_queries:
            print(f"Searching for laws with query: {query}")
            laws = self.rag_storage.search_laws(query, n_results=2)  # Fewer results per query
            all_relevant_laws.extend(laws)

            print(f"Searching for judgements with query: {query}")
            judgements = self.rag_storage.search_judgements(query, n_results=1)  # Fewer results per query
            all_relevant_judgements.extend(judgements)

        # Remove duplicates based on ID or content
        unique_laws = []
        seen_law_ids = set()
        for law in all_relevant_laws:
            law_id = law.get('id') or law.get('metadata', {}).get('id')
            if law_id and law_id not in seen_law_ids:
                unique_laws.append(law)
                seen_law_ids.add(law_id)
            elif not law_id and law not in unique_laws:  # Fallback for items without ID
                unique_laws.append(law)

        unique_judgements = []
        seen_judgement_ids = set()
        for judgement in all_relevant_judgements:
            judgement_id = judgement.get('id') or judgement.get('metadata', {}).get('id')
            if judgement_id and judgement_id not in seen_judgement_ids:
                unique_judgements.append(judgement)
                seen_judgement_ids.add(judgement_id)
            elif not judgement_id and judgement not in unique_judgements:  # Fallback for items without ID
                unique_judgements.append(judgement)

        # Format the laws and judgements for the prompt
        formatted_laws = self.rag_storage.format_laws_for_prompt(unique_laws)
        formatted_judgements = self.rag_storage.format_judgements_for_prompt(unique_judgements)

        # Combine formatted content
        combined_context = f"{formatted_laws}\n\n{formatted_judgements}"
        print(f"Retrieved {len(unique_laws)} laws and {len(unique_judgements)} judgements")

        print(combined_context)
        # Build the analysis prompt based on type
        prompt = await self._build_analysis_prompt(
            analysis_type, documents, combined_context, context
        )

        try:
            # Generate analysis using PydanticAI
            try:
                result = await self.agent.run(prompt,model_settings=ModelSettings(temperature=0.0))
                analysis = result.output
            except Exception as pydantic_error:
                    raise pydantic_error
            print("Analysis generated successfully.")
            print(f"Analysis type: {analysis}")
            # --- Review the analysis for legal and calculation correctness ---
            # try:
            #     reviewed_analysis = await self.review_analysis(formatted_laws, formatted_judgements, analysis, documents,context)
            #     print("Review agent completed successfully.")
            # except Exception as review_error:
            #     print(f"Review agent failed: {review_error}")
            #     reviewed_analysis = analysis  # fallback to original if review fails

            return DocumentAnalysisResponse(
                legal_analysis=analysis,
                formatted_laws=formatted_laws,
                formatted_judgements=formatted_judgements,
                documents=documents,
                status="success",
                analysis_type=analysis_type
            )
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error generating legal analysis: {get_error_detail(e)}"
            )

    async def _build_analysis_prompt(self, analysis_type: str, documents: Dict, 
                                   laws_and_judgements: str, 
                                   context: str = None) -> str:
        """Build analysis prompt based on type"""
        
        base_prompt = f"""
RELEVANT LABOR LAWS and JUDGEMENTS (Retrieved from legal database):
{laws_and_judgements}


DOCUMENTS PROVIDED FOR ANALYSIS:
{', '.join(documents.keys())}

ADDITIONAL CONTEXT:
{context if context else 'No additional context provided.'}
"""
        # Add document contents to prompt
        for doc_type, content in documents.items():
            base_prompt += f"\n{doc_type.upper()} CONTENT:\n{content}\n"
        
        # Add specific instructions based on analysis type
        if analysis_type == 'report':
            base_prompt += self._get_report_instructions()
        elif analysis_type == 'profitability':
            base_prompt += self._get_profitability_instructions()
        # elif analysis_type == 'professional':
        #     base_prompt += self._get_professional_instructions()
        elif analysis_type == 'warning_letter':
            base_prompt += self._get_warning_letter_instructions()
        elif analysis_type == 'easy':
            base_prompt += self._get_easy_instructions()
        elif analysis_type == 'table':
            base_prompt += self._get_table_instructions()
        elif analysis_type == 'claim':
            base_prompt += self._get_claim_instructions()
        elif analysis_type == 'combined':
            base_prompt += self._get_combined_instructions()
        elif analysis_type == 'violation_count_table':
            base_prompt += self._get_violation_count_table_instructions()
        elif analysis_type == 'violations_list':
            base_prompt += self._get_violations_list_instructions()
        
        return base_prompt

    def _get_report_instructions(self) -> str:
        return """
INSTRUCTIONS:
1. If no labor laws are provided, respond with: "אין חוקים לעבודה זמינים לניתוח התאמה." in Hebrew.
2. If labor laws exist, analyze the documents ONLY against the provided laws.
3. ONLY refer to the judgements and their results provided above for legal analysis - do not use external cases or knowledge.
4. If no judgements are provided, respond with: "לא קיימות החלטות משפטיות זמינות לניתוח." in Hebrew.

🎯 GOAL: Help the business owner understand exactly what went wrong and where they made mistakes in their employment practices.

For each payslip provided, analyze and identify violations. For each violation found in each payslip, format the response EXACTLY as shown below, with each section on a new line and proper spacing:

Violation Format Template:

🚨 [VIOLATION TITLE - What the employer did wrong]

📋 מה קרה בפועל:
[Describe exactly what the employer did or failed to do, with specific details from the documents]

⚖️ מה היה צריך לקרות לפי החוק:
[Explain what should have been done according to the law, with specific legal requirements]

📖 בסיס חוקי:
[LAW REFERENCE AND YEAR FROM PROVIDED LAWS - cite the specific law that was violated]

💰 הנזק הכספי:
[Calculate the exact financial impact - how much the employee lost due to this violation]

🏛️ תקדימים משפטיים:
[SIMILAR CASES OR PRECEDENTS FROM PROVIDED JUDGEMENTS](Refer *only* to the retrieved judgements. If a relevant judgement is found, describe the case and its result. If no relevant judgement is found, state "לא נמצאו תקדימים רלוונטיים בפסקי הדין שסופקו.")

⚠️ השלכות אפשריות:
[Explain potential legal consequences and risks for the employer]

✅ מה לעשות כדי לתקן:
[Specific actionable steps the employer should take to fix this violation and prevent it in the future]

---

SUMMARY FOR EMPLOYER:
After completing the violation analysis, provide a clear summary for the business owner:

=== סיכום למעסיק - איפה טעיתם ומה לעשות ===

📊 טבלת הפרות שזוהו:
תלוש/מסמך | מה עשיתם לא נכון | כמה זה עולה (₪) | מה לעשות עכשיו
[Add rows with actual data showing: document name | specific mistake made | financial cost | corrective action needed]

💸 סה"כ עלות הטעויות: [total amount] ₪

🔧 צעדים מיידיים לתיקון:
1. [First immediate action needed]
2. [Second immediate action needed]
3. [Third immediate action needed]

📋 איך למנוע טעויות בעתיד:
• [Prevention measure 1]
• [Prevention measure 2]
• [Prevention measure 3]

IMPORTANT:
- Always Respond in Hebrew
- Focus on helping the employer understand their mistakes and how to fix them
- Use clear, business-friendly language that explains the "why" behind each violation
- Format each violation with proper spacing and line breaks as shown above
- Analyze each payslip separately and clearly indicate which payslip the violations belong to
- Separate multiple violations with '---'
- If no violations are found against the provided laws in a payslip, respond with: "לא נמצאו הפרות בתלוש מספר" followed by the payslip number in hebrew
- If no violations are found in any payslip, respond with: "לא נמצאו הפרות נגד חוקי העבודה שסופקו." in hebrew
- DO NOT output template text or placeholders - use real data from the analysis
- Replace ALL placeholders with actual information from the documents
- Make it clear to the employer what they did wrong and how to prevent it happening again
"""

    def _get_violations_list_instructions(self) -> str:
        return """
CRITICAL INSTRUCTIONS - FOLLOW EXACTLY:

You MUST respond with ONLY the format below. NO additional text, explanations, analysis, or commentary.

If violations found, use EXACTLY this format:

**הפרות שזוהו:**

• [הפרה קצרה] - [חוק]
• [הפרה קצרה] - [חוק]
• [הפרה קצרה] - [חוק]

If NO violations found, respond ONLY with: "לא נמצאו הפרות"

ABSOLUTE RESTRICTIONS:
❌ NO "Legal analysis" headers
❌ NO explanations or breakdowns  
❌ NO calculations or formulas
❌ NO "Analysis of the violation" sections
❌ NO "method of calculating" text
❌ NO English words
❌ NO additional paragraphs
❌ NO template text like [Violation Title]

✅ ONLY Hebrew
✅ ONLY the exact format above
✅ Maximum 4 violations
✅ Each violation = ONE short line

EXAMPLE of correct output:
**הפרות שזוהו:**

• לא שולם שכר מינימום - חוק שכר מינימום
• לא שולמו שעות נוספות - חוק שעות עבודה ומנוחה
• לא הופקדה פנסיה - צו הרחבה פנסיה

If you write ANYTHING beyond this format, you FAILED.
"""

    def _get_profitability_instructions(self) -> str:
        return """
INSTRUCTIONS:
1. Analyze the provided documents and identify potential labor law violations, based *exclusively on the retrieved LABOR LAWS and JUDGEMENTS*.
2. For each violation, refer *only* to the retrieved judgements to identify similar legal cases and their outcomes (both successful and unsuccessful).
3. If similar cases (from the retrieved judgements) were unsuccessful:
   - Explain why the cases were unsuccessful based on the retrieved judgement information.
   - Provide a clear recommendation against pursuing legal action based on this.
   - List potential risks and costs.

4. If similar cases (from the retrieved judgements) were successful, calculate using information from those judgements if available:
   - Average compensation amount from successful retrieved cases.
   - Estimated legal fees (30% of potential compensation).
   - Tax implications (25% of net compensation).
   - Time and effort cost estimation.

Provide the analysis in the following format:

ניתוח כדאיות כלכלית:

הפרות שזוהו (על בסיס החוקים שסופקו):
[List identified violations]

תקדימים משפטיים (מתוך פסקי הדין שסופקו):
[Similar cases from retrieved judgements with outcomes - both successful and unsuccessful]

במקרה של תקדימים שליליים (מתוך פסקי הדין שסופקו):
- סיבות לדחיית התביעות: [REASONS BASED ON RETRIEVED JUDGEMENTS]
- סיכונים אפשריים: [RISKS]
- המלצה: לא מומלץ להגיש תביעה בשל [EXPLANATION BASED ON RETRIEVED JUDGEMENTS]

במקרה של תקדימים חיוביים (מתוך פסקי הדין שסופקו):
ניתוח כספי:
- סכום פיצוי ממוצע (מפסיקות דין שסופקו): [AMOUNT] ₪
- עלות משוערת של עורך דין (30%): [AMOUNT] ₪
- השלכות מס (25% מהסכום נטו): [AMOUNT] ₪
- סכום נטו משוער: [AMOUNT] ₪

המלצה סופית:
[Based on analysis of both successful and unsuccessful cases from the retrieved judgements, provide clear recommendation]

SUMMARY TABLE:
After completing the profitability analysis, provide a summary table:
- Use the heading: === טבלת סיכום כלכלי ===
- Create columns for: פריט | סכום (₪)
- Add rows with actual calculated amounts for:
  * סה"כ פיצוי משוער
  * עלות עורך דין (30%)
  * מס (25%)
- End with: סכום נטו משוער: [calculated amount]

CRITICAL: Replace ALL placeholders with actual calculated values from the analysis. Do not output template text.
""""""
"""

#     def _get_professional_instructions(self) -> str:
#         return """
# Analyze the provided documents for labor law violations based strictly on the retrieved Israeli labor laws and the content of the documents. For each violation, calculate the monetary differences using only those laws.
# Provide your analysis in the following format, entirely in Hebrew:

# ניתוח מקצועי של הפרות שכר:

# הפרה: [כותרת ההפרה]
# [תיאור מפורט של ההפרה, כולל תאריכים רלוונטיים, שעות עבודה, שכר שעתי וחישובים, בהתבסס אך ורק על החוקים הישראליים שנמצאו והמסמכים שסופקו.
# דוגמה: העובד עבד X שעות נוספות בין [חודש שנה] ל-[חודש שנה]. לפי שכר שעתי בסיסי של [שכר] ₪ ושיעורי תשלום שעות נוספות ([שיעור1]% עבור X השעות הראשונות, [שיעור2]% לאחר מכן) כפי שמופיע בחוקי העבודה שנמצאו, העובד היה זכאי ל-[סכום] ₪ לחודש. בפועל קיבל רק [סכום שקיבל] ₪ למשך X חודשים ו-[סכום] ₪ בחודש [חודש].]
# סה"כ חוב: [סכום ההפרש עבור הפרה זו] ₪

# הפרה: [כותרת ההפרה]
# [תיאור מפורט של ההפרה, כולל תאריכים וחישובים, בהתבסס אך ורק על החוקים שנמצאו והמסמכים. 
# דוגמה: בחודש [חודש שנה] לא בוצעה הפקדה לפנסיה. המעסיק מחויב להפקיד [אחוז]% מהשכר בגובה [שכר] ₪ = [סכום] ₪ בהתאם לחוק/צו הרחבה שנמצא.]
# סה"כ חוב פנסיה: [סכום חוב הפנסיה להפרה זו] ₪

# ---

# סה"כ תביעה משפטית (לא כולל ריבית): [הסכום הכולל לתביעה מכלל ההפרות] ₪  
# אסמכתאות משפטיות: [רשימת שמות החוק הרלוונטיים מתוך החוקים הישראליים שנמצאו. לדוגמה: חוק שעות עבודה ומנוחה, צו הרחבה לפנסיה חובה]

# SUMMARY TABLE:
# After completing the professional analysis, provide a summary table with actual data:
# - Use the heading: === טבלת סיכום הפרות מקצועי ===
# - Create columns for: סוג הפרה | תקופה | סכום (₪)
# - Add rows with actual violation types, periods, and amounts from your analysis
# - End with a total line showing the total amount in ₪
# - Include legal references from retrieved laws

# CRITICAL: Replace ALL placeholders with actual data from the analysis. Do not output template text.
# """

    def _get_warning_letter_instructions(self) -> str:
        format_content = self.letter_format.get_format().get('content', '')

        return f"""
INSTRUCTIONS:
1. Analyze the provided documents for labor law violations *based exclusively on the retrieved LABOR LAWS and JUDGEMENTS*.
2. If violations are found, generate a formal warning letter using the provided template.
3. If no violations are found, respond with: "לא נמצאו הפרות המצדיקות מכתב התראה." in Hebrew.

Warning Letter Template:
{format_content}

Please generate the warning letter in Hebrew with the following guidelines:
- Replace [שם המעסיק] with the employer's name from the documents
- Replace [פרטי ההפרות] with specific details of each violation found (based on retrieved laws)
- Replace [הפניות לחוקים] with relevant labor law citations *from the retrieved LABOR LAWS*.
- Replace [מועד] with a reasonable timeframe for corrections (typically 14 days)
- Maintain a professional and formal tone throughout
- Include all violations found in the analysis (based on retrieved laws)
- Format the letter according to the provided template structure
"""

    def _get_easy_instructions(self) -> str:
        return """
🔒 מטרה: צור סיכום קצר וברור של ההפרות בתלושי השכר של העובד.
📌 כללים מחייבים:
    1. כתוב בעברית בלבד – אל תשתמש באנגלית בכלל.
    2. עבור כל חודש הצג את ההפרות בשורות נפרדות, כל שורה בפורמט הבא:
❌ [סוג ההפרה בקצרה] – [סכום בש"ח עם ₪, כולל פסיק לאלפים]
לדוגמה: ❌ לא שולם החזר נסיעות בפברואר 2025 – 250 ₪
    3. אם יש מספר רכיבי פנסיה (עובד/מעסיק/בריאות) בחודש מסוים – חבר אותם לסכום אחד של פנסיה באותו החודש.
    4. כל הסכומים יוצגו עם פסיקים לאלפים ועם ₪ בסוף.
    5. חישוב הסכום הכולל יופיע בשורה נפרדת:
💰 סה"כ: [סכום כולל] ₪
    6. הוסף המלצה בסוף:
📝 מה לעשות עכשיו:
פנה/י למעסיק עם דרישה לתשלום הסכומים.
אם אין מענה – מומלץ לפנות לייעוץ משפטי.
📍 הנחיות נוספות:
    • אין לכתוב מספרים בלי הקשר, כל שורה חייבת להיות מלווה בחודש.
    • מיזוג שורות: אם באותו חודש יש כמה רכיבים של פנסיה – מיזג אותם לשורה אחת.
    • הסר שורות ללא סכום ברור.
    • ניסוח פשוט, ללא מינוחים משפטיים, הבהרות או הסברים טכניים.
    • אין לציין "רכיב עובד", "רכיב מעסיק", "לא הופקד" – במקום זאת כתוב: "לא שולמה פנסיה".
🎯 פלט רצוי:
    • שורות מסודרות לפי חודשים
    • אין כפילויות
    • סכומים מדויקים בלבד
    • ניסוח ברור ומובן
    • עברית בלבד

🧪 Example of desired output:
📢 סיכום ההפרות:
❌ לא שולם עבור שעות נוספות בנובמבר 2024 – 495 ₪
❌ לא שולמה פנסיה בנובמבר 2024 – 750 ₪
❌ לא שולמה פנסיה בדצמבר 2024 – 1,221 ₪
❌ לא שולמה פנסיה בינואר 2025 – 831 ₪
❌ לא שולם החזר נסיעות בפברואר 2025 – 250 ₪
❌ לא שולמה פנסיה בפברואר 2025 – 858 ₪
❌ לא שולמה פנסיה במרץ 2025 – 866 ₪
💰 סה"כ: 5,271 ₪
📝 מה לעשות עכשיו:
פנה/י למעסיק עם דרישה לתשלום הסכומים.
אם אין מענה – מומלץ לפנות לייעוץ משפטי.

SUMMARY TABLE:
After completing the easy summary, provide a simple visual table:

=== טבלת סיכום ויזואלי ===
חודש/שנה            | סוג הפרה           | סכום (₪)
נובמבר 2024          | שעות נוספות        | 495
נובמבר 2024          | פנסיה             | 750
דצמבר 2024           | פנסיה             | 1,221
ינואר 2025           | פנסיה             | 831
פברואר 2025          | נסיעות            | 250
פברואר 2025          | פנסיה             | 858
מרץ 2025             | פנסיה             | 866
-----------------------------------------
סה"כ: 5,271 ₪
"""

    def _get_table_instructions(self) -> str:
        return """
אתה עוזר משפטי מיומן. עליך לנתח את רשימת ההפרות ולהפיק רשימת תביעות מסודרת לפי מסמך (לדוגמה: תלוש שכר מס' 1, מסמך שימוע, מכתב פיטורין וכו').

הנחיות:

1. סדר את התביעות לפי מסמך: כל קבוצה מתחילה בכותרת כמו "תלוש שכר מס' 4 – 05/2024:".

2. תחת כל כותרת, צור רשימה ממוספרת באותיות עבריות (א., ב., ג. וכו').

3. השתמש במבנה הקבוע הבא:
   א. סכום של [amount] ש"ח עבור [תיאור קצר של ההפרה].

4. השתמש בפורמט מספרים עם פסיקים לאלפים ושתי ספרות אחרי הנקודה (למשל: 1,618.75 ש"ח).

5. כתוב "ש"ח" אחרי הסכום — לא ₪.

6. אל תסכם את כלל ההפרות – הצג רק את הפירוט לפי מסמך.

7. אל תוסיף בולטים, טבלאות או הערות.

נתוני הקלט לדוגמה:
❌ אי תשלום שעות נוספות (תלוש 6, 11/2024) – 495 ש"ח
❌ העדר רכיב מעביד לפנסיה (תלוש 6, 11/2024) – 750 ש"ח
❌ העדר רכיב עובד לפנסיה (תלוש 7, 12/2024) – 396 ש"ח
❌ העדר נסיעות (תלוש 9, 02/2025) – 250 ש"ח

פלט נדרש לדוגמה:

תלוש שכר מס' 6 – 11/2024:
א. סכום של 495.00 ש"ח עבור אי תשלום שעות נוספות.
ב. סכום של 750.00 ש"ח עבור העדר רכיב מעביד לפנסיה.

תלוש שכר מס' 7 – 12/2024:
א. סכום של 396.00 ש"ח עבור העדר רכיב עובד לפנסיה.

תלוש שכר מס' 9 – 02/2025:
א. סכום של 250.00 ש"ח עבור העדר תשלום עבור נסיעות.

החזר את הפלט בפורמט זה בלבד, מופרד לפי כל מסמך.

SUMMARY TABLE:
After completing the organized document list, provide a simple summary table:

=== טבלת סיכום לפי מסמך ===
מסמך                | מספר הפרות         | סכום כולל (ש"ח)
תלוש שכר מס' 6       | 2                  | 1,245.00
תלוש שכר מס' 7       | 1                  | 396.00  
תלוש שכר מס' 9       | 1                  | 250.00
------------------------------------------
סה"כ                | 4                  | 1,891.00
"""

    def _get_claim_instructions(self) -> str:
        return """
משימה:
כתוב טיוטת כתב תביעה לבית הדין האזורי לעבודה, בהתאם למבנה המשפטי הנהוג בישראל.

נתונים:
השתמש במידע מתוך המסמכים שצורפו (כגון תלושי שכר, הסכמי עבודה, הודעות פיטורין, שיחות עם המעסיק) ובממצאים שנמצאו בניתוח קודם של ההפרות.

פורמט לכתיבה:
1. כותרת: "בית הדין האזורי לעבודה ב[שם עיר]"
2. פרטי התובע/ת:
   - שם מלא, ת"ז, כתובת, טלפון, מייל
3. פרטי הנתבע/ת (מעסיק):
   - שם החברה/מעסיק, ח.פ./ת"ז, כתובת, טלפון, מייל (אם קיים)
4. כותרת: "כתב תביעה"
5. פתיח משפטי:
   בית הדין הנכבד מתבקש לזמן את הנתבע/ת לדין ולחייבו/ה לשלם לתובע/ת את הסכומים המפורטים, מהנימוקים הבאים:
6. סעיף 1 – רקע עובדתי:
   תיאור תקופת העבודה, תפקידים, תאריך תחילה וסיום (אם רלוונטי), מהות יחסי העבודה, מקום העבודה.
7. סעיף 2 – עילות התביעה (לפי הפרות):
   לכל הפרה:
     - איזה חוק הופֵר (לדוג': חוק שכר מינימום, חוק שעות עבודה ומנוחה וכו')
     - פירוט העובדות והתקופה הרלוונטית
     - סכום הפיצוי או הנזק
     - אסמכתאות משפטיות אם רלוונטי
8. סעיף 3 – נזקים שנגרמו:
   סכומים כספיים (בפירוט), נזק לא ממוני אם קיים (עוגמת נפש).
9. סעיף 4 – סעדים מבוקשים:
   תשלום הפרשים, פיצויים, ריבית והצמדה, הוצאות משפט, שכר טרחת עו"ד, וכל סעד אחר.
10. סעיף 5 – סמכות שיפוט:
   ציין סמכות בית הדין לעבודה לפי חוק בית הדין לעבודה תשכ"ט–1969.
11. סעיף 6 – דיון מקדים/הליך מהיר (אם רלוונטי).
12. סעיף 7 – כתובות להמצאת כתבי בי-דין:
   כתובת התובע והנתבע.
13. סיום:
   חתימה, תאריך, רשימת נספחים תומכים (תלושי שכר, מכתבים וכו').

שפה:
- כתוב בעברית משפטית, רשמית ומסודרת.
- סדר את התביעה עם כותרות, סעיפים ממוספרים ורווחים ברורים.
- אל תמציא עובדות או חוקים. אם חסר מידע – השאר שדה ריק לצורך השלמה ע"י המשתמש.

אזהרה:
בסיום, כתוב משפט: "הטיוטה נכתבה אוטומטית לצורך עזרה ראשונית בלבד. מומלץ לפנות לעורך/ת דין מוסמך/ת לפני הגשה לבית הדין."

SUMMARY TABLE:
After completing the claim draft, provide a simple summary table:

=== טבלת סיכום כתב התביעה ===
עילת תביעה           | תקופה               | סכום נתבע (₪)
אי תשלום שעות נוספות  | נובמבר 2024          | 495
אי הפקדת פנסיה       | נובמבר 2024-מרץ 2025  | 3,525
אי החזר נסיעות       | פברואר 2025          | 250
---------------------------------------------
סה"כ תביעה ראשית: 4,270 ₪
ריבית והצמדה: לפי החלטת בית הדין
הוצאות משפט: לפי החלטת בית הדין
"""

    def _get_combined_instructions(self) -> str:
        return """
COMPREHENSIVE ANALYSIS INSTRUCTIONS:

📌 OVERVIEW
You are analyzing employee documents(payslips, employment contracts, and attendance records) for legal violations based **only** on the provided labor laws and judgements. Do **not** use any default knowledge, assumptions, or examples from the law unless explicitly confirmed in the payslip.
and explain step by step your thought process, calculations, and legal references.
labor laws and judgements are different from employee documents dont use example values from the legal texts or examples. use only real data from the documents provided.

PART 1 - VIOLATION IDENTIFICATION AND ANALYSIS:

1. If no labor laws are provided, respond with: "אין חוקים לעבודה זמינים לניתוח התאמה." in Hebrew.
2. If labor laws exist, analyze the payslip ONLY against the laws provided.
3. Use **only** the documents provided (e.g., payslip data, employment contracts data, and attendance records data). **Do not extract or reuse any example values (e.g., 6000 ₪, 186 hours, 14 hours overtime) that appear in the legal texts or examples.**
4. Do **not invent** missing data. If the document does not include sufficient detail for a violation (e.g., no overtime hours), **do not report a violation**.
5. If no judgements are provided, respond with: "לא קיימות החלטות משפטיות זמינות לניתוח." in Hebrew.


PART 2 - DETAILED PROFESSIONAL ANALYSIS FORMAT:

For each violation found, provide the following, using actual values from the document:

הפרה: [כותרת ההפרה]
[תיאור מפורט של ההפרה, כולל תאריכים רלוונטיים, שעות עבודה בפועל, שכר שעתי, ימי מחלה מאושרים וכדומה. התבסס רק על הנתונים הקיימים בתלוש, ובאופן בלעדי על החוקים שהוצגו.]

[הפניה לחוק הרלוונטי ושנה מתוך החוקים שסופקו]

[אם קיימים – תאר תקדימים משפטיים מתוך פסקי הדין שסופקו. אם אין – כתוב "לא נמצאו תקדימים רלוונטיים בפסקי הדין שסופקו."]

[השלכות משפטיות ופעולות מומלצות]

סה"כ חוב להפרה זו: [סכום] ₪

---

PART 3 - ORGANIZED TABLE FORMAT BY DOCUMENT:

רשימת תביעות מסודרת לפי מסמך:

תלוש שכר מס' [מספר] – [חודש/שנה]:
א. סכום של [amount] ש"ח עבור [תיאור קצר של ההפרה].
ב. סכום של [amount] ש"ח עבור [תיאור קצר של ההפרה].

---

PART 4 - FINAL SUMMARY:

סה"כ תביעה משפטית (לא כולל ריבית): [סכום כולל] ₪  
אסמכתאות משפטיות: [רשימת שמות החוקים הרלוונטיים מתוך החוקים שסופקו]

---

PART 5 - COMPREHENSIVE SUMMARY TABLE:

=== טבלת סיכום מקיפה ===

Create columns:
- תלוש/מסמך
- סוג הפרה
- תקופה
- סכום (₪)

Add rows using actual data per document and violation.

🚨 TOTAL LINE:
Show total number of violations, total documents, and total amount claimed.

🔍 Breakdown by violation type:
- Total number of documents where each violation occurred
- Total amount per violation type

📘 Legal references:
List names of all laws used in calculations.

---

⚠️ FORMATTING & ACCURACY RULES:

- Always respond in Hebrew
- NEVER use sample values or formulas from laws directly
- Use only data provided in the payslip, employment contracts, and attendance records
- Replace ALL placeholders with actual values
- Do not hallucinate sick days, overtime hours, or absences
- If no violations found for a payslip, write:
  "לא נמצאו הפרות בתלוש מספר [X]"
- If no violations in any payslip, write:
  "לא נמצאו הפרות נגד חוקי העבודה שסופקו."
"""

    def _get_violation_count_table_instructions(self) -> str:
        return """
INSTRUCTIONS:
Analyze the provided documents and identify potential labor law violations, based *exclusively on the retrieved LABOR LAWS and JUDGEMENTS*.
Return ONLY a summary table in Hebrew, with NO additional analysis, commentary, or template text. Do not include any explanations, recommendations, or placeholders.

Provide a comprehensive summary table with actual data:
- Use the heading: === טבלת סיכום מקיפה ===
- Create columns for: תלוש/מסמך | סוג הפרה | תקופה | סכום (₪) | מספר הפרות
- Add rows with actual document names, violation types, periods, amounts, and the count of each violation type per document
- For each violation type, calculate and display the total amount for that violation type across all documents, and the total count of documents/slips affected
- For each violation type, explicitly state how many documents/slips had that violation (e.g., "6 מתוך 10 תלושים נמצאה הפרה של אי תשלום פנסיה")
- End with a total line showing total violations, months, and total amount
- Include a breakdown by violation type, showing the total count and total amount for each type of violation
- All amount calculations must be thorough and accurate: sum the amounts for each violation type, each document, and overall totals. Do not estimate or round; use the actual amounts found in the analysis.
- Include legal references from retrieved laws

Formatting requirements:
- Always respond in Hebrew
- Use actual violation types, counts, and amounts from the analysis
- Do not output any template text or placeholders
- Do not include any additional commentary or explanations
- If no violations are found, respond with: "לא נמצאו הפרות נגד חוקי העבודה שסופקו."
"""

    def _extract_text2(self, content: bytes, filename: str, compress: bool = False) -> str:
        """Extract text from various document formats using AgenticDoc for PDFs and images"""
        print("Extracting text from file:", filename.lower())
        
        try:
            if filename.lower().endswith(('.pdf', '.png', '.jpg', '.jpeg')):
                # Pass bytes directly to AgenticDoc
                print("Extracting text from image or PDF...")
                
                result = parse(content)
                # Return markdown or structured, as you prefer
                return result[0].markdown
            elif filename.lower().endswith('.docx'):
                doc_file = BytesIO(content)
                doc = Document(doc_file)
                return '\n'.join([p.text for p in doc.paragraphs])
            elif filename.lower().endswith('.xlsx'):
                excel_file = BytesIO(content)
                df = pd.read_excel(excel_file, sheet_name=None)
                text = ""
                for sheet, data in df.items():
                    text += f"## Sheet: {sheet}\n"
                    # Convert DataFrame to Markdown table
                    text += data.to_markdown(index=False) + "\n\n"
                return text
            else:
                # Unknown file type, treat as plain text
                result = parse(content)      
                print(f"Extracting text :{result}")
                # Return markdown or structured, as you prefer
                return result[0].markdown
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Extraction error: {get_error_detail(e)}")

    def _extract_text_legacy(self, content: bytes, filename: str, compress: bool = False) -> str:
        """Legacy text extraction method (kept for fallback)"""
        if filename.lower().endswith('.pdf'):
            return self._extract_pdf_text(content)
        elif filename.lower().endswith('.docx'):
            return self._extract_docx_text(content)
        elif filename.lower().endswith('.xlsx'):
            return self._extract_excel_text(content)
        else:
            return self._extract_image_text(content, compress)

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF using pdfplumber and Vision API"""
        try:
            pdf_file = BytesIO(content)
            text = ''
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += page_text + "\n"
                    else:
                        # Convert PDF page to image and use Vision API
                        img = page.to_image(resolution=300).original
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='PNG')
                        img_content = img_byte_arr.getvalue()
                        
                        vision_image = vision.Image(content=img_content)
                        response = self.vision_client.text_detection(image=vision_image, image_context=self.image_context)
                        if response.text_annotations:
                            text_list = [text_annotation.description for text_annotation in response.text_annotations]
                            text += " ".join(text_list) + "\n"
            return text
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error processing PDF: {get_error_detail(e)}")

    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX files"""
        doc_file = BytesIO(content)
        doc = Document(doc_file)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    def _extract_excel_text(self, content: bytes) -> str:
        """Extract text from Excel files"""
        try:
            excel_file = BytesIO(content)
            df = pd.read_excel(excel_file, sheet_name=None)  # Load all sheets
            
            text = ""
            for sheet_name, sheet_data in df.items():
                text += f"Sheet: {sheet_name}\n"
                text += sheet_data.to_string(index=False) + "\n\n"
            return text
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error processing Excel file: {get_error_detail(e)}")

    def _extract_image_text(self, content: bytes, compress: bool = False) -> str:
        """Extract text from images using Vision API"""
        try:
            if compress:
                content = self._compress_image(content)
            
            vision_image = vision.Image(content=content)
            response = self.vision_client.document_text_detection(image=vision_image, image_context=self.image_context)
            
            if response.error.message:
                raise HTTPException(status_code=500, detail=f"Vision API Error: {response.error.message}")
            
            if response.text_annotations:
                return " ".join([text_annotation.description for text_annotation in response.text_annotations])
            else:
                return ""
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error processing image: {get_error_detail(e)}")

    def _compress_image(self, image_bytes: bytes, max_size_mb: int = 4) -> bytes:
        """Compress image to reduce size for API limits"""
        # Convert size to bytes
        max_size_bytes = max_size_mb * 1024 * 1024
        
        # If image is already smaller than max size, return original
        if len(image_bytes) <= max_size_bytes:
            return image_bytes
        
        # Open image using PIL
        img = Image.open(BytesIO(image_bytes))
        
        # Get original dimensions
        width, height = img.size
        
        # Scale down large images while maintaining aspect ratio
        max_dimension = 3000  # Maximum dimension for width or height
        if width > max_dimension or height > max_dimension:
            scale = min(max_dimension / width, max_dimension / height)
            new_width = int(width * scale)
            new_height = int(height * scale)
            img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
        
        # Preserve original format if possible
        output_format = img.format or 'JPEG'
        
        # Initial quality and compression settings
        quality = 95
        output = BytesIO()
        
        # Progressive quality reduction with format-specific handling
        while quality > 20:  # Increased minimum quality threshold
            output.seek(0)
            output.truncate(0)
            
            if output_format == 'JPEG':
                img.save(output, format=output_format, quality=quality, optimize=True)
            elif output_format == 'PNG':
                img.save(output, format=output_format, optimize=True)
            else:
                # Default to JPEG for unsupported formats
                img.save(output, format='JPEG', quality=70, optimize=True)
            
            if output.tell() <= max_size_bytes:
                break
            
            # More gradual quality reduction
            quality -= 5
        
        # If still too large, convert to JPEG as last resort
        if output.tell() > max_size_bytes and output_format != 'JPEG':
            output.seek(0)
            output.truncate(0)
            img.save(output, format='JPEG', quality=70, optimize=True)
        
        return output.getvalue()

    async def summarise(self, ai_content_text: str) -> str:
        """
        Summarizes the given text using PydanticAI.
        """
        try:
            prompt = f"Please summarize the following text concisely in Hebrew:\n\n{ai_content_text}"
            
            # Use the same agent for summarization
            try:
                result = await self.agent.run(prompt)
                return result.output if hasattr(result, 'output') else str(result)
            except Exception as pydantic_error:
                # If PydanticAI fails, check if it's an event loop issue
                if "Event loop is closed" in str(pydantic_error) or "event loop" in str(pydantic_error).lower():
                    # Try to recreate the agent with a fresh context
                    print("Attempting to recreate PydanticAI agent for summarization due to event loop issue...")
                    
                    # Create a new agent instance
                    temp_agent = Agent(
                        model=self.model,
                        output_type=str,
                        system_prompt="You are a helpful assistant that summarizes text concisely."
                    )
                    
                    # Try the request again with the fresh agent
                    result = await temp_agent.run(prompt)
                    return result.output if hasattr(result, 'output') else str(result)
                else:
                    # Re-raise the original error if it's not event loop related
                    raise pydantic_error
            
        except Exception as e:
            error_detail = f"Error generating summary with PydanticAI: {get_error_detail(e)}"
            raise HTTPException(
                status_code=500,
                detail=error_detail
            )
    
    # Sync wrapper methods for backward compatibility
    def create_report_sync(self, payslip_text: str = None, contract_text: str = None, 
                          attendance_text: str = None, type: str = "report", 
                          context: str = None) -> Dict:
        """Synchronous wrapper for create_report method with improved event loop handling"""
        # Handle backward compatibility parameter mapping
        try:
            # Try to get existing event loop
            loop = asyncio.get_event_loop()
            if loop.is_running():
                # If loop is already running, create a new thread
                import concurrent.futures
                import threading
                
                def run_in_thread():
                    # Create new event loop for this thread
                    new_loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(new_loop)
                    try:
                        return new_loop.run_until_complete(self.create_report(
                            payslip_text=payslip_text,
                            contract_text=contract_text, 
                            attendance_text=attendance_text,
                            analysis_type=type,  # Map 'type' to 'analysis_type'
                            context=context
                        ))
                    finally:
                        new_loop.close()
                
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(run_in_thread)
                    result = future.result()
            else:
                # No running loop, use asyncio.run
                result = asyncio.run(self.create_report(
                    payslip_text=payslip_text,
                    contract_text=contract_text, 
                    attendance_text=attendance_text,
                    analysis_type=type,  # Map 'type' to 'analysis_type'
                    context=context
                ))
        except RuntimeError as e:
            if "no current event loop" in str(e).lower() or "event loop is closed" in str(e).lower():
                # Create new event loop
                result = asyncio.run(self.create_report(
                    payslip_text=payslip_text,
                    contract_text=contract_text, 
                    attendance_text=attendance_text,
                    analysis_type=type,  # Map 'type' to 'analysis_type'
                    context=context
                ))
            else:
                raise e
        
        # Convert response to dict for backward compatibility
        return {
            "legal_analysis": result.legal_analysis,
            "status": result.status,
            # "relevant_laws": result.relevant_laws,
            # "relevant_judgements": result.relevant_judgements
        }
    
    def create_report_with_optimization_sync(self, payslip_text: str = None, contract_text: str = None, 
                          attendance_text: str = None, type: str = "report", 
                          context: str = None, use_optimizer: bool = True, target_score: float = 0.85) -> Dict:
        """Synchronous wrapper for create_report_with_optimization method"""
        try:
            # Try to get existing event loop
            loop = asyncio.get_event_loop()
            if loop.is_running():
                # If loop is already running, create a new thread
                import concurrent.futures
                
                def run_in_thread():
                    # Create new event loop for this thread
                    new_loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(new_loop)
                    try:
                        return new_loop.run_until_complete(self.create_report_with_optimization(
                            payslip_text=payslip_text,
                            contract_text=contract_text, 
                            attendance_text=attendance_text,
                            analysis_type=type,  # Map 'type' to 'analysis_type'
                            context=context,
                            use_optimizer=use_optimizer,
                            target_score=target_score
                        ))
                    finally:
                        new_loop.close()
                
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(run_in_thread)
                    return future.result()
            else:
                # No running loop, use asyncio.run
                return asyncio.run(self.create_report_with_optimization(
                    payslip_text=payslip_text,
                    contract_text=contract_text, 
                    attendance_text=attendance_text,
                    analysis_type=type,  # Map 'type' to 'analysis_type'
                    context=context,
                    use_optimizer=use_optimizer,
                    target_score=target_score
                ))
        except RuntimeError as e:
            if "no current event loop" in str(e).lower() or "event loop is closed" in str(e).lower():
                # Create new event loop
                return asyncio.run(self.create_report_with_optimization(
                    payslip_text=payslip_text,
                    contract_text=contract_text, 
                    attendance_text=attendance_text,
                    analysis_type=type,  # Map 'type' to 'analysis_type'
                    context=context,
                    use_optimizer=use_optimizer,
                    target_score=target_score
                ))
            else:
                raise e
    
    def summarise_sync(self, ai_content_text: str) -> str:
        """Synchronous wrapper for summarise method with improved event loop handling"""
        try:
            # Try to get existing event loop
            loop = asyncio.get_event_loop()
            if loop.is_running():
                # If loop is already running, create a new thread
                import concurrent.futures
                
                def run_in_thread():
                    # Create new event loop for this thread
                    new_loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(new_loop)
                    try:
                        return new_loop.run_until_complete(self.summarise(ai_content_text))
                    finally:
                        new_loop.close()
                
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(run_in_thread)
                    return future.result()
            else:
                # No running loop, use asyncio.run
                return asyncio.run(self.summarise(ai_content_text))
        except RuntimeError as e:
            if "no current event loop" in str(e).lower() or "event loop is closed" in str(e).lower():
                # Create new event loop
                return asyncio.run(self.summarise(ai_content_text))
            else:
                raise e
