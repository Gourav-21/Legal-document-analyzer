from fastapi import UploadFile, HTTPException
from PIL import Image
import pytesseract
import google.generativeai as genai
import os
from dotenv import load_dotenv
import pdfplumber
from io import BytesIO
from typing import List, Dict, Union
from docx import Document
from labour_law import LaborLawStorage
from routers.letter_format_api import LetterFormatStorage
from judgement import JudgementStorage
from google.cloud import vision
import io
from google.cloud.vision import ImageAnnotatorClient
import pandas as pd

# Load environment variables from .env file
load_dotenv()

class DocumentProcessor:
    def __init__(self):
        # Initialize Gemini AI
        genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
        self.model = genai.GenerativeModel('gemini-2.0-flash')
        
        # Initialize Labor Law Storage
        self.law_storage = LaborLawStorage()
        self.letter_format = LetterFormatStorage()
        self.judgement_storage = JudgementStorage()

        # Configure Tesseract path
        # tesseract_path = os.getenv("TESSERACT_CMD", r"C:\Program Files\Tesseract-OCR\tesseract.exe")
        # if not os.path.exists(tesseract_path):
        #     raise Exception("Tesseract not found. Please install Tesseract and set TESSERACT_CMD environment variable.")
        # pytesseract.pytesseract.tesseract_cmd = tesseract_path

        # Configure Tesseract path
        # tesseract_path = os.getenv("TESSERACT_CMD", r"C:\Program Files\Tesseract-OCR\tesseract.exe")
        # if not os.path.exists(tesseract_path):
        #     raise Exception("Tesseract not found. Please install Tesseract and set TESSERACT_CMD environment variable.")
        # pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Initialize Vision client with API key
        vision_api_key = os.getenv("GOOGLE_CLOUD_VISION_API_KEY")
        if not vision_api_key:
            raise Exception("Google Cloud Vision API key not found. Please set GOOGLE_CLOUD_VISION_in your .env filevariable.")
        
        self.vision_client = ImageAnnotatorClient(client_options={"api_key": vision_api_key})
        self.image_context = {"language_hints": ["he"]} 


    def process_document(self, files: Union[UploadFile, List[UploadFile]], doc_types: Union[str, List[str]], compress: bool = False) -> Dict[str, str]:
        if not files:
            raise HTTPException(
                status_code=400,
                detail="At least one document must be provided"
            )
        
        # Handle single file case
        if isinstance(files, UploadFile):
            files = [files]
            doc_types = [doc_types]
        elif not isinstance(doc_types, list):
            doc_types = [doc_types] * len(files)
        
        # Initialize text storage
        payslip_text = None
        contract_text = None
        attendance_text = None
        
        # Initialize payslip counter
        payslip_counter = 0
        contract_counter = 0
        attendance_counter = 0
        
        # Process each file based on its type
        for file, doc_type in zip(files, doc_types):
            extracted_text = self._extract_text2(file.file.read(), file.filename, compress=compress)
            if doc_type.lower() == "payslip":
                payslip_counter += 1
                payslip_text = f"Payslip {payslip_counter}:\n{extracted_text}" if payslip_text is None else f"{payslip_text}\n\nPayslip {payslip_counter}:\n{extracted_text}"
            elif doc_type.lower() == "contract":
                contract_counter += 1
                contract_text = f"Contract {contract_counter}:\n{extracted_text}" if contract_text is None else f"{contract_text}\n\nContract {contract_counter}:\n{extracted_text}"
            elif doc_type.lower() == "attendance":
                attendance_counter += 1
                attendance_text = f"Attendance {attendance_counter}:\n{extracted_text}" if attendance_text is None else f"{attendance_text}\n\nAttendance {attendance_counter}:\n{extracted_text}"
        # Return the extracted texts
        return {
            "payslip_text": payslip_text,
            "contract_text": contract_text,
            "attendance_text": attendance_text
        }
    
    def create_report(self, payslip_text: str = None, contract_text: str = None, attendance_text: str = None, type: str = "report", context: str = None) -> Dict:
        # Prepare documents for analysis
        documents = {}
        if payslip_text:
            documents['payslip'] = payslip_text
        if contract_text:
            documents['contract'] = contract_text
        if attendance_text:
            documents['attendance report'] = attendance_text
            
        # Prepare the prompt for Gemini AI
        # Get formatted labor laws
        labor_laws = self.law_storage.format_laws_for_prompt()
        # Get formatted judgements
        judgements = self.judgement_storage.format_judgements_for_prompt()
        
        prompt = f"""You are a legal document analyzer specializing in Israeli labor law compliance based *only* on user-provided information.

    LABOR LAWS TO CHECK AGAINST(Your ONLY reference point):
    {labor_laws if labor_laws else 'No labor laws provided for analysis.'}

    JUDGEMENTS TO CONSIDER (Your ONLY reference point):
    {judgements if judgements else 'No judgements provided for analysis.'}

    DOCUMENTS PROVIDED FOR ANALYSIS:
    {', '.join(documents.keys())}

    ADDITIONAL CONTEXT:
    {context if context else 'No additional context provided.'}

    """

        # Add document contents to prompt
        for doc_type, content in documents.items():
            prompt += f"\n{doc_type.upper()} CONTENT:\n{content}\n"

            prompt += f"""
    INSTRUCTIONS:
    1. If no labor laws are provided, respond with: "אין חוקים לעבודה זמינים לניתוח התאמה." in Hebrew.
    2. If labor laws exist, analyze the documents ONLY against the provided laws.
    3. ONLY refer to the judgements and their results provided above for legal analysis - do not use external cases or knowledge.
    4. If no judgements are provided, respond with: "לא קיימות החלטות משפטיות זמינות לניתוח." in Hebrew."""

        if(type=='report'):
            prompt += f"""
    3. For each payslip provided, analyze and identify violations. For each violation found in each payslip, format the response EXACTLY as shown below, with each section on a new line and proper spacing:

    Violation Format Template:

    [VIOLATION TITLE]

    [SPECIFIC VIOLATION DETAILS]

    [LAW REFERENCE AND YEAR FROM PROVIDED LAWS]

    [SIMILAR CASES OR PRECEDENTS FROM PROVIDED JUDGEMENTS](Refer *only* to the 'JUDGEMENTS TO CONSIDER' section provided earlier. If a relevant judgement is found among those provided, describe the case and its result. If no relevant judgement is found, state "לא נמצאו תקדימים רלוונטיים בפסקי הדין שסופקו.")

    [LEGAL IMPLICATIONS BASED ON PROVIDED INFORMATION]

    [RECOMMENDED ACTIONS]

    ---

    Example of correctly formatted violation:

    Payslip 1:

    Possible Violation of Mandatory Pension Expansion Order

    The payslip shows no pension contribution, despite over 6 months of employment.

    According to the Mandatory Pension Expansion Order (2008) (assuming this law was provided).

    In the [NAME OF RULING FROM PROVIDED JUDGEMENTS] ruling, there was a similar case and the employee won 10,000 thousand shekels there (assuming this judgement was provided).

    Lack of contribution may entitle the employee to retroactive compensation or legal action.

    It is recommended to review the pension fund details, start date of employment, and contract terms.

    ---

    IMPORTANT:
    - Always Respond in Hebrew
    - Format each violation with proper spacing and line breaks as shown above
    - Analyze each payslip separately and clearly indicate which payslip the violations belong to
    - Separate multiple violations with '---'
    - If no violations are found against the provided laws in a payslip, respond with: "לא נמצאו הפרות בתלוש מספר [X]" in hebrew
    - If no violations are found in any payslip, respond with: "לא נמצאו הפרות נגד חוקי העבודה שסופקו." in hebrew
    - Do not include any additional commentary or explanations outside of the violation format"""
        elif(type=='profitability'):
            prompt += f"""
    INSTRUCTIONS:
    1. Analyze the provided documents and identify potential labor law violations, based *exclusively on the LABOR LAWS and JUDGEMENTS provided above*.
    2. For each violation, refer *only* to the `JUDGEMENTS TO CONSIDER` (provided earlier in the main prompt) to identify similar legal cases and their outcomes (both successful and unsuccessful). Do not use external sources or knowledge for this.
    3. If similar cases (from the provided judgements) were unsuccessful:
       - Explain why the cases were unsuccessful based on the provided judgement information.
       - Provide a clear recommendation against pursuing legal action based on this.
       - List potential risks and costs.

    4. If similar cases (from the provided judgements) were successful, calculate using information from those judgements if available:
       - Average compensation amount from successful provided cases.
       - Estimated legal fees (30% of potential compensation).
       - Tax implications (25% of net compensation).
       - Time and effort cost estimation.

    Provide the analysis in the following format:

    ניתוח כדאיות כלכלית:

    הפרות שזוהו (על בסיס החוקים שסופקו):
    [List identified violations]

    תקדימים משפטיים (מתוך פסקי הדין שסופקו):
    [Similar cases from provided judgements with outcomes - both successful and unsuccessful]

    במקרה של תקדימים שליליים (מתוך פסקי הדין שסופקו):
    - סיבות לדחיית התביעות: [REASONS BASED ON PROVIDED JUDGEMENTS]
    - סיכונים אפשריים: [RISKS]
    - המלצה: לא מומלץ להגיש תביעה בשל [EXPLANATION BASED ON PROVIDED JUDGEMENTS]

    במקרה של תקדימים חיוביים (מתוך פסקי הדין שסופקו):
    ניתוח כספי:
    - סכום פיצוי ממוצע (מפסיקות דין שסופקו): [AMOUNT] ₪
    - עלות משוערת של עורך דין (30%): [AMOUNT] ₪
    - השלכות מס (25% מהסכום נטו): [AMOUNT] ₪
    - סכום נטו משוער: [AMOUNT] ₪

    המלצה סופית:
    [Based on analysis of both successful and unsuccessful cases from the provided judgements, provide clear recommendation]
    """

        elif(type=='professional'):
            prompt += f"""
    INSTRUCTIONS:
    Analyze the provided documents for labor law violations based ONLY on the provided labor laws and document content. Calculate monetary differences for each violation using ONLY the provided laws.

    Provide the analysis in the following format in Hebrew:

    ניתוח מקצועי של הפרות שכר:

    הפרה: [VIOLATION TITLE]
    [Detailed description of the violation, including relevant dates, hours, wage rates, and calculations based *strictly on the provided laws and documents*. Example: The employee worked X overtime hours in [Month Year] through [Month Year]. Based on a base hourly wage of [Wage] ILS and legal overtime rates ([Rate1]% for first X hours, [Rate2]% thereafter) *as specified in the provided labor laws*, the employee was entitled to [Amount] ILS/month. Received only [Amount Received] ILS for X months and [Amount Received] ILS in [Month].]
    סה"כ חוב: [Total underpaid amount for this specific violation] ILS

    הפרה: [VIOLATION TITLE]
    [Detailed description of the violation, including relevant dates, calculations based *strictly on the provided laws and documents*. Example: In [Month Year], no pension contribution was made. Employer must contribute [Percentage]% of [Salary] ILS = [Amount] ILS *as per the provided pension law/order*.]
    סה"כ חוב פנסיה: [Total unpaid pension for this specific violation] ILS

    ---

    סה"כ תביעה משפטית (לא כולל ריבית): [Total combined legal claim amount from all violations] ILS
    אסמכתאות משפטיות: [List relevant law names *from the 'LABOR LAWS TO CHECK AGAINST' section provided earlier*, e.g., חוק שעות עבודה ומנוחה, צו הרחבה לפנסיה חובה]

    IMPORTANT:
    - Respond ONLY in Hebrew.
    - Base all analysis and calculations STRICTLY on the provided labor laws and document content. Do NOT use external knowledge or make assumptions.
    - Show clear calculations within the violation description where applicable.
    - Calculate the total amount owed for EACH violation separately.
    - Calculate the final TOTAL legal claim by summing up all individual violation amounts.
    - List the names of the relevant laws used (from those provided) as legal references at the end.
    - If no violations are found, respond with: "לא נמצאו הפרות בהתאם לחוקי העבודה והמסמכים שסופקו."
    """

        elif(type=='warning_letter'):
            format_content = self.letter_format.get_format().get('content', '')
            prompt += f"""
    INSTRUCTIONS:
    1. Analyze the provided documents for labor law violations *based exclusively on the LABOR LAWS and JUDGEMENTS provided above*.
    2. If violations are found, generate a formal warning letter using the provided template.
    3. If no violations are found, respond with: "לא נמצאו הפרות המצדיקות מכתב התראה." in Hebrew.

    Warning Letter Template:
    {format_content}

    Please generate the warning letter in Hebrew with the following guidelines:
    - Replace [EMPLOYER_NAME] with the employer's name from the documents
    - Replace [VIOLATION_DETAILS] with specific details of each violation found (based on provided laws)
    - Replace [LAW_REFERENCES] with relevant labor law citations *from the 'LABOR LAWS TO CHECK AGAINST' section provided earlier*.
    - Replace [REQUIRED_ACTIONS] with clear corrective actions needed
    - Replace [DEADLINE] with a reasonable timeframe for corrections (typically 14 days)
    - Maintain a professional and formal tone throughout
    - Include all violations found in the analysis (based on provided laws)
    - Format the letter according to the provided template structure
    """

        elif(type=='easy'):
            prompt += f"""
    **Objective:** Generate a report of potential labor law violations based on provided analysis and calculations. This analysis and these calculations *must be based strictly on the LABOR LAWS and JUDGEMENTS provided in the initial sections of this prompt*. Adhere strictly to the specified format and rules below.

    **Input:** You will be given information identifying potential labor law violations (derived *only* from provided laws/judgements) and the corresponding calculated compensation amount (represented by 'X') for each violation found.

    **Output Instructions:**

    1. **Language:** Respond **exclusively in Hebrew**.

    2. **Violation Format:** For each identified violation, structure the output precisely as follows:

    [Title of Violation]

    [Simple explanation of what the employer may be liable for, based on provided laws]


    * Replace `[Title of Violation]` with the specific title of the violation (e.g., "Lower Wage than Minimum Wage" - if minimum wage law was provided).
    * Replace `[Simple explanation of what the employer may be liable for]` with a simple explanation of the potential obligation (e.g., "You may be entitled to compensation from your employer.").
    * Replace 'X' with the actual calculated compensation amount provided for that violation (e.g., "5,000 NIS").
    * Maintain the exact line breaks and spacing shown.
    3. **Multiple Violations:** If more than one violation is found, separate each complete violation block (as formatted above) with a line containing only `---`.
    4. **No Violations:** If the analysis indicates that no violations were found against the specific laws checked, respond **only** with the exact Hebrew phrase:
    ``No violations were found against the provided labor laws.''
    * Do not add any other text before or after this phrase if no violations are found.
    5. **Strict Adherence:** Do **not** include any introductory text, concluding remarks, summaries, additional commentary, or any explanations outside of the defined format for each violation or the "no violations" message.

    **Example of Correctly Formatted Output for a Single Violation (assuming relevant laws were provided for this analysis):**

    It seems that your hourly rate is less than the legal minimum wage (e.g., 32.79 NIS per hour, if this figure was in the provided laws).

    You may be entitled to compensation from your employer.

    According to my calculations, you may be entitled to compensation of 2,345 NIS.

    **Example of Correctly Formatted Output for Multiple Violations (assuming relevant laws were provided for this analysis):**

    Failure to pay for overtime as required by law.

    You may be entitled to additional payment for hours worked in excess of the daily/weekly quota.

    According to my calculations, you may be entitled to compensation of 1,800 NIS.

    ---

    Dismissal without due process.

    Your dismissal may not have been made in accordance with the procedure required by law.

    According to my calculations, you may be entitled to compensation of 15,000 NIS.

    **Now, process the following violation data (which was derived strictly from the provided labor laws and judgements) and generate the response according to these instructions:**

    [Here you would insert the specific violation details and calculated 'X' amounts based on your analysis]"""
            
        try:
            # Generate analysis using Gemini AI
            response = self.model.generate_content(prompt)
            analysis = response.text
            
            # Structure the result
            result = {
                "legal_analysis": analysis,
                "status": "success"
            }
            
            return result
            
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error generating legal analysis: {str(e)}"
            )

            
    def _extract_text(self, content: bytes, filename: str) -> str:
        if filename.lower().endswith('.pdf'):
            try:
                pdf_file = BytesIO(content)
                text = ''
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += page_text + "\n"
                        else:
                            img = page.to_image(resolution=300).original
                            text += pytesseract.image_to_string(img) + "\n"
            except Exception as e:
                raise HTTPException(
                    status_code=500,
                    detail=f"Error processing PDF: {str(e)}"
                )
        elif filename.lower().endswith('.docx'):
            doc_file = BytesIO(content)
            doc = Document(doc_file)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        else:
            image = Image.open(BytesIO(content))
            text = pytesseract.image_to_string(image, lang='eng+heb')
        return text
    # google vision api
    def _compress_image(self, image_bytes: bytes, max_size_mb: int = 4) -> bytes:
        # Convert size to bytes
        max_size_bytes = max_size_mb * 1024 * 1024
        
        # If image is already smaller than max size, return original
        if len(image_bytes) <= max_size_bytes:
            return image_bytes
        
        # Open image using PIL
        img = Image.open(BytesIO(image_bytes))
        
        # Get original dimensions
        width, height = img.size
        
        # Scale down large images while maintaining aspect ratio
        max_dimension = 3000  # Maximum dimension for width or height
        if width > max_dimension or height > max_dimension:
            scale = min(max_dimension / width, max_dimension / height)
            new_width = int(width * scale)
            new_height = int(height * scale)
            img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
        
        # Preserve original format if possible
        output_format = img.format or 'JPEG'
        
        # Initial quality and compression settings
        quality = 95
        output = BytesIO()
        
        # Progressive quality reduction with format-specific handling
        while quality > 20:  # Increased minimum quality threshold
            output.seek(0)
            output.truncate(0)
            
            if output_format == 'JPEG':
                img.save(output, format=output_format, quality=quality, optimize=True)
            elif output_format == 'PNG':
                img.save(output, format=output_format, optimize=True)
            else:
                # Default to JPEG for unsupported formats
                img.save(output, format='JPEG', quality=quality, optimize=True)
            
            if output.tell() <= max_size_bytes:
                break
            
            # More gradual quality reduction
            quality -= 5
        
        # If still too large, convert to JPEG as last resort
        if output.tell() > max_size_bytes and output_format != 'JPEG':
            output.seek(0)
            output.truncate(0)
            img.save(output, format='JPEG', quality=70, optimize=True)
        
        return output.getvalue()

    def _extract_text2(self, content: bytes, filename: str, compress: bool = False) -> str:
        if filename.lower().endswith('.pdf'):
            try:
                pdf_file = BytesIO(content)
                text = ''
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += page_text + "\n"
                        else:
                            # Convert PDF page to image and use Vision API
                            img = page.to_image(resolution=300).original
                            img_byte_arr = io.BytesIO()
                            img.save(img_byte_arr, format='PNG')
                            img_content = img_byte_arr.getvalue()
                            
                            vision_image = vision.Image(content=img_content)
                            response = self.vision_client.text_detection(image=vision_image,image_context=self.image_context)
                            if response.text_annotations:
                                text_list = [text_annotation.description for text_annotation in response.text_annotations]
                                text += " ".join(text_list) + "\n"
            except Exception as e:
                raise HTTPException(
                    status_code=500,
                    detail=f"Error processing PDF: {str(e)}"
                )
        elif filename.lower().endswith('.docx'):
            doc_file = BytesIO(content)
            doc = Document(doc_file)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        elif filename.lower().endswith('.xlsx'):
            try:
                # Load Excel file into a pandas DataFrame
                excel_file = BytesIO(content)
                df = pd.read_excel(excel_file, sheet_name=None)  # Load all sheets
                
                # Extract text from all sheets
                text = ""
                for sheet_name, sheet_data in df.items():
                    text += f"Sheet: {sheet_name}\n"
                    text += sheet_data.to_string(index=False) + "\n\n"
            except Exception as e:
                raise HTTPException(
                    status_code=500,
                    detail=f"Error processing Excel file: {str(e)}"
                )
        else:
            # Compress image if needed
            if(compress):
                content = self._compress_image(content)
            vision_image = vision.Image(content=content)
            response = self.vision_client.document_text_detection(image=vision_image,image_context=self.image_context)
            
            if response.error.message:
                raise HTTPException(
                    status_code=500,
                    detail=f"Vision API Error: {response.error.message}"
                )
            
            if response.text_annotations:
                text = " ".join([text_annotation.description for text_annotation in response.text_annotations])
            else:
                text = ""
       
        return text