
from fastapi import UploadFile, HTTPException
from PIL import Image
import os
import json
from dotenv import load_dotenv
from io import BytesIO
from typing import List, Dict, Union, Optional, Any
from docx import Document
import io
import pandas as pd
from letter_format import LetterFormatStorage
from pydantic_ai import Agent
from pydantic_ai.models.google import GoogleModel
from pydantic_ai.providers.google import GoogleProvider
from pydantic_ai.settings import ModelSettings
from pydantic import BaseModel
import asyncio
from agentic_doc.parse import parse
# nest_asyncio imported conditionally inside __init__ to avoid uvloop conflicts
# Load environment variables from .env 
load_dotenv()



class DocumentAnalysisRequest(BaseModel):
    """Request model for document analysis"""
    payslip_text: Optional[str] = None
    contract_text: Optional[str] = None
    attendance_text: Optional[str] = None
    analysis_type: str = "report"
    context: Optional[str] = None

class DocumentAnalysisResponse(BaseModel):
    """Response model for document analysis"""
    legal_analysis: str
    # relevant_laws: List[Dict]
    # relevant_judgements: List[Dict]
    status: str
    analysis_type: str
    formatted_laws: Optional[str] = None
    formatted_judgements: Optional[str] = None
    documents: Optional[Dict] = None

def get_error_detail(e):
    try:
        # Try accessing the structured body if it's JSON-like
        if hasattr(e, "body"):
            import json
            body = e.body
            # If it's a string, parse it
            if isinstance(body, str):
                body = json.loads(body)
            return body.get("error", {}).get("message", str(e))

        # Fallback to using .message
        if hasattr(e, "message"):
            return str(e.message)

        return str(e)

    except Exception:
        return repr(e)


class DocumentProcessor:

    def _is_streamlit(self) -> bool:
        """Check if the code is running in a Streamlit environment."""
        try:
            import streamlit as st
            return hasattr(st, 'runtime') and st.runtime.exists()
        except (ImportError, AttributeError):
            return False
 
    def __init__(self):
        self.letter_format = LetterFormatStorage()
        
        # Initialize PydanticAI model - use Google Gemini
        gemini_api_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GOOGLE_CLOUD_VISION_API_KEY")
        # openai_api_key = os.getenv("OPENAI_API_KEY")
        # Use ModelSettings to set temperature to 0.0 (deterministic output)
        model_settings = ModelSettings(temperature=0.0)
        # if openai_api_key:
            # try:
            #     import openai
            #     openai.api_key = openai_api_key
            #     self.model_type = "gpt-5"
            #     self.model = "gpt-5"  # Use process_with_gpt5 for calls
            # except ImportError:
            #     print("⚠️ openai package not installed. GPT-5 support unavailable.")
        if gemini_api_key:
            # Create Google provider with API key
            google_provider = GoogleProvider(api_key=gemini_api_key)
            self.model = GoogleModel('gemini-2.5-pro', provider=google_provider, settings=model_settings)
            self.flashmodel = GoogleModel('gemini-2.5-flash', provider=google_provider, settings=model_settings)
            self.model_type = "gemini"
        else:
            raise Exception("GOOGLE_API_KEY or GOOGLE_CLOUD_VISION_API_KEY must be set in environment variables")
        # # Initialize PydanticAI Agent
        # # Configure Vision API
        # vision_api_key = os.getenv("GOOGLE_CLOUD_VISION_API_KEY")
        # if not vision_api_key:
        #     raise Exception("Google Cloud Vision API key not found. Please set GOOGLE_CLOUD_VISION_API_KEY in your .env file")
        
        # self.vision_client = ImageAnnotatorClient(client_options={"api_key": vision_api_key})
        # self.image_context = {"language_hints": ["he"]} 
        # Only apply nest_asyncio if not running under uvloop (e.g., not under uvicorn)
        _can_patch = True
        try:
            import asyncio
            loop = asyncio.get_event_loop()
            if 'uvloop' in str(type(loop)):
                _can_patch = False
        except Exception:
            pass
        if _can_patch and os.environ.get("USE_NEST_ASYNCIO", "0") == "1":
            import nest_asyncio
            nest_asyncio.apply()
        self.agent = Agent(
            model=self.model,
            output_type=str,
            system_prompt="""You are an expert legal document analyzer specializing in Israeli labor law compliance.

You will be provided with:
1. Relevant labor laws retrieved from a legal database
2. Relevant legal judgements and precedents retrieved from a legal database  
3. Document content to analyze (payslips, contracts, attendance records)
4. Additional context if provided

Your analysis must be based STRICTLY on the provided laws and judgements. Do not use external legal knowledge.

🚫 VERY IMPORTANT:
- Read the document carefully and remember its contents like wage per hour, working hours, etc. and Do not recalculate data like wages per hour, sick days, etc. unless the document provides exact values.
- Do not infer or estimate violations without clear proof in the payslip.
- Use **only** the documents provided (e.g., payslip data, employment contracts data, and attendance records data). **Do not extract or reuse any example values (e.g., 6000 ₪, 186 hours, 14 hours overtime) that appear in the legal texts or examples.**
- Do **not invent** missing data. If the document does not include sufficient detail for a violation (e.g., no overtime hours), **do not report a violation**.
- Do not hallucinate sick days, overtime hours, or absences
- think step by step and analyze the documents carefully. do not rush to conclusions.
- while calculating read the whole law and dont miss anything and explain the calculations step by step and how you arrived at the final amounts.

Always respond in Hebrew and follow the specific formatting requirements for each analysis type.

CRITICAL: When providing analysis, do NOT output template text or placeholders. Always replace ALL placeholders with real data from the analysis.""",
        )

        # Initialize Hebrew content extraction agent
        self.hebrew_agent = Agent(
            model=self.flashmodel,
            output_type=str,
            system_prompt="""You are an expert at extracting structured data from Hebrew text documents, particularly legal and administrative documents.

Your task is to extract specific fields from provided Hebrew text content. You must:
- Carefully read the text and identify values for the requested fields.
- Return only valid JSON with the exact field names as keys.
- If a field is not found, set its value to null.
- Do not add extra text or explanations - only the JSON object.
- Be precise and accurate in extracting values, especially for Hebrew text with special characters.""",
        )

        # Initialize Rule Engine components
        try:
            from engine.loader import RuleLoader
            from engine.evaluator import RuleEvaluator
            
            self.rule_loader = RuleLoader
            self.rule_evaluator = RuleEvaluator
            self.rules_data = None  # Will be loaded when needed
            print("✅ Rule Engine classes initialized successfully")
        except Exception as e:
            print(f"⚠️ Failed to initialize Rule Engine: {e}")
            self.rule_loader = None
            self.rule_evaluator = None
            self.rules_data = None
        
    async def export_to_excel(self, processed_result: dict) -> bytes:
        """
        Process the processed_result, use AI agent to extract employee name, overtime hours, salary, and all relevant data, and generate an Excel file for download.
        Returns the Excel file as bytes.
        """
        import pandas as pd
        import io
        import asyncio

        # Compose a prompt for the AI agent to extract structured data
        prompt = f"""
        Extract a table of all employees from the following legal document analysis results. For each employee, extract the following fields (if available):
        - שם העובד (Employee Name)
        - חודש עבודה (Work Month)
        - שנה (Year)
        - סה"כ ימי עבודה (Total Work Days)
        - סה"כ שעות עבודה (Total Work Hours)
        - תעריף שעתי (Hourly Rate)
        - משכורת 100 % (100% Salary)
        - ש.נ. 125% (Overtime 125%)
        - ש.נ. 150% (Overtime 150%)
        - ש.נ. 175% (Overtime 175%)
        - ש.נ. 200% (Overtime 200%)
        - שעות שבת (Shabbat Hours)
        - שעות חג (Holiday Hours)
        - חופשה (Vacation)
        - ימי מחלה (Sick Days)
        - הבראה (Convalescence)
        - נסיעות (Travel)
        - הפרשות מעסיק לפנסיה (Employer Pension Contributions)
        - הפרשות מעסיק קרן השתלמות (Employer Study Fund Contributions)
        - ניכוי עובד גמל (Employee Pension Deduction)
        - ניכוי עובד קרן התשלמות (Employee Study Fund Deduction)
        - תוספת וותק (Seniority Bonus)
        - ש.נ. גלובלי (Global Overtime)
        - כמות ש.נ גלובלי (Global Overtime Amount)
        - עמלה (Commission)
        - בונוס (Bonus)
        - עמלה (Commission)

        Return the result as a markdown table with columns for all the above fields (in the order listed). If there are multiple payslips or employees, include all rows.

        Payslip Text:
        {processed_result.get('payslip_text', '')}

        Contract Text:
        {processed_result.get('contract_text', '')}

        Attendance Text:
        {processed_result.get('attendance_text', '')}
        """

        # Simplified agent call
        async def get_table_sync():
            try:
                # For FastAPI/uvicorn - direct call (no nest_asyncio needed)
                if not self._is_streamlit():
                    return await self.agent.run(prompt, model_settings=ModelSettings(temperature=0.0))

                # For Streamlit - use nest_asyncio
                import nest_asyncio
                nest_asyncio.apply()
                loop = asyncio.get_event_loop()
                return loop.run_until_complete(self.agent.run(prompt, model_settings=ModelSettings(temperature=0.0)))
            except Exception as e:
                raise e

        result = await get_table_sync()
        print(result)
        table_md = result.output if hasattr(result, 'output') else str(result)

        # Parse the markdown table to a DataFrame
        import re
        from io import StringIO
        def markdown_table_to_df(md):
            # Find the first markdown table in the string
            lines = md.splitlines()
            table_lines = []
            in_table = False
            for line in lines:
                if '|' in line:
                    in_table = True
                    table_lines.append(line)
                elif in_table and line.strip() == '':
                    break
            if not table_lines:
                raise ValueError("No markdown table found in AI output.")
            # Remove markdown separator lines (e.g., | :--- | ... |)
            cleaned_lines = []
            for i, line in enumerate(table_lines):
                # Remove lines where all cells are dashes or colons (markdown separator)
                cells = [c.strip() for c in line.strip().split('|') if c.strip()]
                if all(re.match(r'^:?[-]+:?$', c) for c in cells):
                    continue
                cleaned_lines.append(line)
            table_str = '\n'.join(cleaned_lines)
            # Use pandas to read the markdown table
            try:
                import pandas as pd
                from io import StringIO
                df = pd.read_csv(StringIO(table_str), sep='|', engine='python')
                # Remove unnamed columns and whitespace
                df = df.loc[:, ~df.columns.str.contains('^Unnamed')]
                df.columns = [c.strip() for c in df.columns]
                # Use DataFrame.map for element-wise string strip (applymap is deprecated)
                for col in df.columns:
                    if df[col].dtype == object:
                        df[col] = df[col].map(lambda x: x.strip() if isinstance(x, str) else x)
                return df
            except Exception as e:
                raise ValueError(f"Failed to parse markdown table: {e}\nTable string:\n{table_str}")

        df = markdown_table_to_df(table_md)

        # Write DataFrame to Excel in memory
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name='EmployeeData')
        output.seek(0)
        return output.read()
    
    async def generate_ai_rule_checks(self, rule_description: str) -> List[Dict]:
        """
        Generate AI-powered rule checks based on description using dynamic parameters and available functions from engine
        
        Args:
            rule_description: Description of the labor law rule to implement
        
        Returns:
            List of check dictionaries, each containing id, condition, amount_owed, and violation_message
        """
        
        # Load dynamic parameters from engine
        from engine.dynamic_params import DynamicParams
        dynamic_params = DynamicParams.load()
        
        # Extract available parameters as lists of param names
        available_params = {}
        for section, params_list in dynamic_params.items():
            available_params[section] = [p['param'] for p in params_list]
        
        # Get available functions from engine (matching RuleEvaluator)
        available_functions = ["min", "max", "abs", "round"]
        
        # Format available parameters for the AI prompt
        params_summary = ""
        for section, params in available_params.items():
            params_summary += f"- {section}: {', '.join(params)}\n"
        
        functions_summary = ", ".join(available_functions)
        
        prompt = f"""
אתה מומחה ליצירת כללי בדיקה משפטיים עבור חוקי עבודה ישראליים.

עליך ליצור מערך של בדיקות מפורטות המבוסס על התיאור הבא:
{rule_description}

פרמטרים זמינים:
{params_summary}

פונקציות מתמטיות זמינות: {functions_summary}

הנחיות ליצירת הבדיקות:

1. **בדיקות (checks)** - צור מערך של בדיקות עם השדות הבאים לכל בדיקה:
   - id: מזהה ייחודי לבדיקה (באנגלית, עם קווים תחתונים)
   - condition: תנאי בדיקה באמצעות הפרמטרים הזמינים (שימוש בתחביר Python)
   - amount_owed: נוסחה לחישוב הסכום החסר (השתמש בפרמטרים ופונקציות זמינות)
   - violation_message: הודעת הפרה ברורה בעברית

2. **דוגמאות לתחביר**:
   - condition: "attendance.overtime_hours > 2"
   - amount_owed: "(contract.hourly_rate * 1.25 - payslip.overtime_rate) * min(attendance.overtime_hours, 2)"
   - השתמש ב-min(), max(), abs() לפי הצורך

3. **חשוב**:
   - השתמש רק בפרמטרים מהרשימה הזמינה
   - וודא שהנוסחאות מתמטיות נכונות
   - התמקד בחוקי עבודה ישראליים
   - צור בדיקות מרובות אם נדרש לבדוק מקרים שונים

החזר את התוצאה כ-JSON תקין בפורמט הבא:
[
  {{
    "id": "overtime_first_2_hours",
    "condition": "attendance.overtime_hours > 0",
    "amount_owed": "min(attendance.overtime_hours, 2) * (contract.hourly_rate * 0.25)",
    "violation_message": "השעות הנוספות הראשונות לא שולמו בתעריף הנכון"
  }},
  {{
    "id": "overtime_additional_hours", 
    "condition": "attendance.overtime_hours > 2",
    "amount_owed": "(attendance.overtime_hours - 2) * (contract.hourly_rate * 0.5)",
    "violation_message": "השעות הנוספות הנוספות לא שולמו בתעריף הנכון"
  }}
]
"""
        
        result = await self.agent.run(prompt, model_settings=ModelSettings(temperature=0.0))
        response_text = result.output if hasattr(result, 'output') else str(result)
        
        # Try to parse the JSON response
        try:
            generated_checks = json.loads(response_text)
            # Ensure it's a list
            if isinstance(generated_checks, list):
                return generated_checks
            else:
                # If it's a dict with checks key, extract the checks
                if isinstance(generated_checks, dict) and 'checks' in generated_checks:
                    return generated_checks['checks']
                else:
                    raise ValueError("AI response is not a valid checks array")
        except json.JSONDecodeError:
            # If JSON parsing fails, try to extract JSON array from the response
            import re
            json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
            if json_match:
                try:
                    generated_checks = json.loads(json_match.group())
                    if isinstance(generated_checks, list):
                        return generated_checks
                except json.JSONDecodeError:
                    pass
            
            # If all parsing fails, raise an exception
            raise ValueError(f"Failed to parse AI response as JSON array: {response_text}")

    async def create_report_with_rule_engine(self, payslip_data: List[Dict], attendance_data: List[Dict] = None, 
                                            contract_data: Dict = None, analysis_type: str = "rule_based") -> Dict:
        """
        Create legal analysis report using the rule engine for structured violation detection
        
        Args:
            payslip_data: List of payslip data dictionaries
            attendance_data: List of attendance data dictionaries
            contract_data: Contract data dictionary
            analysis_type: Type of analysis to perform
            
        Returns:
            Dictionary with rule-based analysis results
        """
        
        if not self.rule_evaluator:
            raise HTTPException(
                status_code=500,
                detail="Rule Engine not properly initialized"
            )
        
        try:
            print("🔄 Running rule-based analysis...")
            print(analysis_type)
            
            # Reload rules to stay updated with the data
            rules_path = os.path.join(os.path.dirname(__file__), 'rules/labor_law_rules.json')
            self.rules_data = self.rule_loader.load_rules(rules_path)
            
            # Validate and sanitize input data
            if not payslip_data or not isinstance(payslip_data, list):
                payslip_data = []
            
            if attendance_data is None or not isinstance(attendance_data, list):
                attendance_data = []
            
            if contract_data is None or not isinstance(contract_data, dict):
                contract_data = {}
            
            print(f"Processing {len(payslip_data)} payslips, {len(attendance_data)} attendance records")
            print(payslip_data)
            print(attendance_data)
            
            # Import required components
            from engine.evaluator import RuleEvaluator
            
            # Define context builder function (creates nested structure expected by rules)
            def build_context(payslip, attendance, contract):
                # Use dynamic param config to build context (matches main.py logic)
                from engine.dynamic_params import DynamicParams
                params = DynamicParams.load()
                context = {
                    'payslip': payslip,
                    'attendance': attendance,
                    'contract': contract
                }
                
                def coerce_value(value, param_type):
                    if value is None:
                        return None
                    if param_type == "number":
                        if isinstance(value, str):
                            value = value.strip()
                            # Remove commas for numbers like "5,300"
                            value = value.replace(',', '')
                            try:
                                # Try int first
                                if '.' not in value:
                                    return int(value)
                                else:
                                    return float(value)
                            except ValueError:
                                return value  # Keep as string if can't convert
                        elif isinstance(value, (int, float)):
                            return value
                        else:
                            return value
                    else:  # string
                        return str(value) if value is not None else None
                
                # Flatten all dynamic params for direct access with type coercion
                for section in ['payslip', 'attendance', 'contract']:
                    for p in params[section]:
                        raw_value = (locals()[section] or {}).get(p['param'], None)
                        param_type = p.get('type', 'number')
                        # Only set if we have a value (don't overwrite with None)
                        if raw_value is not None:
                            context[p['param']] = coerce_value(raw_value, param_type)
                # Add employee_id and month for legacy compatibility
                context['employee_id'] = context.get('employee_id', payslip.get('employee_id', None))
                context['month'] = context.get('month', payslip.get('month', None))
                return context
            
            # Step 1: Evaluate rules against the provided data using improved logic
            results = []
            
            # Check if we have any data to process
            if not payslip_data:
                print("⚠️ No payslip data provided for analysis")
                return {
                    'legal_analysis': 'לא סופקו נתוני תלושי שכר לניתוח',
                    'status': 'no_data',
                    'analysis_type': analysis_type,
                    'violations_count': 0,
                    'inconclusive_count': 0,
                    'compliant_count': 0,
                    'total_amount_owed': 0.0,
                    'violations_by_law': {},
                    'inconclusive_results': [],
                    'all_results': [],
                }
            
            # Process each payslip with month-based attendance matching
            for payslip in payslip_data:
                emp_id = payslip.get('employee_id', 'unknown')
                month = payslip.get('month', 'unknown')
                print(f"Evaluating payslip for employee {emp_id}, month {month}...")
                
                # Aggregate all attendance records for this month (improved logic)
                attendance_records = []
                if attendance_data:
                    attendance_records = [a for a in attendance_data if a.get('month') == month]
                
                if attendance_records:
                    # Sum numeric fields, keep others from the first record
                    aggregated = dict(attendance_records[0])
                    for key in aggregated:
                        if isinstance(aggregated[key], (int, float)):
                            aggregated[key] = sum(a.get(key, 0) for a in attendance_records 
                                                if isinstance(a.get(key, 0), (int, float)))
                    attendance = aggregated
                else:
                    attendance = {}
                
                # Get contract data (use first contract if multiple)
                contract = contract_data if contract_data else {}
                
                # Build context for rule evaluation
                context = build_context(payslip, attendance, contract)
                
                # Evaluate each rule
                for rule in self.rules_data['rules']:
                    # Check if rule is applicable for this month
                    if not RuleEvaluator.is_rule_applicable(rule, month):
                        print(rule["name"], "skipped for month", month)
                        continue
                    
                    # Evaluate rule checks
                    check_results, named_results = RuleEvaluator.evaluate_checks(rule['checks'], context)
                    
                    # Find violations (checks with amount >= 0)
                    violations = [cr for cr in check_results if cr['amount'] >= 0]
                    
                    # Check if any checks have missing fields
                    missing_fields_results = [cr for cr in check_results if cr.get('missing_fields', [])]
                    has_missing_fields = len(missing_fields_results) > 0
                    
                    # Collect all missing fields for this rule
                    all_missing_fields = set()
                    for cr in check_results:
                        if cr.get('missing_fields'):
                            all_missing_fields.update(cr['missing_fields'])
                    
                    # Include rule in results if there are violations OR missing fields
                    if violations or has_missing_fields:
                        result = {
                            'rule_id': rule['rule_id'],
                            'employee_id': emp_id,
                            'period': month,
                            'violations': violations,
                            'total_amount_owed': 0 if has_missing_fields else sum(v.get('amount', 0) for v in violations),
                            'check_results': check_results,
                            'has_missing_fields': has_missing_fields,
                            'missing_fields': sorted(list(all_missing_fields)),
                            'rule_name': rule.get('name', rule['rule_id']),
                            'law_reference': rule.get('law_reference', ''),
                            'description': rule.get('description', ''),
                            # Add rule metadata for UI display
                            'rule_checks': rule.get('checks', []),
                            'context_used': context,
                            'named_results': named_results
                        }
                        
                        # Add compliance status
                        if has_missing_fields:
                            result['compliance_status'] = 'inconclusive'
                        elif len(violations) == 0:
                            result['compliance_status'] = 'compliant'
                        else:
                            result['compliance_status'] = 'violation'
                        
                        results.append(result)
            
            # Step 2: Generate report based on analysis_type
            violation_results = [r for r in results if r['compliance_status'] == 'violation']
            inconclusive_results = [r for r in results if r['compliance_status'] == 'inconclusive']
            compliant_results = [r for r in results if r['compliance_status'] == 'compliant']
            
             # Calculate combined total (amount owed + penalties)
            total_amount_owed = sum(result.get('total_amount_owed', 0) for result in violation_results)
            
            # Handle different analysis types
            if analysis_type in ["violations_list", "easy", "table", "violation_count_table"]:
                # Use non-AI formatting for these types
                legal_analysis = self._format_rule_engine_results_non_ai(
                    violation_results, inconclusive_results, compliant_results, 
                    total_amount_owed, analysis_type
                )
            else:
                # Use AI formatting for complex types like "report", "combined", etc.
                    legal_analysis = await self._format_rule_engine_results_with_ai(
                        violation_results, inconclusive_results, compliant_results, 
                        total_amount_owed, analysis_type
                    )
            
            # Group results by law reference for return data
            violations_by_law = {}
            if violation_results:
                for result in violation_results:
                    law_ref = result.get('law_reference', 'לא מוגדר')
                    if law_ref not in violations_by_law:
                        violations_by_law[law_ref] = []
                    violations_by_law[law_ref].append(result)
            
           
            
            return {
                'legal_analysis': legal_analysis,
                'analysis_type': analysis_type,
                'violations_count': len(violation_results),
                'inconclusive_count': len(inconclusive_results),
                'compliant_count': len(compliant_results),
                'total_amount_owed': total_amount_owed,
                'violations_by_law': violations_by_law,
                'inconclusive_results': inconclusive_results,
            }
            
        except Exception as e:
            print(f"⚠️ Rule engine analysis failed: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Error in rule engine analysis: {get_error_detail(e)}"
            )

    def _format_rule_engine_results_non_ai(self, violation_results: List[Dict], 
                                           inconclusive_results: List[Dict], 
                                           compliant_results: List[Dict], 
                                           total_amount_owed: float, 
                                           analysis_type: str) -> str:
        """Format rule engine results without AI for specific analysis types"""
        
        if analysis_type == "violations_list":
            output = ""
            if violation_results:
                output += "## הפרות שזוהו\n\n"
                for result in violation_results[:4]:  # Limit to 4 violations
                    rule_name = result.get('rule_name', 'הפרה לא מזוהה')
                    law_ref = result.get('law_reference', 'לא מוגדר')
                    output += f"- **{rule_name}** - {law_ref}\n"
            if inconclusive_results:
                output += "\n## בדיקות לא חד-משמעיות (חסרים נתונים)\n\n"
                output += "החוקים הבאים לא חושבו בשל משתנים חסרים:\n"
                for result in inconclusive_results:
                    rule_name = result.get('rule_name', 'הפרה לא מזוהה')
                    law_ref = result.get('law_reference', 'לא מוגדר')
                    missing_fields = result.get('missing_fields', [])
                    output += f"- **{rule_name}** ({law_ref}): חסרים משתנים: {', '.join(missing_fields)}\n"
            if not violation_results and not inconclusive_results:
                output += "לא נמצאו הפרות\n"
            return output.strip()
        
        elif analysis_type == "easy":
            output = "# 📢 סיכום ההפרות\n\n"
            if violation_results:
                total_amount = 0
                for result in violation_results:
                    emp_id = result['employee_id']
                    period = result['period']
                    rule_name = result.get('rule_name', 'הפרה')
                    amount = result.get('total_amount_owed', 0)
                    if amount >= 0:
                        output += f"- ❌ **{rule_name}** ב{period} - **{amount:,.0f} ₪**\n"
                        total_amount += amount
                output += f"\n## 💰 סה\"כ: {total_amount:,.0f} ₪\n\n"
                output += "## 📝 מה לעשות עכשיו\n\n"
                output += "1. פנה/י למעסיק עם דרישה לתשלום הסכומים\n"
                output += "2. אם אין מענה – מומלץ לפנות לייעוץ משפטי\n"
            if inconclusive_results:
                output += "\n## בדיקות לא חד-משמעיות\n\n"
                output += "החוקים הבאים לא חושבו בשל משתנים חסרים:\n"
                for result in inconclusive_results:
                    rule_name = result.get('rule_name', 'הפרה')
                    law_ref = result.get('law_reference', 'לא מוגדר')
                    missing_fields = result.get('missing_fields', [])
                    output += f"- **{rule_name}** ({law_ref}): חסרים משתנים: {', '.join(missing_fields)}\n"
            if not violation_results and not inconclusive_results:
                output += "לא נמצאו הפרות\n"
            return output.strip()
        
        elif analysis_type == "table":
            output = "# טבלת הפרות\n\n"
            if violation_results:
                grouped = {}
                for result in violation_results:
                    key = f"{result['employee_id']} - {result['period']}"
                    if key not in grouped:
                        grouped[key] = []
                    grouped[key].append(result)
                for group_key, group_results in grouped.items():
                    output += f"## עובד {group_key}\n\n"
                    for i, result in enumerate(group_results, 1):
                        rule_name = result.get('rule_name', 'הפרה')
                        amount = result.get('total_amount_owed', 0)
                        hebrew_letter = chr(ord('א') + i - 1)
                        output += f"{hebrew_letter}. סכום של **{amount:.2f} ש\"ח** עבור **{rule_name}**\n"
                    output += "\n"
            if inconclusive_results:
                output += "## בדיקות לא חד-משמעיות\n\n"
                output += "החוקים הבאים לא חושבו בשל משתנים חסרים:\n"
                for result in inconclusive_results:
                    rule_name = result.get('rule_name', 'הפרה')
                    law_ref = result.get('law_reference', 'לא מוגדר')
                    missing_fields = result.get('missing_fields', [])
                    output += f"- **{rule_name}** ({law_ref}): חסרים משתנים: {', '.join(missing_fields)}\n"
            if not violation_results and not inconclusive_results:
                output += "לא נמצאו הפרות לארגון בטבלה\n"
            return output.strip()
        
        elif analysis_type == "violation_count_table":
            output = "# טבלת סיכום מקיפה\n\n"
            if violation_results:
                output += "| תלוש/מסמך | סוג הפרה | תקופה | סכום (₪) | מספר הפרות |\n"
                output += "---|---|---|---|---\n"
                total_amount = 0
                violation_counts = {}
                for result in violation_results:
                    emp_id = result['employee_id']
                    period = result['period']
                    rule_name = result.get('rule_name', 'הפרה')
                    amount = result.get('total_amount_owed', 0)
                    output += f"עובד {emp_id} | {rule_name} | {period} | {amount:,.2f} | 1\n"
                    total_amount += amount
                    if rule_name not in violation_counts:
                        violation_counts[rule_name] = {'count': 0, 'amount': 0}
                    violation_counts[rule_name]['count'] += 1
                    violation_counts[rule_name]['amount'] += amount
                output += f"---|---|---|---|---\n"
                output += f"סה\"כ | {len(violation_results)} הפרות | - | {total_amount:,.2f} | {len(violation_results)}\n\n"
                output += "פירוט לפי סוג הפרה:\n"
                for vtype, vdata in violation_counts.items():
                    output += f"• {vtype}: {vdata['count']} הפרות, סה\"כ {vdata['amount']:,.2f} ₪\n"
            if inconclusive_results:
                output += "\n## בדיקות לא חד-משמעיות\n\n"
                output += "החוקים הבאים לא חושבו בשל משתנים חסרים:\n"
                for result in inconclusive_results:
                    rule_name = result.get('rule_name', 'הפרה')
                    law_ref = result.get('law_reference', 'לא מוגדר')
                    missing_fields = result.get('missing_fields', [])
                    output += f"- **{rule_name}** ({law_ref}): חסרים משתנים: {', '.join(missing_fields)}\n"
            if not violation_results and not inconclusive_results:
                output += "לא נמצאו הפרות נגד חוקי העבודה שסופקו.\n"
            return output.strip()
        
        else:
            # Default rule-based report format
            return self._format_default_rule_engine_report(
                violation_results, inconclusive_results, compliant_results, total_amount_owed
            )
    
    def _format_default_rule_engine_report(self, violation_results: List[Dict], 
                                          inconclusive_results: List[Dict], 
                                          compliant_results: List[Dict], 
                                          total_amount_owed: float) -> str:
        """Format default rule engine report"""
        report_sections = []
        
        # Summary section
        report_sections.append("# דוח ניתוח משפטי - מנוע חוקים")
        report_sections.append(f"## סיכום כללי")
        report_sections.append(f"- **סה\"כ הפרות שזוהו:** {len(violation_results)}")
        report_sections.append(f"- **בדיקות לא חד-משמעיות (חסרים נתונים):** {len(inconclusive_results)}")
        report_sections.append(f"- **בדיקות תקינות:** {len(compliant_results)}")
        report_sections.append(f"- **סה\"כ חסר בתשלום:** {total_amount_owed:,.2f} ₪")
        
        # Violations section
        if violation_results:
            report_sections.append(f"\n## הפרות שזוהו")
            
            for result in violation_results:
                report_sections.append(f"\n### {result.get('rule_name', 'הפרה לא מזוהה')}")
                report_sections.append(f"**עובד:** {result['employee_id']}")
                report_sections.append(f"**תקופה:** {result['period']}")
                report_sections.append(f"**תיאור:** {result.get('description', 'לא זמין')}")
                if result.get('law_reference'):
                    report_sections.append(f"**בסיס חוקי:** {result['law_reference']}")
                
                # Show violation details
                if result['violations']:
                    report_sections.append("**פרטי ההפרות:**")
                    for violation in result['violations']:
                        if violation['amount'] >= 0:
                            report_sections.append(f"  - {violation.get('description', 'הפרה')}: {violation['amount']:,.2f} ₪")
                
                if result['total_amount_owed'] >= 0:
                    report_sections.append(f"**סכום חסר בתשלום:** {result['total_amount_owed']:,.2f} ₪")
                
                # if result['penalty_amount'] > 0:
                #     report_sections.append(f"**סכום חסר:** {result['total_amount_owed']:,.2f} ₪")
        
        # Missing data section
        if inconclusive_results:
            report_sections.append(f"\n## בדיקות לא חד-משמעיות (חסרים נתונים)")
            report_sections.append("הבדיקות הבאות לא ניתן היה להשלים בשל חוסר במידע:")
            
            for result in inconclusive_results:
                report_sections.append(f"\n### {result.get('rule_name', 'בדיקה לא מזוהה')}")
                report_sections.append(f"**עובד:** {result['employee_id']}")
                report_sections.append(f"**תקופה:** {result['period']}")
                if result['missing_fields']:
                    report_sections.append(f"**שדות חסרים:** {', '.join(result['missing_fields'])}")
        
        # Compliant section (optional, brief)
        if compliant_results:
            report_sections.append(f"\n## בדיקות תקינות")
            report_sections.append(f"נמצאו {len(compliant_results)} בדיקות שעברו בהצלחה ללא הפרות.")
        
        # No issues found
        if not violation_results and not inconclusive_results:
            report_sections.append("\n## לא זוהו הפרות")
            report_sections.append("על פי הבדיקה שבוצעה באמצעות מנוע החוקים, לא זוהו הפרות של חוקי העבודה במסמכים שנבדקו.")
        
        # Recommendations section
        if violation_results:
            report_sections.append("\n## המלצות לתיקון")
            recommendations = []
            
            for result in violation_results:
                if result['total_amount_owed'] >= 0:
                    recommendations.append(f"להשלים תשלום חסר של {result['total_amount_owed']:,.2f} ₪ לעובד {result['employee_id']} עבור תקופה {result['period']}")
            
            for i, recommendation in enumerate(recommendations, 1):
                report_sections.append(f"{i}. {recommendation}")
        
        return "\n".join(report_sections)
    
    async def _format_rule_engine_results_with_ai(self, violation_results: List[Dict], 
                                                 inconclusive_results: List[Dict], 
                                                 compliant_results: List[Dict], 
                                                 total_amount_owed: float, 
                                                 analysis_type: str) -> str:
        """Format rule engine results using AI for complex analysis types"""
        
        # Build prompt based on analysis type
        if analysis_type == "report":
            instructions = self._get_report_instructions()
        elif analysis_type == "combined":
            instructions = self._get_combined_instructions()
        elif analysis_type == "claim":
            instructions = self._get_claim_instructions()
        elif analysis_type == "warning_letter":
            instructions = self._get_warning_letter_instructions()
        else:
            # Default to report format
            instructions = self._get_report_instructions()
        
        prompt = f"""
נתוני הפרות ממנוע החוקים:

סיכום כמותי:
- סה"כ הפרות: {len(violation_results)}
- בדיקות לא חד-משמעיות: {len(inconclusive_results)}
- בדיקות תקינות: {len(compliant_results)}
- סה"כ חסר בתשלום: {total_amount_owed:,.2f} ₪
פירוט הפרות:
{self._format_violations_for_ai_prompt(violation_results)}

בדיקות לא חד-משמעיות:
{self._format_inconclusive_for_ai_prompt(inconclusive_results)}

הנחיות נוספות:
- אם אין הפרות אך יש בדיקות לא חד-משמעיות, ציין 'לא נמצאו הפרות עד כה, אך החוקים הבאים לא חושבו בשל משתנים חסרים:' ופרט את החוקים והשדות החסרים.
- תמיד כלול את הבדיקות לא חד-משמעיות בדוח, גם אם יש הפרות.
- השתמש בכותרות ברורות ובמרווחים מתאימים בין חלקי הדוח השונים.

---

{instructions}
"""
        
        try:
            result = await self.agent.run(prompt, model_settings=ModelSettings(temperature=0.0))
            return result.output if hasattr(result, 'output') else str(result)
        except Exception as e:
            # Fallback to non-AI formatting if AI fails
            print(f"AI formatting failed, using fallback: {e}")
            return self._format_default_rule_engine_report(
                violation_results, inconclusive_results, compliant_results, total_amount_owed
            )
    
    def _format_violations_for_ai_prompt(self, violation_results: List[Dict]) -> str:
        """Format violations data for AI prompt"""
        if not violation_results:
            return "לא נמצאו הפרות"
        formatted = []
        for result in violation_results:
            underpaid = result.get('total_amount_owed', 0)
            formatted.append(f"""
הפרה: {result.get('rule_name', 'לא מוגדר')}
עובד: {result['employee_id']}
תקופה: {result['period']}
בסיס חוקי: {result.get('law_reference', 'לא מוגדר')}
תיאור: {result.get('description', 'לא זמין')}
סכום חסר: {underpaid:,.2f} ₪
פרטי הפרות: {', '.join([f"{v.get('description', 'הפרה')}: {v['amount']:,.2f} ₪" for v in result.get('violations', []) if v['amount'] >= 0])}
""")
        return "\n".join(formatted)
    
    def _format_inconclusive_for_ai_prompt(self, inconclusive_results: List[Dict]) -> str:
        """Format inconclusive results for AI prompt"""
        if not inconclusive_results:
            return "אין בדיקות לא חד-משמעיות"
        
        formatted = []
        for result in inconclusive_results:
            formatted.append(f"""
בדיקה: {result.get('rule_name', 'לא מוגדר')}
עובד: {result['employee_id']}
תקופה: {result['period']}
שדות חסרים: {', '.join(result.get('missing_fields', []))}
""")
        
        return "\n".join(formatted)

    async def process_document(self, files: Union[UploadFile, List[UploadFile]], doc_types: Union[str, List[str]], compress: bool = False) -> Dict[str, Union[List[Dict], Dict]]:
        """Process uploaded documents and extract text concurrently"""
        if not files:
            raise HTTPException(
                status_code=400,
                detail="At least one document must be provided"
            )
        # Handle single file case
        if isinstance(files, UploadFile):
            files = [files]
            doc_types = [doc_types]
        elif not isinstance(doc_types, list):
            doc_types = [doc_types] * len(files)

        async def process_single(file, doc_type, idx):
            print(f"Processing file: {file.filename} as type: {doc_type}")
            # Properly read the file content from UploadFile
            if hasattr(file, 'read'):
                content = await file.read()  # Await the async read method
                if isinstance(content, bytes):
                    pass  # Good, we got bytes
                else:
                    # If read() doesn't return bytes, try file.file.read()
                    if hasattr(file, 'file'):
                        content = file.file.read()
                    else:
                        raise ValueError(f"Cannot read content from file {file.filename}")
            else:
                raise ValueError(f"File object {file.filename} does not support reading")

            print(f"Read {len(content)} bytes from {file.filename}")
            # Call updated _extract_text2 with doc_type parameter
            extracted_data = await self._extract_text2(content, file.filename, doc_type, compress)
            return (doc_type.lower(), idx, extracted_data)

        tasks = [process_single(file, doc_type, idx+1) for idx, (file, doc_type) in enumerate(zip(files, doc_types))]
        results = await asyncio.gather(*tasks)

        # Initialize arrays for structured data
        payslip_data = []
        contract_data = []
        attendance_data = []
        
        # Initialize separate counters for each document type
        payslip_counter = 0
        contract_counter = 0
        attendance_counter = 0

        for doc_type, _, extracted_data in results:
            if doc_type == "payslip":
                payslip_counter += 1
                structured_payslip = extracted_data.get('structured_data', {})
                # Ensure required fields for rule engine
                structured_payslip['employee_id'] = structured_payslip.get('employee_id')
                structured_payslip['month'] = structured_payslip.get('month')  # Default fallback
                structured_payslip['document_number'] = payslip_counter
                payslip_data.append(structured_payslip)
                
            elif doc_type == "contract":
                contract_counter += 1
                structured_contract = extracted_data.get('structured_data', {})
                structured_contract['employee_id'] = structured_contract.get('employee_id')
                structured_contract['document_number'] = contract_counter
                contract_data.append(structured_contract)
                
            elif doc_type == "attendance":
                attendance_counter += 1
                structured_attendance = extracted_data.get('structured_data', {})
                structured_attendance['employee_id'] = structured_attendance.get('employee_id')
                structured_attendance['month'] = structured_attendance.get('month')  # Default fallback
                structured_attendance['document_number'] = attendance_counter
                attendance_data.append(structured_attendance)

        # Return in the format expected by create_report_with_rule_engine
        return {
            "payslip_data": payslip_data,
            "contract_data": contract_data[0] if contract_data else {},  # Single contract dict as expected
            "attendance_data": attendance_data
        }

    async def qna(self, report: str, questions: str) -> str:
        """Generate answer of queries based on the provided document content."""
        prompt = f"""
להלן דוח ניתוח משפטי משולב של מסמכי העובד:

{report}

---

שאלות:
{questions}

---

אנא ספק תשובות לשאלות לעיל.
"""
        result = await self.agent.run(prompt, model_settings=ModelSettings(temperature=0.0))
        return result.output

    def _get_report_instructions(self) -> str:
        return """
INSTRUCTIONS:
1. If no labor laws are provided, respond with: "אין חוקים לעבודה זמינים לניתוח התאמה." in Hebrew.
2. If labor laws exist, analyze the documents ONLY against the provided laws.
3. ONLY refer to the judgements and their results provided above for legal analysis - do not use external cases or knowledge.
4. If no judgements are provided, respond with: "לא קיימות החלטות משפטיות זמינות לניתוח." in Hebrew.

🎯 GOAL: Help the business owner understand exactly what went wrong and where they made mistakes in their employment practices.

For each payslip provided, analyze and identify violations. For each violation found in each payslip, format the response EXACTLY as shown below, with each section on a new line and proper spacing:

Violation Format Template:

🚨 [VIOLATION TITLE - What the employer did wrong]

📋 מה קרה בפועל:
[Describe exactly what the employer did or failed to do, with specific details from the documents]

⚖️ מה היה צריך לקרות לפי החוק:
[Explain what should have been done according to the law, with specific legal requirements]

📖 בסיס חוקי:
[LAW REFERENCE AND YEAR FROM PROVIDED LAWS - cite the specific law that was violated]

💰 הנזק הכספי:
[Calculate the exact financial impact - how much the employee lost due to this violation]

🏛️ תקדימים משפטיים:
[SIMILAR CASES OR PRECEDENTS FROM PROVIDED JUDGEMENTS](Refer *only* to the retrieved judgements. If a relevant judgement is found, describe the case and its result. If no relevant judgement is found, state "לא נמצאו תקדימים רלוונטיים בפסקי הדין שסופקו.")

⚠️ השלכות אפשריות:
[Explain potential legal consequences and risks for the employer]

✅ מה לעשות כדי לתקן:
[Specific actionable steps the employer should take to fix this violation and prevent it in the future]

---

SUMMARY FOR EMPLOYER:
After completing the violation analysis, provide a clear summary for the business owner:

=== סיכום למעסיק - איפה טעיתם ומה לעשות ===

📊 טבלת הפרות שזוהו:
תלוש/מסמך | מה עשיתם לא נכון | כמה זה עולה (₪) | מה לעשות עכשיו
[Add rows with actual data showing: document name | specific mistake made | financial cost | corrective action needed]

💸 סה"כ עלות הטעויות: [total amount] ₪

🔧 צעדים מיידיים לתיקון:
1. [First immediate action needed]
2. [Second immediate action needed]
3. [Third immediate action needed]

📋 איך למנוע טעויות בעתיד:
• [Prevention measure 1]
• [Prevention measure 2]
• [Prevention measure 3]

IMPORTANT:
- Always Respond in Hebrew
- Focus on helping the employer understand their mistakes and how to fix them
- Use clear, business-friendly language that explains the "why" behind each violation
- Format each violation with proper spacing and line breaks as shown above
- Analyze each payslip separately and clearly indicate which payslip the violations belong to
- Separate multiple violations with '---'
- If no violations are found against the provided laws in a payslip, respond with: "לא נמצאו הפרות בתלוש מספר" followed by the payslip number in hebrew
- If no violations are found in any payslip, respond with: "לא נמצאו הפרות נגד חוקי העבודה שסופקו." in hebrew
- DO NOT output template text or placeholders - use real data from the analysis
- Replace ALL placeholders with actual information from the documents
- Make it clear to the employer what they did wrong and how to prevent it happening again
"""

    def _get_violations_list_instructions(self) -> str:
        return """
CRITICAL INSTRUCTIONS - FOLLOW EXACTLY:

You MUST respond with ONLY the format below. NO additional text, explanations, analysis, or commentary.

If violations found, use EXACTLY this format:

**הפרות שזוהו:**

• [הפרה קצרה] - [חוק]
• [הפרה קצרה] - [חוק]
• [הפרה קצרה] - [חוק]

If NO violations found, respond ONLY with: "לא נמצאו הפרות"

ABSOLUTE RESTRICTIONS:
❌ NO "Legal analysis" headers
❌ NO explanations or breakdowns  
❌ NO calculations or formulas
❌ NO "Analysis of the violation" sections
❌ NO "method of calculating" text
❌ NO English words
❌ NO additional paragraphs
❌ NO template text like [Violation Title]

✅ ONLY Hebrew
✅ ONLY the exact format above
✅ Maximum 4 violations
✅ Each violation = ONE short line

EXAMPLE of correct output:
**הפרות שזוהו:**

• לא שולם שכר מינימום - חוק שכר מינימום
• לא שולמו שעות נוספות - חוק שעות עבודה ומנוחה
• לא הופקדה פנסיה - צו הרחבה פנסיה חובה

If you write ANYTHING beyond this format, you FAILED.
"""

    def _get_profitability_instructions(self) -> str:
        return """
INSTRUCTIONS:
1. Analyze the provided documents and identify potential labor law violations, based *exclusively on the retrieved LABOR LAWS and JUDGEMENTS*.
2. For each violation, refer *only* to the retrieved judgements to identify similar legal cases and their outcomes (both successful and unsuccessful).
3. If similar cases (from the retrieved judgements) were unsuccessful:
   - Explain why the cases were unsuccessful based on the retrieved judgement information.
   - Provide a clear recommendation against pursuing legal action based on this.
   - List potential risks and costs.

4. If similar cases (from the retrieved judgements) were successful, calculate using information from those judgements if available:
   - Average compensation amount from successful retrieved cases.
   - Estimated legal fees (30% of potential compensation).
   - Tax implications (25% of net compensation).
   - Time and effort cost estimation.

Provide the analysis in the following format:

ניתוח כדאיות כלכלית:

הפרות שזוהו (על בסיס החוקים שסופקו):
[List identified violations]

תקדימים משפטיים (מתוך פסקי הדין שסופקו):
[Similar cases from retrieved judgements with outcomes - both successful and unsuccessful]

במקרה של תקדימים שליליים (מתוך פסקי הדין שסופקו):
- סיבות לדחיית התביעות: [REASONS BASED ON RETRIEVED JUDGEMENTS]
- סיכונים אפשריים: [RISKS]
- המלצה: לא מומלץ להגיש תביעה בשל [EXPLANATION BASED ON RETRIEVED JUDGEMENTS]

במקרה של תקדימים חיוביים (מתוך פסקי הדין שסופקו):
ניתוח כספי:
- סכום פיצוי ממוצע (מפסיקות דין שסופקו): [AMOUNT] ₪
- עלות משוערת של עורך דין (30%): [AMOUNT] ₪
- השלכות מס (25% מהסכום נטו): [AMOUNT] ₪
- סכום נטו משוער: [AMOUNT] ₪

המלצה סופית:
[Based on analysis of both successful and unsuccessful cases from the retrieved judgements, provide clear recommendation]

SUMMARY TABLE:
After completing the profitability analysis, provide a summary table:
- Use the heading: === טבלת סיכום כלכלי ===
- Create columns for: פריט | סכום (₪)
- Add rows with actual calculated amounts for:
  * סה"כ פיצוי משוער
  * עלות עורך דין (30%)
  * מס (25%)
- End with: סכום נטו משוער: [calculated amount]

CRITICAL: Replace ALL placeholders with actual calculated values from the analysis. Do not output template text.
""""""
"""

    def _get_warning_letter_instructions(self) -> str:
        format_content = self.letter_format.get_format().get('content', '')

        return f"""
INSTRUCTIONS:
1. Analyze the provided documents for labor law violations *based exclusively on the retrieved LABOR LAWS and JUDGEMENTS*.
2. If violations are found, generate a formal warning letter using the provided template.
3. If no violations are found, respond with: "לא נמצאו הפרות המצדיקות מכתב התראה." in Hebrew.

Warning Letter Template:
{format_content}

Please generate the warning letter in Hebrew with the following guidelines:
- Replace [שם המעסיק] with the employer's name from the documents
- Replace [פרטי ההפרות] with specific details of each violation found (based on retrieved laws)
- Replace [הפניות לחוקים] with relevant labor law citations *from the retrieved LABOR LAWS*.
- Replace [מועד] with a reasonable timeframe for corrections (typically 14 days)
- Maintain a professional and formal tone throughout
- Include all violations found in the analysis (based on retrieved laws)
- Format the letter according to the provided template structure
"""

    def _get_easy_instructions(self) -> str:
        return """
🔒 מטרה: צור סיכום קצר וברור של ההפרות בתלושי השכר של העובד.
📌 כללים מחייבים:
    1. כתוב בעברית בלבד – אל תשתמש באנגלית בכלל.
    2. עבור כל חודש הצג את ההפרות בשורות נפרדות, כל שורה בפורמט הבא:
❌ [סוג ההפרה בקצרה] – [סכום בש"ח עם ₪, כולל פסיק לאלפים]
לדוגמה: ❌ לא שולם החזר נסיעות בפברואר 2025 – 250 ₪
    3. אם יש מספר רכיבי פנסיה (עובד/מעסיק/בריאות) בחודש מסוים – חבר אותם לסכום אחד של פנסיה באותו החודש.
    4. כל הסכומים יוצגו עם פסיקים לאלפים ועם ₪ בסוף.
    5. חישוב הסכום הכולל יופיע בשורה נפרדת:
💰 סה"כ: [total amount] ₪
    6. הוסף המלצה בסוף:
📝 מה לעשות עכשיו:
פנה/י למעסיק עם דרישה לתשלום הסכומים.
אם אין מענה – מומלץ לפנות לייעוץ משפטי.
📍 הנחיות נוספות:
    • אין לכתוב מספרים בלי הקשר, כל שורה חייבת להיות מלווה בחודש.
    • מיזוג שורות: אם באותו חודש יש כמה רכיבים של פנסיה – מיזג אותם לשורה אחת.
    • הסר שורות ללא סכום ברור.
    • ניסוח פשוט, ללא מינוחים משפטיים, הבהרות או הסברים טכניים.
    • אין לציין "רכיב עובד", "רכיב מעסיק", "לא הופקד" – במקום זאת כתוב: "לא שולמה פנסיה".
🎯 פלט רצוי:
    • שורות מסודרות לפי חודשים
    • אין כפילויות
    • סכומים מדויקים בלבד
    • ניסוח ברור ומובן
    • עברית בלבד

🧪 Example of desired output:
📢 סיכום ההפרות:
❌ לא שולם עבור שעות נוספות בנובמבר 2024 – 495 ₪
❌ לא שולמה פנסיה בנובמבר 2024 – 750 ₪
❌ לא שולמה פנסיה בדצמבר 2024 – 1,221 ₪
❌ לא שולמה פנסיה בינואר 2025 – 831 ₪
❌ לא שולם החזר נסיעות בפברואר 2025 – 250 ₪
❌ לא שולמה פנסיה בפברואר 2025 – 858 ₪
❌ לא שולמה פנסיה במרץ 2025 – 866 ₪
💰 סה"כ: 5,271 ₪
📝 מה לעשות עכשיו:
פנה/י למעסיק עם דרישה לתשלום הסכומים.
אם אין מענה – מומלץ לפנות לייעוץ משפטי.

SUMMARY TABLE:
After completing the easy summary, provide a simple visual table:

=== טבלת סיכום ויזואלי ===
חודש/שנה            | סוג הפרה           | סכום (₪)
נובמבר 2024          | שעות נוספות        | 495
נובמבר 2024          | פנסיה             | 750
דצמבר 2024           | פנסיה             | 1,221
ינואר 2025           | פנסיה             | 831
פברואר 2025          | נסיעות            | 250
פברואר 2025          | פנסיה             | 858
מרץ 2025             | פנסיה             | 866
-----------------------------------------
סה"כ: 5,271 ₪
"""

    def _get_table_instructions(self) -> str:
        return """
אתה עוזר משפטי מיומן. עליך לנתח את רשימת ההפרות ולהפיק רשימת תביעות מסודרת לפי מסמך (לדוגמה: תלוש שכר מס' 1, מסמך שימוע, מכתב פיטורין וכו').

הנחיות:

1. סדר את התביעות לפי מסמך: כל קבוצה מתחילה בכותרת כמו "תלוש שכר מס' 4 – 05/2024:".

2. תחת כל כותרת, צור רשימה ממוספרת באותיות עבריות (א., ב., ג. וכו').

3. השתמש במבנה הקבוע הבא:
   א. סכום של [amount] ש"ח עבור [תיאור קצר של ההפרה].

4. השתמש בפורמט מספרים עם פסיקים לאלפים ושתי ספרות אחרי הנקודה (למשל: 1,618.75 ש"ח).

5. כתוב "ש"ח" אחרי הסכום — לא ₪.

6. אל תסכם את כלל ההפרות – הצג רק את הפירוט לפי מסמך.

7. אל תוסיף בולטים, טבלאות או הערות.

נתוני הקלט לדוגמה:
❌ אי תשלום שעות נוספות (תלוש 6, 11/2024) – 495 ש"ח
❌ העדר רכיב מעביד לפנסיה (תלוש 6, 11/2024) – 750 ש"ח
❌ העדר רכיב עובד לפנסיה (תלוש 7, 12/2024) – 396 ש"ח
❌ העדר נסיעות (תלוש 9, 02/2025) – 250 ש"ח

פלט נדרש לדוגמה:

תלוש שכר מס' 6 – 11/2024:
א. סכום של 495.00 ש"ח עבור אי תשלום שעות נוספות.
ב. סכום של 750.00 ש"ח עבור העדר רכיב מעביד לפנסיה.

תלוש שכר מס' 7 – 12/2024:
א. סכום של 396.00 ש"ח עבור העדר רכיב עובד לפנסיה.

תלוש שכר מס' 9 – 02/2025:
א. סכום של 250.00 ש"ח עבור העדר תשלום עבור נסיעות.

החזר את הפלט בפורמט זה בלבד, מופרד לפי כל מסמך.

SUMMARY TABLE:
After completing the organized document list, provide a simple summary table:

=== טבלת סיכום לפי מסמך ===
מסמך                | מספר הפרות         | סכום כולל (ש"ח)
תלוש שכר מס' 6       | 2                  | 1,245.00
תלוש שכר מס' 7       | 1                  | 396.00  
תלוש שכר מס' 9       | 1                  | 250.00
------------------------------------------
סה"כ                | 4                  | 1,891.00
"""

    def _get_claim_instructions(self) -> str:
        return """
משימה:
כתוב טיוטת כתב תביעה לבית הדין האזורי לעבודה, בהתאם למבנה המשפטי הנהוג בישראל.

נתונים:
השתמש במידע מתוך המסמכים שצורפו (כגון תלושי שכר, הסכמי עבודה, הודעות פיטורין, שיחות עם המעסיק) ובממצאים שנמצאו בניתוח קודם של ההפרות.

פורמט לכתיבה:
1. כותרת: "בית הדין האזורי לעבודה ב[שם עיר]"
2. פרטי התובע/ת:
   - שם מלא, ת"ז, כתובת, טלפון, מייל
3. פרטי הנתבע/ת (מעסיק):
   - שם החברה/מעסיק, ח.פ./ת"ז, כתובת, טלפון, מייל (אם קיים)
4. כותרת: "כתב תביעה"
5. פתיח משפטי:
   בית הדין הנכבד מתבקש לזמן את הנתבע/ת לדין ולחייבו/ה לשלם לתובע/ת את הסכומים המפורטים, מהנימוקים הבאים:
6. סעיף 1 – רקע עובדתי:
   תיאור תקופת העבודה, תפקידים, תאריך תחילה וסיום (אם רלוונטי), מהות יחסי העבודה, מקום העבודה.
7. סעיף 2 – עילות התביעה (לפי הפרות):
   לכל הפרה:
     - איזה חוק הופֵר (לדוג': חוק שכר מינימום, חוק שעות עבודה ומנוחה וכו')
     - פירוט העובדות והתקופה הרלוונטית
     - סכום הפיצוי או הנזק
     - אסמכתאות משפטיות אם רלוונטי
8. סעיף 3 – נזקים שנגרמו:
   סכומים כספיים (בפירוט), נזק לא ממוני אם קיים (עוגמת נפש).
9. סעיף 4 – סעדים מבוקשים:
   תשלום הפרשים, פיצויים, ריבית והצמדה, הוצאות משפט, שכר טרחת עו"ד, וכל סעד אחר.
10. סעיף 5 – סמכות שיפוט:
   ציין סמכות בית הדין לעבודה לפי חוק בית הדין לעבודה תשכ"ט–1969.
11. סעיף 6 – דיון מקדים/הליך מהיר (אם רלוונטי).
12. סעיף 7 – כתובות להמצאת כתבי בי-דין:
   כתובת התובע והנתבע.
13. סיום:
   חתימה, תאריך, רשימת נספחים תומכים (תלושי שכר, מכתבים וכו').

שפה:
- כתוב בעברית משפטית, רשמית ומסודרת.
- סדר את התביעה עם כותרות, סעיפים ממוספרים ורווחים ברורים.
- אל תמציא עובדות או חוקים. אם חסר מידע – השאר שדה ריק לצורך השלמה ע"י המשתמש.

אזהרה:
בסיום, כתוב משפט: "הטיוטה נכתבה אוטומטית לצורך עזרה ראשונית בלבד. מומלץ לפנות לעורך/ת דין מוסמך/ת לפני הגשה לבית הדין."

SUMMARY TABLE:
After completing the claim draft, provide a simple summary table:

=== טבלת סיכום כתב התביעה ===
עילת תביעה           | תקופה               | סכום נתבע (₪)
אי תשלום שעות נוספות  | נובמבר 2024          | 495
אי הפקדת פנסיה       | נובמבר 2024-מרץ 2025  | 3,525
אי החזר נסיעות       | פברואר 2025          | 250
---------------------------------------------
סה"כ תביעה ראשית: 4,270 ₪
ריבית והצמדה: לפי החלטת בית הדין
הוצאות משפט: לפי החלטת בית הדין
"""

    def _get_combined_instructions(self) -> str:
        return """
COMPREHENSIVE ANALYSIS INSTRUCTIONS:

📌 OVERVIEW
You are analyzing employee documents(payslips, employment contracts, and attendance records) for legal violations based **only** on the provided labor laws and judgements. Do **not** use any default knowledge, assumptions, or examples from the law unless explicitly confirmed in the payslip.
and explain step by step your thought process, calculations, and legal references.
labor laws and judgements are different from employee documents dont use example values from the legal texts or examples. use only real data from the documents provided.

PART 1 - VIOLATION IDENTIFICATION AND ANALYSIS:

1. If no labor laws are provided, respond with: "אין חוקים לעבודה זמינים לניתוח התאמה." in Hebrew.
2. If labor laws exist, analyze the payslip ONLY against the laws provided.
3. Use **only** the documents provided (e.g., payslip data, employment contracts data, and attendance records data). **Do not extract or reuse any example values (e.g., 6000 ₪, 186 hours, 14 hours overtime) that appear in the legal texts or examples.**
4. Do **not invent** missing data. If the document does not include sufficient detail for a violation (e.g., no overtime hours), **do not report a violation**.
5. If no judgements are provided, respond with: "לא קיימות החלטות משפטיות זמינות לניתוח." in Hebrew.


PART 2 - DETAILED PROFESSIONAL ANALYSIS FORMAT:

For each violation found, provide the following, using actual values from the document:

הפרה: [כותרת ההפרה]
[תיאור מפורט של ההפרה, כולל תאריכים רלוונטיים, שעות עבודה בפועל, שכר שעתי, ימי מחלה מאושרים וכדומה. התבסס רק על הנתונים הקיימים בתלוש, ובאופן בלעדי על החוקים שהוצגו.]

[הפניה לחוק הרלוונטי ושנה מתוך החוקים שסופקו]

[אם קיימים – תאר תקדימים משפטיים מתוך פסקי הדין שסופקו. אם אין – כתוב "לא נמצאו תקדימים רלוונטיים בפסקי הדין שסופקו."]

[השלכות משפטיות ופעולות מומלצות]

סה"כ חוב להפרה זו: [סכום] ₪

---

PART 3 - ORGANIZED TABLE FORMAT BY DOCUMENT:

רשימת תביעות מסודרת לפי מסמך:

תלוש שכר מס' [מספר] – [חודש/שנה]:
א. סכום של [amount] ש"ח עבור [תיאור קצר של ההפרה].
ב. סכום של [amount] ש"ח עבור [תיאור קצר של ההפרה].

---

PART 4 - FINAL SUMMARY:

סה"כ תביעה משפטית (לא כולל ריבית): [סכום כולל] ₪  
אסמכתאות משפטיות: [רשימת שמות החוקים הרלוונטיים מתוך החוקים שסופקו]

---

PART 5 - COMPREHENSIVE SUMMARY TABLE:

=== טבלת סיכום מקיפה ===

Create columns:
- תלוש/מסמך
- סוג הפרה
- תקופה
- סכום (₪)

Add rows using actual data per document and violation.

🚨 TOTAL LINE:
Show total number of violations, total documents, and total amount claimed.

🔍 Breakdown by violation type:
- Total number of documents where each violation occurred
- Total amount per violation type

📘 Legal references:
List names of all laws used in calculations.

---

⚠️ FORMATTING & ACCURACY RULES:

- Always respond in Hebrew
- NEVER use sample values or formulas from laws directly
- Use only data provided in the payslip, employment contracts, and attendance records
- Replace ALL placeholders with actual values
- Do not hallucinate sick days, overtime hours, or absences
- If no violations found for a payslip, write:
  "לא נמצאו הפרות בתלוש מספר [X]"
- If no violations in any payslip, write:
  "לא נמצאו הפרות נגד חוקי העבודה שסופקו."
"""

    def _get_violation_count_table_instructions(self) -> str:
        return """
INSTRUCTIONS:
Analyze the provided documents and identify potential labor law violations, based *exclusively on the retrieved LABOR LAWS and JUDGEMENTS*.
Return ONLY a summary table in Hebrew, with NO additional analysis, commentary, or template text. Do not include any explanations, recommendations, or placeholders.

Provide a comprehensive summary table with actual data:
- Use the heading: === טבלת סיכום מקיפה ===
- Create columns for: תלוש/מסמך | סוג הפרה | תקופה | סכום (₪) | מספר הפרות
- Add rows with actual document names, violation types, periods, amounts, and the count of each violation type per document
- For each violation type, calculate and display the total amount for that violation type across all documents, and the total count of documents/slips affected
- For each violation type, explicitly state how many documents/slips had that violation (e.g., "6 מתוך 10 תלושים נמצאה הפרה של אי תשלום פנסיה")
- End with a total line showing total violations, months, and total amount
- Include a breakdown by violation type, showing the total count and total amount for each type of violation
- All amount calculations must be thorough and accurate: sum the amounts for each violation type, each document, and overall totals. Do not estimate or round; use the actual amounts found in the analysis.
- Include legal references from retrieved laws

Formatting requirements:
- Always respond in Hebrew
- Use actual violation types, counts, and amounts from the analysis
- Do not output any template text or placeholders
- Do not include any additional commentary or explanations
- If no violations are found, respond with: "לא נמצאו הפרות נגד חוקי העבודה שסופקו."
"""

    async def _extract_text2(self, content: bytes, filename: str, doc_type: str = None, compress: bool = False) -> Dict:
        """Extract structured data from various document formats using AgenticDoc for PDFs and images"""
        import re
        print("Extracting structured data from file:", filename.lower(), "with doc_type:", doc_type)
        try:
            # Load dynamic parameters for extraction schema
            dynamic_params_path = os.path.join(os.path.dirname(__file__), 'data', 'dynamic_parameters.json')
            with open(dynamic_params_path, 'r', encoding='utf-8') as f:
                dynamic_params = json.load(f)

            # Create extraction schema based on document type
            extraction_model = None
            if doc_type and doc_type.lower() in dynamic_params:
                doc_params = dynamic_params[doc_type.lower()]

                # Try creating a Pydantic model instead of schema
                from pydantic import BaseModel, Field, create_model

                # Create fields dict for dynamic model
                fields = {}
                for param in doc_params:
                    # Use parameter type to determine field type
                    param_type = param.get("type", "number")
                    if param_type == "number":
                        fields[param["param"]] = (Optional[Union[int, float]], Field(default=None, description=param["description"]))
                    else:  # string type
                        fields[param["param"]] = (Optional[str], Field(default=None, description=param["description"]))

                # Create dynamic model
                extraction_model = create_model(f'{doc_type.title()}Model', **fields)
            else:
                print(f"No parameters found for doc_type: {doc_type} in dynamic_params keys: {list(dynamic_params.keys()) if dynamic_params else 'None'}")

            # DOCX extraction block
            if filename.lower().endswith('.docx'):
                from docx import Document
                try:
                    docx_file = BytesIO(content)
                    doc = Document(docx_file)
                    # Extract paragraphs and tables separately
                    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                    tables = []
                    for table in doc.tables:
                        for row in table.rows:
                            row_data = [cell.text.strip() for cell in row.cells]
                            if any(row_data):
                                tables.append(row_data)
                    # Combine all text for extraction
                    full_text = '\n'.join(paragraphs)
                    for row in tables:
                        full_text += '\n' + '\t'.join(row)
                    print(f"Extracted text from DOCX: {full_text[:200]}...")

                    if extraction_model:
                        try:
                            print("Calling Gemini AI for DOCX extraction...")
                            # Build prompt for Gemini to extract fields using Hebrew labels
                            doc_params = dynamic_params[doc_type.lower()]
                            fields_description = "\n".join([f"{param['label_he']} -> {param['param']}: {param['description']} (type: {param.get('type', 'number')})" for param in doc_params])
                            return_format = "{" + ", ".join([f'"{param["param"]}": "value"' for param in doc_params]) + "}"
                            prompt = f"""
Extract the following fields from the provided document text. Use the parameter names (after ->) as JSON keys.

{fields_description}

Document text:
{full_text}

Return only a valid JSON object with the parameter names as keys and extracted values in their appropriate types (numbers as numbers, strings as strings, or null if not found).

For number fields, return numeric values (integers or decimals).
For string fields, return text values.
Use null for missing or unextractable values.

Return format: {return_format}
"""
                            result = await self.hebrew_agent.run(prompt, model_settings=ModelSettings(temperature=0.0))
                            response_text = result.output if hasattr(result, 'output') else str(result)
                            
                            # Try to parse JSON response
                            try:
                                print(f"Raw Gemini response: {response_text}")
                                response_text = response_text.strip()
                                # Handle markdown code blocks
                                if response_text.startswith('```json') and response_text.endswith('```'):
                                    response_text = response_text[7:-3].strip()
                                elif response_text.startswith('```') and response_text.endswith('```'):
                                    response_text = response_text[3:-3].strip()
                                
                                extracted_data = json.loads(response_text)
                            except json.JSONDecodeError:
                                print(f"Failed to parse Gemini response as JSON: {response_text}")
                                extracted_data = {}
                        except Exception as e:
                            print(f"Gemini AI extraction failed: {e}")
                            extracted_data = {}
                    print(f"Extracted structured data from DOCX: {extracted_data}")
                    return {'structured_data': extracted_data}
                except Exception as e:
                    print(f"DOCX extraction failed: {e}")
                    return {'structured_data': {}}

            if filename.lower().endswith(('.pdf', '.png', '.jpg', '.jpeg')):
                # For PDF/images: Use extraction model with dynamic params to get structured data
                print("Extracting structured data from image or PDF using extraction model...")

                if extraction_model:
                    try:
                        print(f"Calling agentic_doc parse with extraction_model...")
                        print(f"Model: {extraction_model}")
                        result = parse(content, extraction_model=extraction_model)
                        print(f"Parse result: {len(result)} documents")

                        if result and len(result) > 0:
                            print(f"First result has extraction: {hasattr(result[0], 'extraction')}")
                            if hasattr(result[0], 'extraction'):
                                print(f"Extraction data type: {type(result[0].extraction)}")
                                print(f"Extraction data: {result[0].extraction}")

                        # Extract structured data using the model
                        extracted_data = result[0].extraction if (result and hasattr(result[0], 'extraction') and result[0].extraction) else {}
                        print(f"Final extracted data: {extracted_data}")

                        # Convert Pydantic model to dict if needed
                        if hasattr(extracted_data, 'model_dump'):
                            extracted_data = extracted_data.model_dump()
                        elif hasattr(extracted_data, 'dict'):
                            extracted_data = extracted_data.dict()


                        return {
                            'structured_data': extracted_data  # Only return structured data
                        }

                    except Exception as e:
                        print(f"Extraction model failed: {e}")
                        # Return empty structured data when extraction fails
                        return {
                            'structured_data': {}
                        }
                else:
                    # No extraction model available
                    print(f"No extraction model available for {doc_type}")
                    return {
                        'structured_data': {}
                    }

            elif filename.lower().endswith('.xlsx') and doc_type == 'attendance':
                # Special handling for attendance Excel files
                print("Processing attendance Excel file with date, check-in, and departure time columns...")
                excel_file = BytesIO(content)
                df = pd.read_excel(excel_file, sheet_name=None)

                if df:
                    first_sheet = next(iter(df.values()))
                    print(f"Processing attendance Excel sheet with {len(first_sheet)} rows")
                    print(f"Columns found: {list(first_sheet.columns)}")

                    # Calculate attendance parameters from date, check-in, and departure times
                    attendance_data = self._extract_attendance_from_excel(first_sheet)

                    return {
                        'structured_data': attendance_data
                    }
                else:
                    return {
                        'structured_data': {}
                    }
            else:
                # Unknown file type: Try extraction model first, fallback to OCR
                print(f"Unknown file type, trying extraction model...")

                if extraction_model:
                    try:
                        result = parse(content, extraction_model=extraction_model)

                        if result and hasattr(result[0], 'extraction') and result[0].extraction:
                            extracted_data = result[0].extraction

                            # Convert Pydantic model to dict if needed
                            if hasattr(extracted_data, 'model_dump'):
                                extracted_data = extracted_data.model_dump()
                            elif hasattr(extracted_data, 'dict'):
                                extracted_data = extracted_data.dict()


                            return {
                                'structured_data': extracted_data
                            }
                    except Exception as e:
                        print(f"Extraction model failed for unknown type: {e}")

                # Fallback to empty data
                return {
                    'structured_data': {}
                }

        except Exception as e:
            print(f"Extraction error: {e}")
            # Fallback - return empty data
            return {
                'structured_data': {}
            }

    def _extract_attendance_from_excel(self, df) -> Dict:
        """Calculate attendance parameters from Excel with date, check-in, and departure time columns"""
        import pandas as pd
        from datetime import datetime, timedelta

        attendance_data = {}

        if df.empty:
            print("Empty DataFrame provided for attendance calculation")
            return attendance_data

        print(f"Processing attendance DataFrame with {len(df)} rows")
        print(f"Columns: {list(df.columns)}")

        # Find the relevant columns (case-insensitive search for English and Hebrew)
        date_col = None
        checkin_col = None
        departure_col = None

        for col in df.columns:
            col_str = str(col).strip()
            col_lower = col_str.lower()
            # English column names
            if 'date' in col_lower:
                date_col = col
            elif 'check' in col_lower and 'in' in col_lower:
                checkin_col = col
            elif 'departure' in col_lower or 'depart' in col_lower or 'out' in col_lower:
                departure_col = col
            # Hebrew column names - direct string matching
            elif 'תאריך' in col_str:
                date_col = col
            elif 'כניסה' in col_str or 'הגעה' in col_str:
                checkin_col = col
            elif 'יציאה' in col_str or 'עזיבה' in col_str:
                departure_col = col

        print(f"Identified columns - Date: {date_col}, Check-in: {checkin_col}, Departure: {departure_col}")

        if not all([date_col, checkin_col, departure_col]):
            print("Missing required columns for attendance calculation")
            return attendance_data

        # Clean and process the data
        df_clean = df.dropna(subset=[date_col, checkin_col, departure_col]).copy()
        print(f"After cleaning: {len(df_clean)} valid rows")

        if df_clean.empty:
            print("No valid rows after cleaning")
            return attendance_data

        # Calculate days_worked (unique dates)
        try:
            # Convert date column to datetime
            df_clean[date_col] = pd.to_datetime(df_clean[date_col], errors='coerce')
            valid_dates = df_clean[date_col].dropna()
            if not valid_dates.empty:
                days_worked = len(valid_dates.unique())
                attendance_data['days_worked'] = str(days_worked)
                print(f"Calculated days_worked: {days_worked}")
        except Exception as e:
            print(f"Error calculating days_worked: {e}")

        # Calculate month from dates
        try:
            if not valid_dates.empty:
                # Get the most common month from the dates
                months = valid_dates.dt.strftime('%Y-%m')
                if not months.empty:
                    most_common_month = months.mode().iloc[0]
                    attendance_data['month'] = most_common_month
                    print(f"Calculated month: {most_common_month}")
        except Exception as e:
            print(f"Error calculating month: {e}")

        # Calculate total_hours from check-in and departure times
        try:
            total_hours = 0.0
            valid_rows_count = 0

            for idx, row in df_clean.iterrows():
                try:
                    checkin_time = row[checkin_col]
                    departure_time = row[departure_col]

                    # Handle different time formats
                    if pd.isna(checkin_time) or pd.isna(departure_time):
                        continue

                    # Convert to datetime if they're not already
                    if not isinstance(checkin_time, datetime):
                        if isinstance(checkin_time, str):
                            # Try different time formats
                            for fmt in ['%H:%M:%S', '%H:%M', '%I:%M %p', '%I:%M:%S %p']:
                                try:
                                    checkin_time = datetime.strptime(checkin_time, fmt).time()
                                    break
                                except:
                                    continue
                        elif hasattr(checkin_time, 'hour'):  # Already a time object
                            pass
                        else:
                            continue

                    if not isinstance(departure_time, datetime):
                        if isinstance(departure_time, str):
                            # Try different time formats
                            for fmt in ['%H:%M:%S', '%H:%M', '%I:%M %p', '%I:%M:%S %p']:
                                try:
                                    departure_time = datetime.strptime(departure_time, fmt).time()
                                    break
                                except:
                                    continue
                        elif hasattr(departure_time, 'hour'):  # Already a time object
                            pass
                        else:
                            continue

                    # Calculate hours worked for this day
                    if hasattr(checkin_time, 'hour') and hasattr(departure_time, 'hour'):
                        # Create datetime objects for today to calculate difference
                        today = datetime.now().date()
                        checkin_dt = datetime.combine(today, checkin_time)
                        departure_dt = datetime.combine(today, departure_time)

                        # Handle overnight shifts (departure next day)
                        if departure_dt < checkin_dt:
                            departure_dt += timedelta(days=1)

                        hours_worked = (departure_dt - checkin_dt).total_seconds() / 3600
                        total_hours += hours_worked
                        valid_rows_count += 1
                        print(f"Row {idx}: Check-in {checkin_time}, Departure {departure_time}, Hours: {hours_worked:.2f}")

                except Exception as e:
                    print(f"Error processing row {idx}: {e}")
                    continue

            if valid_rows_count > 0:
                attendance_data['total_hours'] = str(round(total_hours, 2))
                print(f"Calculated total_hours: {total_hours:.2f}")

        except Exception as e:
            print(f"Error calculating total_hours: {e}")

        print(f"Final attendance data: {attendance_data}")
        return attendance_data

    def _compress_image(self, image_bytes: bytes, max_size_mb: int = 4) -> bytes:
        """Compress image to reduce size for API limits"""
        # Convert size to bytes
        max_size_bytes = max_size_mb * 1024 * 1024
        
        # If image is already smaller than max size, return original
        if len(image_bytes) <= max_size_bytes:
            return image_bytes
        
        # Open image using PIL
        img = Image.open(BytesIO(image_bytes))
        
        # Get original dimensions
        width, height = img.size
        
        # Scale down large images while maintaining aspect ratio
        max_dimension = 3000  # Maximum dimension for width or height
        if width > max_dimension or height > max_dimension:
            scale = min(max_dimension / width, max_dimension / height)
            new_width = int(width * scale)
            new_height = int(height * scale)
            img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
        
        # Preserve original format if possible
        output_format = img.format or 'JPEG'
        
        # Initial quality and compression settings
        quality = 95
        output = BytesIO()
        
        # Progressive quality reduction with format-specific handling
        while quality > 20:  # Increased minimum quality threshold
            output.seek(0)
            output.truncate(0)
            
            if output_format == 'JPEG':
                img.save(output, format=output_format, quality=quality, optimize=True)
            elif output_format == 'PNG':
                img.save(output, format=output_format, optimize=True)
            else:
                # Default to JPEG for unsupported formats
                img.save(output, format='JPEG', quality=70, optimize=True)
            
            if output.tell() <= max_size_bytes:
                break
            
            # More gradual quality reduction
            quality -= 5
        
        # If still too large, convert to JPEG as last resort
        if output.tell() > max_size_bytes and output_format != 'JPEG':
            output.seek(0)
            output.truncate(0)
            img.save(output, format='JPEG', quality=70, optimize=True)
        
        return output.getvalue()

    async def summarise(self, ai_content_text: str) -> str:
        """
        Summarizes the given text using PydanticAI.
        """
        try:
            prompt = f"Please summarize the following text concisely in Hebrew:\n\n{ai_content_text}"
            
            # Use the same agent for summarization
            try:
                result = await self.agent.run(prompt)
                return result.output if hasattr(result, 'output') else str(result)
            except Exception as pydantic_error:
                # If PydanticAI fails, check if it's an event loop issue
                if "Event loop is closed" in str(pydantic_error) or "event loop" in str(pydantic_error).lower():
                    # Try to recreate the agent with a fresh context
                    print("Attempting to recreate PydanticAI agent for summarization due to event loop issue...")
                    
                    # Create a new agent instance
                    temp_agent = Agent(
                        model=self.model,
                        output_type=str,
                        system_prompt="You are a helpful assistant that summarizes text concisely."
                    )
                    
                    # Try the request again with the fresh agent
                    result = await temp_agent.run(prompt)
                    return result.output if hasattr(result, 'output') else str(result)
                else:
                    # Re-raise the original error if it's not event loop related
                    raise pydantic_error
            
        except Exception as e:
            error_detail = f"Error generating summary with PydanticAI: {get_error_detail(e)}"
            raise HTTPException(
                status_code=500,
                detail=error_detail
            )
    
    # Sync wrapper methods for backward compatibility
    def create_report_sync(self, payslip_text: str = None, contract_text: str = None, 
                          attendance_text: str = None, type: str = "report", 
                          context: str = None) -> Dict:
        """Synchronous wrapper for create_report method with improved event loop handling"""
        # Handle backward compatibility parameter mapping
        try:
            # Try to get existing event loop
            loop = asyncio.get_event_loop()
            if loop.is_running():
                # If loop is already running, create a new thread
                import concurrent.futures
                import threading
                
                def run_in_thread():
                    # Create new event loop for this thread
                    new_loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(new_loop)
                    try:
                        return new_loop.run_until_complete(self.create_report(
                            payslip_text=payslip_text,
                            contract_text=contract_text, 
                            attendance_text=attendance_text,
                            analysis_type=type,  # Map 'type' to 'analysis_type'
                            context=context
                        ))
                    finally:
                        new_loop.close()
                
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(run_in_thread)
                    result = future.result()
            else:
                # No running loop, use asyncio.run
                result = asyncio.run(self.create_report(
                    payslip_text=payslip_text,
                    contract_text=contract_text, 
                    attendance_text=attendance_text,
                    analysis_type=type,  # Map 'type' to 'analysis_type'
                    context=context
                ))
        except RuntimeError as e:
            if "no current event loop" in str(e).lower() or "event loop is closed" in str(e).lower():
                # Create new event loop
                result = asyncio.run(self.create_report(
                    payslip_text=payslip_text,
                    contract_text=contract_text, 
                    attendance_text=attendance_text,
                    analysis_type=type,  # Map 'type' to 'analysis_type'
                    context=context
                ))
            else:
                raise e
        
        # Convert response to dict for backward compatibility
        return {
            "legal_analysis": result.legal_analysis,
            "status": result.status,
            # "relevant_laws": result.relevant_laws,
            # "relevant_judgements": result.relevant_judgements
        }
    
    def summarise_sync(self, ai_content_text: str) -> str:
        """Synchronous wrapper for summarise method with improved event loop handling"""
        try:
            # Try to get existing event loop
            loop = asyncio.get_event_loop()
            if loop.is_running():
                # If loop is already running, create a new thread
                import concurrent.futures
                
                def run_in_thread():
                    # Create new event loop for this thread
                    new_loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(new_loop)
                    try:
                        return new_loop.run_until_complete(self.summarise(ai_content_text))
                    finally:
                        new_loop.close()
                
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(run_in_thread)
                    return future.result()
            else:
                # No running loop, use asyncio.run
                return asyncio.run(self.summarise(ai_content_text))
        except RuntimeError as e:
            if "no current event loop" in str(e).lower() or "event loop is closed" in str(e).lower():
                # Create new event loop
                return asyncio.run(self.summarise(ai_content_text))
            else:
                raise e

    def qna_sync(self, report: str, questions: str) -> str:
        """Synchronous wrapper for the async qna method."""
        import asyncio
        try:
            return asyncio.run(self.qna(report, questions))
        except RuntimeError as e:
            # If there's already a running event loop (e.g. in Streamlit), use alternative
            if self._is_streamlit():
                try:
                    import nest_asyncio
                    nest_asyncio.apply()
                    loop = asyncio.get_event_loop()
                    return loop.run_until_complete(self.qna(report, questions))
                except Exception as inner_e:
                    raise RuntimeError(f"Failed to run qna async in Streamlit: {inner_e}") from e
            else:
                # In uvicorn/FastAPI, this shouldn't happen since we should use the async method directly
                raise RuntimeError(f"Cannot run async method in sync context. Use the async qna method instead: {e}") from e
            