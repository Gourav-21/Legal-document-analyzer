from fastapi import UploadFile, HTTPException
from PIL import Image
import os
from dotenv import load_dotenv
import pdfplumber
from io import BytesIO
from typing import List, Dict, Union, Optional
from docx import Document
from google.cloud import vision
import io
from google.cloud.vision import ImageAnnotatorClient
import pandas as pd
from letter_format import LetterFormatStorage
# from rag_storage import RAGLegalStorage
from rag_storage_local import RAGLegalStorage
from pydantic_ai import Agent
from pydantic_ai.models.openai import OpenAIModel
from pydantic_ai.models.gemini import GeminiModel
from pydantic import BaseModel
import asyncio
from agentic_doc.parse import parse
from judgement import JudgementStorage
from labour_law import LaborLawStorage
import nest_asyncio
# Load environment variables from .env 
load_dotenv()



class DocumentAnalysisRequest(BaseModel):
    """Request model for document analysis"""
    payslip_text: Optional[str] = None
    contract_text: Optional[str] = None
    attendance_text: Optional[str] = None
    analysis_type: str = "report"
    context: Optional[str] = None

class DocumentAnalysisResponse(BaseModel):
    """Response model for document analysis"""
    legal_analysis: str
    # relevant_laws: List[Dict]
    # relevant_judgements: List[Dict]
    status: str
    analysis_type: str

class DocumentProcessor:
    def __init__(self):
        # Initialize RAG storage
        self.rag_storage = RAGLegalStorage()
        
        # Initialize backward compatibility storages
        self.letter_format = LetterFormatStorage()
        # self.law_storage = LaborLawStorage()
        # self.judgement_storage = JudgementStorage()
        self.law_storage = self.rag_storage
        self.judgement_storage = self.rag_storage
        
        # Initialize PydanticAI model - use Gemini only
        gemini_api_key = os.getenv("GOOGLE_CLOUD_VISION_API_KEY")
        if gemini_api_key:
            self.model = GeminiModel('gemini-2.5-pro', api_key=gemini_api_key)
            self.model_type = "gemini"
        else:
            raise Exception("GEMINI_API_KEY must be set in environment variables")
        # Initialize PydanticAI Agent
        nest_asyncio.apply()
        self.agent = Agent(
            model=self.model,
            result_type=str,
            system_prompt="""You are an expert legal document analyzer specializing in Israeli labor law compliance.

You will be provided with:
1. Relevant labor laws retrieved from a legal database
2. Relevant legal judgements and precedents retrieved from a legal database  
3. Document content to analyze (payslips, contracts, attendance records)
4. Additional context if provided

Your analysis must be based STRICTLY on the provided laws and judgements. Do not use external legal knowledge.

Always respond in Hebrew and follow the specific formatting requirements for each analysis type."""        )
        
        nest_asyncio.apply()
        self.context_agent = Agent(
            model=self.model,
            result_type=str,
            system_prompt="""You are an expert legal document analyzer specializing in Israeli labor law compliance.
                    you have to collect all the context of laws and judgements from the RAG storage to give ai for correct anaslysis.
                    create search queries based on law summaries and the document content provided and use the get_laws tool to retrieve relevant laws and use the get_judgements tool to retrieve relevant judgements related to the case.
                    Always respond in Hebrew.
            """
        )
        
        @self.context_agent.system_prompt
        def dynamic_system_prompt():
            """Dynamic system prompt for context agent"""
            # Get all law summaries from RAG storage
            law_summaries = self.rag_storage.get_all_law_summaries()
            
            if law_summaries:
                print(f"Found {len(law_summaries)} law summaries in the database.")
                summaries_text = "\n".join([f"- {summary}" for summary in law_summaries])
                return f"""
הינך מנתח מסמכים משפטיים מומחה המתמחה בהתאמה לחוקי העבודה הישראליים.

הנה רשימה של סיכומי חוקים הזמינים במאגר הנתונים שניתן לחפש בהם:

{summaries_text}

על בסיס תוכן המסמכים שסופק, זהה את המושגים המשפטיים הרלוונטיים ביותר וצור שאילתות חיפוש מתאימות למציאת החוקים והפסיקות הרלוונטיים.
"""
            else:
                return """
הינך מנתח מסמכים משפטיים מומחה המתמחה בהתאמה לחוקי העבודה הישראליים.

אין כרגע סיכומי חוקים זמינים במאגר הנתונים. נתח את תוכן המסמכים וצור שאילתות חיפוש על בסיס המושגים המשפטיים הכלליים הקיימים.
"""
        @self.context_agent.tool_plain
        async def get_laws(query: str, max_laws: int = 5) -> List[Dict]:
            """Retrieve relevant labor laws based on query"""
            print(f"Searching for laws with query: {query}")
            return self.rag_storage.search_laws(query, n_results=max_laws)
        
        @self.context_agent.tool_plain
        async def get_judgements(query: str, max_judgements: int = 3) -> List[Dict]:
            """Retrieve relevant legal judgements based on query"""
            print(f"Searching for judgements with query: {query}")
            return self.rag_storage.search_judgements(query, n_results=max_judgements)
        
        

        # Configure Vision API
        vision_api_key = os.getenv("GOOGLE_CLOUD_VISION_API_KEY")
        if not vision_api_key:
            raise Exception("Google Cloud Vision API key not found. Please set GOOGLE_CLOUD_VISION_API_KEY in your .env file")
        
        self.vision_client = ImageAnnotatorClient(client_options={"api_key": vision_api_key})
        self.image_context = {"language_hints": ["he"]} 

    def process_document(self, files: Union[UploadFile, List[UploadFile]], doc_types: Union[str, List[str]], compress: bool = False) -> Dict[str, str]:
        """Process uploaded documents and extract text"""
        if not files:
            raise HTTPException(
                status_code=400,
                detail="At least one document must be provided"
            )
          # Handle single file case
        if isinstance(files, UploadFile):
            files = [files]
            doc_types = [doc_types]
        elif not isinstance(doc_types, list):
            doc_types = [doc_types] * len(files)
        
        # Initialize text storage
        payslip_text = None
        contract_text = None
        attendance_text = None
        
        # Initialize counters
        payslip_counter = 0
        contract_counter = 0
        attendance_counter = 0
        # Process each file based on its type
        for file, doc_type in zip(files, doc_types):
            print(f"Processing file: {file.filename} as type: {doc_type}")
            extracted_text = self._extract_text2(file.file.read(), file.filename, compress=compress)
            if doc_type.lower() == "payslip":
                payslip_counter += 1
                payslip_text = f"Payslip {payslip_counter}:\n{extracted_text}" if payslip_text is None else f"{payslip_text}\n\nPayslip {payslip_counter}:\n{extracted_text}"
            elif doc_type.lower() == "contract":
                contract_counter += 1
                contract_text = f"Contract {contract_counter}:\n{extracted_text}" if contract_text is None else f"{contract_text}\n\nContract {contract_counter}:\n{extracted_text}"
            elif doc_type.lower() == "attendance":
                attendance_counter += 1
                attendance_text = f"Attendance {attendance_counter}:\n{extracted_text}" if attendance_text is None else f"{attendance_text}\n\nAttendance {attendance_counter}:\n{extracted_text}"
        
        return {
            "payslip_text": payslip_text,
            "contract_text": contract_text,
            "attendance_text": attendance_text
        }

    async def create_report(self, payslip_text: str = None, contract_text: str = None, 
                          attendance_text: str = None, analysis_type: str = "report", 
                          context: str = None) -> DocumentAnalysisResponse:
        """Create legal analysis report using RAG and PydanticAI"""
        
        # Prepare documents for analysis
        documents = {}
        if payslip_text:
            documents['payslip'] = payslip_text
        if contract_text:
            documents['contract'] = contract_text
        if attendance_text:
            documents['attendance report'] = attendance_text        # Pass the string to the context agent instead of the dictionary
        docs= "\n\n".join([f"{doc_type.upper()} CONTENT:\n{content}" for doc_type, content in documents.items()])
        context_result = await self.context_agent.run(docs)

        # Extract the string data from the result object
        ai_context = str(context_result.data) if hasattr(context_result, 'data') else str(context_result)

        print(f"Retrieving laws for context: {ai_context}")

        # Retrieve relevant laws and judgements using RAG
        # relevant_laws, relevant_judgements = self.rag_storage.get_relevant_context(
        #     query, max_laws=5, max_judgements=3
        # )
        
        # # Format retrieved content for prompt
        # formatted_laws = self.rag_storage.format_laws_for_prompt(relevant_laws)
        # formatted_judgements = self.rag_storage.format_judgements_for_prompt(relevant_judgements)
        
        # Build the analysis prompt based on type
        prompt = await self._build_analysis_prompt(
            analysis_type, documents, ai_context, context
        )
        
        try:
            # Generate analysis using PydanticAI
            try:
                result = await self.agent.run(prompt)
                analysis = result.data
            except Exception as pydantic_error:
                # If PydanticAI fails, check if it's an event loop issue
                if "Event loop is closed" in str(pydantic_error) or "event loop" in str(pydantic_error).lower():
                    # Try to recreate the agent with a fresh context
                    print(pydantic_error)
                    print("Attempting to recreate PydanticAI agent due to event loop issue...")
                      # Create a new agent instance
                    temp_agent = Agent(
                        model=self.model,
                        result_type=str,
                        system_prompt="""You are an expert legal document analyzer specializing in Israeli labor law compliance.

You will be provided with:
1. Relevant labor laws retrieved from a legal database
2. Relevant legal judgements and precedents retrieved from a legal database  
3. Document content to analyze (payslips, contracts, attendance records)
4. Additional context if provided

Your analysis must be based STRICTLY on the provided laws and judgements. Do not use external legal knowledge.

Always respond in Hebrew and follow the specific formatting requirements for each analysis type."""
                    )
                      # Try the request again with the fresh agent
                    result = await temp_agent.run(prompt)
                    analysis = result.data
                else:
                    # Re-raise the original error if it's not event loop related
                    raise pydantic_error
            
            return DocumentAnalysisResponse(
                legal_analysis=analysis,
                # relevant_laws=relevant_laws,
                # relevant_judgements=relevant_judgements,
                status="success",
                analysis_type=analysis_type
            )
            
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error generating legal analysis: {str(e)}"
            )

    async def _build_analysis_prompt(self, analysis_type: str, documents: Dict, 
                                   ai_context: str, 
                                   context: str = None) -> str:
        """Build analysis prompt based on type"""
        
        base_prompt = f"""
RELEVANT LABOR LAWS and JUDGEMENTS (Retrieved from legal database):
{ai_context}


DOCUMENTS PROVIDED FOR ANALYSIS:
{', '.join(documents.keys())}

ADDITIONAL CONTEXT:
{context if context else 'No additional context provided.'}
"""
        # Add document contents to prompt
        for doc_type, content in documents.items():
            base_prompt += f"\n{doc_type.upper()} CONTENT:\n{content}\n"
        
        # Add specific instructions based on analysis type
        if analysis_type == 'report':
            base_prompt += self._get_report_instructions()
        elif analysis_type == 'profitability':
            base_prompt += self._get_profitability_instructions()
        elif analysis_type == 'professional':
            base_prompt += self._get_professional_instructions()
        elif analysis_type == 'warning_letter':
            base_prompt += await self._get_warning_letter_instructions()
        elif analysis_type == 'easy':
            base_prompt += self._get_easy_instructions()
        elif analysis_type == 'table':
            base_prompt += self._get_table_instructions()
        elif analysis_type == 'claim':
            base_prompt += self._get_claim_instructions()
        
        return base_prompt

    def _get_report_instructions(self) -> str:
        return """
INSTRUCTIONS:
1. If no labor laws are provided, respond with: "אין חוקים לעבודה זמינים לניתוח התאמה." in Hebrew.
2. If labor laws exist, analyze the documents ONLY against the provided laws.
3. ONLY refer to the judgements and their results provided above for legal analysis - do not use external cases or knowledge.
4. If no judgements are provided, respond with: "לא קיימות החלטות משפטיות זמינות לניתוח." in Hebrew.

For each payslip provided, analyze and identify violations. For each violation found in each payslip, format the response EXACTLY as shown below, with each section on a new line and proper spacing:

Violation Format Template:

[VIOLATION TITLE]

[SPECIFIC VIOLATION DETAILS]

[LAW REFERENCE AND YEAR FROM PROVIDED LAWS]

[SIMILAR CASES OR PRECEDENTS FROM PROVIDED JUDGEMENTS](Refer *only* to the retrieved judgements. If a relevant judgement is found, describe the case and its result. If no relevant judgement is found, state "לא נמצאו תקדימים רלוונטיים בפסקי הדין שסופקו.")

[LEGAL IMPLICATIONS BASED ON PROVIDED INFORMATION]

[RECOMMENDED ACTIONS]

---

IMPORTANT:
- Always Respond in Hebrew
- Format each violation with proper spacing and line breaks as shown above
- Analyze each payslip separately and clearly indicate which payslip the violations belong to
- Separate multiple violations with '---'
- If no violations are found against the provided laws in a payslip, respond with: "לא נמצאו הפרות בתלוש מספר [X]" in hebrew
- If no violations are found in any payslip, respond with: "לא נמצאו הפרות נגד חוקי העבודה שסופקו." in hebrew
- Do not include any additional commentary or explanations outside of the violation format
"""

    def _get_profitability_instructions(self) -> str:
        return """
INSTRUCTIONS:
1. Analyze the provided documents and identify potential labor law violations, based *exclusively on the retrieved LABOR LAWS and JUDGEMENTS*.
2. For each violation, refer *only* to the retrieved judgements to identify similar legal cases and their outcomes (both successful and unsuccessful).
3. If similar cases (from the retrieved judgements) were unsuccessful:
   - Explain why the cases were unsuccessful based on the retrieved judgement information.
   - Provide a clear recommendation against pursuing legal action based on this.
   - List potential risks and costs.

4. If similar cases (from the retrieved judgements) were successful, calculate using information from those judgements if available:
   - Average compensation amount from successful retrieved cases.
   - Estimated legal fees (30% of potential compensation).
   - Tax implications (25% of net compensation).
   - Time and effort cost estimation.

Provide the analysis in the following format:

ניתוח כדאיות כלכלית:

הפרות שזוהו (על בסיס החוקים שסופקו):
[List identified violations]

תקדימים משפטיים (מתוך פסקי הדין שסופקו):
[Similar cases from retrieved judgements with outcomes - both successful and unsuccessful]

במקרה של תקדימים שליליים (מתוך פסקי הדין שסופקו):
- סיבות לדחיית התביעות: [REASONS BASED ON RETRIEVED JUDGEMENTS]
- סיכונים אפשריים: [RISKS]
- המלצה: לא מומלץ להגיש תביעה בשל [EXPLANATION BASED ON RETRIEVED JUDGEMENTS]

במקרה של תקדימים חיוביים (מתוך פסקי הדין שסופקו):
ניתוח כספי:
- סכום פיצוי ממוצע (מפסיקות דין שסופקו): [AMOUNT] ₪
- עלות משוערת של עורך דין (30%): [AMOUNT] ₪
- השלכות מס (25% מהסכום נטו): [AMOUNT] ₪
- סכום נטו משוער: [AMOUNT] ₪

המלצה סופית:
[Based on analysis of both successful and unsuccessful cases from the retrieved judgements, provide clear recommendation]
"""

    def _get_professional_instructions(self) -> str:
        return """
Analyze the provided documents for labor law violations based strictly on the retrieved Israeli labor laws and the content of the documents. For each violation, calculate the monetary differences using only those laws.
Provide your analysis in the following format, entirely in Hebrew:

ניתוח מקצועי של הפרות שכר:

הפרה: [כותרת ההפרה]
[תיאור מפורט של ההפרה, כולל תאריכים רלוונטיים, שעות עבודה, שכר שעתי וחישובים, בהתבסס אך ורק על החוקים הישראליים שנמצאו והמסמכים שסופקו.
דוגמה: העובד עבד X שעות נוספות בין [חודש שנה] ל-[חודש שנה]. לפי שכר שעתי בסיסי של [שכר] ₪ ושיעורי תשלום שעות נוספות ([שיעור1]% עבור X השעות הראשונות, [שיעור2]% לאחר מכן) כפי שמופיע בחוקי העבודה שנמצאו, העובד היה זכאי ל-[סכום] ₪ לחודש. בפועל קיבל רק [סכום שקיבל] ₪ למשך X חודשים ו-[סכום] ₪ בחודש [חודש].]
סה"כ חוב: [סכום ההפרש עבור הפרה זו] ₪

הפרה: [כותרת ההפרה]
[תיאור מפורט של ההפרה, כולל תאריכים וחישובים, בהתבסס אך ורק על החוקים שנמצאו והמסמכים. 
דוגמה: בחודש [חודש שנה] לא בוצעה הפקדה לפנסיה. המעסיק מחויב להפקיד [אחוז]% מהשכר בגובה [שכר] ₪ = [סכום] ₪ בהתאם לחוק/צו הרחבה שנמצא.]
סה"כ חוב פנסיה: [סכום חוב הפנסיה להפרה זו] ₪

---

סה"כ תביעה משפטית (לא כולל ריבית): [הסכום הכולל לתביעה מכלל ההפרות] ₪  
אסמכתאות משפטיות: [רשימת שמות החוק הרלוונטיים מתוך החוקים הישראליים שנמצאו. לדוגמה: חוק שעות עבודה ומנוחה, צו הרחבה לפנסיה חובה]
"""

    async def _get_warning_letter_instructions(self) -> str:
        # Get letter format from storage (we'll need to adapt this)
        # For now, using a default template
        format_content = """
[תאריך]

לכבוד
[שם המעסיק]
[כתובת המעסיק]

הנדון: התראה בגין הפרות חוקי עבודה

בהתבסס על בדיקת המסמכים שבידינו, נמצאו הפרות של חוקי עבודה כמפורט להלן:

[פרטי ההפרות]

הפרות אלו מהוות הפרה של:
[הפניות לחוקים]

בהתאם לכך, אנו דורשים כי תתקנו את ההפרות הנ"ל תוך [מועד] ימים מקבלת מכתב זה.

אי תיקון ההפרות עלול להוביל לנקיטת הליכים משפטיים.

בכבוד,
[חתימה]
"""
        
        return f"""
INSTRUCTIONS:
1. Analyze the provided documents for labor law violations *based exclusively on the retrieved LABOR LAWS and JUDGEMENTS*.
2. If violations are found, generate a formal warning letter using the provided template.
3. If no violations are found, respond with: "לא נמצאו הפרות המצדיקות מכתב התראה." in Hebrew.

Warning Letter Template:
{format_content}

Please generate the warning letter in Hebrew with the following guidelines:
- Replace [שם המעסיק] with the employer's name from the documents
- Replace [פרטי ההפרות] with specific details of each violation found (based on retrieved laws)
- Replace [הפניות לחוקים] with relevant labor law citations *from the retrieved LABOR LAWS*.
- Replace [מועד] with a reasonable timeframe for corrections (typically 14 days)
- Maintain a professional and formal tone throughout
- Include all violations found in the analysis (based on retrieved laws)
- Format the letter according to the provided template structure
"""

    def _get_easy_instructions(self) -> str:
        return """
🔒 מטרה: צור סיכום קצר וברור של ההפרות בתלושי השכר של העובד.
📌 כללים מחייבים:
	1. כתוב בעברית בלבד – אל תשתמש באנגלית בכלל.
	2. עבור כל חודש הצג את ההפרות בשורות נפרדות, כל שורה בפורמט הבא:
❌ [סוג ההפרה בקצרה] – [סכום בש"ח עם ₪, כולל פסיק לאלפים]
לדוגמה: ❌ לא שולם החזר נסיעות בפברואר 2025 – 250 ₪
	3. אם יש מספר רכיבי פנסיה (עובד/מעסיק/בריאות) בחודש מסוים – חבר אותם לסכום אחד של פנסיה באותו החודש.
	4. כל הסכומים יוצגו עם פסיקים לאלפים ועם ₪ בסוף.
	5. חישוב הסכום הכולל יופיע בשורה נפרדת:
💰 סה"כ: [סכום כולל] ₪
	6. הוסף המלצה בסוף:
📝 מה לעשות עכשיו:
פנה/י למעסיק עם דרישה לתשלום הסכומים.
אם אין מענה – מומלץ לפנות לייעוץ משפטי.
📍 הנחיות נוספות:
	• אין לכתוב מספרים בלי הקשר, כל שורה חייבת להיות מלווה בחודש.
	• מיזוג שורות: אם באותו חודש יש כמה רכיבים של פנסיה – מיזג אותם לשורה אחת.
	• הסר שורות ללא סכום ברור.
	• ניסוח פשוט, ללא מינוחים משפטיים, הבהרות או הסברים טכניים.
	• אין לציין "רכיב עובד", "רכיב מעסיק", "לא הופקד" – במקום זאת כתוב: "לא שולמה פנסיה".
🎯 פלט רצוי:
	• שורות מסודרות לפי חודשים
	• אין כפילויות
	• סכומים מדויקים בלבד
	• ניסוח ברור ומובן
	• עברית בלבד

🧪 Example of desired output:
📢 סיכום ההפרות:
❌ לא שולם עבור שעות נוספות בנובמבר 2024 – 495 ₪
❌ לא שולמה פנסיה בנובמבר 2024 – 750 ₪
❌ לא שולמה פנסיה בדצמבר 2024 – 1,221 ₪
❌ לא שולמה פנסיה בינואר 2025 – 831 ₪
❌ לא שולם החזר נסיעות בפברואר 2025 – 250 ₪
❌ לא שולמה פנסיה בפברואר 2025 – 858 ₪
❌ לא שולמה פנסיה במרץ 2025 – 866 ₪
💰 סה"כ: 5,271 ₪
📝 מה לעשות עכשיו:
פנה/י למעסיק עם דרישה לתשלום הסכומים.
אם אין מענה – מומלץ לפנות לייעוץ משפטי.
"""

    def _get_table_instructions(self) -> str:
        return """
אתה עוזר משפטי מיומן. עליך לנתח את רשימת ההפרות ולהפיק רשימת תביעות מסודרת לפי מסמך (לדוגמה: תלוש שכר מס' 1, מסמך שימוע, מכתב פיטורין וכו').

הנחיות:

1. סדר את התביעות לפי מסמך: כל קבוצה מתחילה בכותרת כמו "תלוש שכר מס' 4 – 05/2024:".

2. תחת כל כותרת, צור רשימה ממוספרת באותיות עבריות (א., ב., ג. וכו').

3. השתמש במבנה הקבוע הבא:
   א. סכום של [amount] ש"ח עבור [תיאור קצר של ההפרה].

4. השתמש בפורמט מספרים עם פסיקים לאלפים ושתי ספרות אחרי הנקודה (למשל: 1,618.75 ש"ח).

5. כתוב "ש"ח" אחרי הסכום — לא ₪.

6. אל תסכם את כלל ההפרות – הצג רק את הפירוט לפי מסמך.

7. אל תוסיף בולטים, טבלאות או הערות.

נתוני הקלט לדוגמה:
❌ אי תשלום שעות נוספות (תלוש 6, 11/2024) – 495 ש"ח
❌ העדר רכיב מעביד לפנסיה (תלוש 6, 11/2024) – 750 ש"ח
❌ העדר רכיב עובד לפנסיה (תלוש 7, 12/2024) – 396 ש"ח
❌ העדר נסיעות (תלוש 9, 02/2025) – 250 ש"ח

פלט נדרש לדוגמה:

תלוש שכר מס' 6 – 11/2024:
א. סכום של 495.00 ש"ח עבור אי תשלום שעות נוספות.
ב. סכום של 750.00 ש"ח עבור העדר רכיב מעביד לפנסיה.

תלוש שכר מס' 7 – 12/2024:
א. סכום של 396.00 ש"ח עבור העדר רכיב עובד לפנסיה.

תלוש שכר מס' 9 – 02/2025:
א. סכום של 250.00 ש"ח עבור העדר תשלום עבור נסיעות.

החזר את הפלט בפורמט זה בלבד, מופרד לפי כל מסמך.
"""

    def _get_claim_instructions(self) -> str:
        return """
משימה:
כתוב טיוטת כתב תביעה לבית הדין האזורי לעבודה, בהתאם למבנה המשפטי הנהוג בישראל.

נתונים:
השתמש במידע מתוך המסמכים שצורפו (כגון תלושי שכר, הסכמי עבודה, הודעות פיטורין, שיחות עם המעסיק) ובממצאים שנמצאו בניתוח קודם של ההפרות.

פורמט לכתיבה:
1. כותרת: "בית הדין האזורי לעבודה ב[שם עיר]"
2. פרטי התובע/ת:
   - שם מלא, ת"ז, כתובת, טלפון, מייל
3. פרטי הנתבע/ת (מעסיק):
   - שם החברה/מעסיק, ח.פ./ת"ז, כתובת, טלפון, מייל (אם קיים)
4. כותרת: "כתב תביעה"
5. פתיח משפטי:
   בית הדין הנכבד מתבקש לזמן את הנתבע/ת לדין ולחייבו/ה לשלם לתובע/ת את הסכומים המפורטים, מהנימוקים הבאים:
6. סעיף 1 – רקע עובדתי:
   תיאור תקופת העבודה, תפקידים, תאריך תחילה וסיום (אם רלוונטי), מהות יחסי העבודה, מקום העבודה.
7. סעיף 2 – עילות התביעה (לפי הפרות):
   לכל הפרה:
     - איזה חוק הופֵר (לדוג': חוק שכר מינימום, חוק שעות עבודה ומנוחה וכו')
     - פירוט העובדות והתקופה הרלוונטית
     - סכום הפיצוי או הנזק
     - אסמכתאות משפטיות אם רלוונטי
8. סעיף 3 – נזקים שנגרמו:
   סכומים כספיים (בפירוט), נזק לא ממוני אם קיים (עוגמת נפש).
9. סעיף 4 – סעדים מבוקשים:
   תשלום הפרשים, פיצויים, ריבית והצמדה, הוצאות משפט, שכר טרחת עו"ד, וכל סעד אחר.
10. סעיף 5 – סמכות שיפוט:
   ציין סמכות בית הדין לעבודה לפי חוק בית הדין לעבודה תשכ"ט–1969.
11. סעיף 6 – דיון מקדים/הליך מהיר (אם רלוונטי).
12. סעיף 7 – כתובות להמצאת כתבי בי-דין:
   כתובת התובע והנתבע.
13. סיום:
   חתימה, תאריך, רשימת נספחים תומכים (תלושי שכר, מכתבים וכו').

שפה:
- כתוב בעברית משפטית, רשמית ומסודרת.
- סדר את התביעה עם כותרות, סעיפים ממוספרים ורווחים ברורים.
- אל תמציא עובדות או חוקים. אם חסר מידע – השאר שדה ריק לצורך השלמה ע"י המשתמש.

אזהרה:
בסיום, כתוב משפט: "הטיוטה נכתבה אוטומטית לצורך עזרה ראשונית בלבד. מומלץ לפנות לעורך/ת דין מוסמך/ת לפני הגשה לבית הדין."
"""

    def _extract_text2(self, content: bytes, filename: str, compress: bool = False) -> str:
        """Extract text from various document formats using AgenticDoc for PDFs and images"""
        print("Extracting text from file:", filename.lower())
        
        try:
            if filename.lower().endswith(('.pdf', '.png', '.jpg', '.jpeg')):
                # Pass bytes directly to AgenticDoc
                print("Extracting text from image or PDF...")
                
                result = parse(content)
                # Return markdown or structured, as you prefer
                return result[0].markdown
            elif filename.lower().endswith('.docx'):
                doc_file = BytesIO(content)
                doc = Document(doc_file)
                return '\n'.join([p.text for p in doc.paragraphs])
            elif filename.lower().endswith('.xlsx'):
                excel_file = BytesIO(content)
                df = pd.read_excel(excel_file, sheet_name=None)
                text = ""
                for sheet, data in df.items():
                    text += f"Sheet: {sheet}\n"
                    text += data.to_string(index=False) + "\n\n"
                return text
            else:
                # Unknown file type, treat as plain text
                result = parse(content)
                # Return markdown or structured, as you prefer
                return result[0].markdown
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Extraction error: {str(e)}")

    def _extract_text_legacy(self, content: bytes, filename: str, compress: bool = False) -> str:
        """Legacy text extraction method (kept for fallback)"""
        if filename.lower().endswith('.pdf'):
            return self._extract_pdf_text(content)
        elif filename.lower().endswith('.docx'):
            return self._extract_docx_text(content)
        elif filename.lower().endswith('.xlsx'):
            return self._extract_excel_text(content)
        else:
            return self._extract_image_text(content, compress)

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF using pdfplumber and Vision API"""
        try:
            pdf_file = BytesIO(content)
            text = ''
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += page_text + "\n"
                    else:
                        # Convert PDF page to image and use Vision API
                        img = page.to_image(resolution=300).original
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='PNG')
                        img_content = img_byte_arr.getvalue()
                        
                        vision_image = vision.Image(content=img_content)
                        response = self.vision_client.text_detection(image=vision_image, image_context=self.image_context)
                        if response.text_annotations:
                            text_list = [text_annotation.description for text_annotation in response.text_annotations]
                            text += " ".join(text_list) + "\n"
            return text
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error processing PDF: {str(e)}")

    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX files"""
        doc_file = BytesIO(content)
        doc = Document(doc_file)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    def _extract_excel_text(self, content: bytes) -> str:
        """Extract text from Excel files"""
        try:
            excel_file = BytesIO(content)
            df = pd.read_excel(excel_file, sheet_name=None)  # Load all sheets
            
            text = ""
            for sheet_name, sheet_data in df.items():
                text += f"Sheet: {sheet_name}\n"
                text += sheet_data.to_string(index=False) + "\n\n"
            return text
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error processing Excel file: {str(e)}")

    def _extract_image_text(self, content: bytes, compress: bool = False) -> str:
        """Extract text from images using Vision API"""
        try:
            if compress:
                content = self._compress_image(content)
            
            vision_image = vision.Image(content=content)
            response = self.vision_client.document_text_detection(image=vision_image, image_context=self.image_context)
            
            if response.error.message:
                raise HTTPException(status_code=500, detail=f"Vision API Error: {response.error.message}")
            
            if response.text_annotations:
                return " ".join([text_annotation.description for text_annotation in response.text_annotations])
            else:
                return ""
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error processing image: {str(e)}")

    def _compress_image(self, image_bytes: bytes, max_size_mb: int = 4) -> bytes:
        """Compress image to reduce size for API limits"""
        # Convert size to bytes
        max_size_bytes = max_size_mb * 1024 * 1024
        
        # If image is already smaller than max size, return original
        if len(image_bytes) <= max_size_bytes:
            return image_bytes
        
        # Open image using PIL
        img = Image.open(BytesIO(image_bytes))
        
        # Get original dimensions
        width, height = img.size
        
        # Scale down large images while maintaining aspect ratio
        max_dimension = 3000  # Maximum dimension for width or height
        if width > max_dimension or height > max_dimension:
            scale = min(max_dimension / width, max_dimension / height)
            new_width = int(width * scale)
            new_height = int(height * scale)
            img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
        
        # Preserve original format if possible
        output_format = img.format or 'JPEG'
        
        # Initial quality and compression settings
        quality = 95
        output = BytesIO()
        
        # Progressive quality reduction with format-specific handling
        while quality > 20:  # Increased minimum quality threshold
            output.seek(0)
            output.truncate(0)
            
            if output_format == 'JPEG':
                img.save(output, format=output_format, quality=quality, optimize=True)
            elif output_format == 'PNG':
                img.save(output, format=output_format, optimize=True)
            else:
                # Default to JPEG for unsupported formats
                img.save(output, format='JPEG', quality=70, optimize=True)
            
            if output.tell() <= max_size_bytes:
                break
            
            # More gradual quality reduction
            quality -= 5
        
        # If still too large, convert to JPEG as last resort
        if output.tell() > max_size_bytes and output_format != 'JPEG':
            output.seek(0)
            output.truncate(0)
            img.save(output, format='JPEG', quality=70, optimize=True)
        
        return output.getvalue()

    async def summarise(self, ai_content_text: str) -> str:
        """
        Summarizes the given text using PydanticAI.
        """
        try:
            prompt = f"Please summarize the following text concisely in Hebrew:\n\n{ai_content_text}"
            
            # Use the same agent for summarization
            try:
                result = await self.agent.run(prompt)
                return result.data if hasattr(result, 'data') else str(result)
            except Exception as pydantic_error:
                # If PydanticAI fails, check if it's an event loop issue
                if "Event loop is closed" in str(pydantic_error) or "event loop" in str(pydantic_error).lower():
                    # Try to recreate the agent with a fresh context
                    print("Attempting to recreate PydanticAI agent for summarization due to event loop issue...")
                    
                    # Create a new agent instance
                    temp_agent = Agent(
                        model=self.model,
                        result_type=str,
                        system_prompt="You are a helpful assistant that summarizes text concisely."
                    )
                    
                    # Try the request again with the fresh agent
                    result = await temp_agent.run(prompt)
                    return result.data if hasattr(result, 'data') else str(result)
                else:
                    # Re-raise the original error if it's not event loop related
                    raise pydantic_error
            
        except Exception as e:
            error_detail = f"Error generating summary with PydanticAI: {str(e)}"
            raise HTTPException(
                status_code=500,
                detail=error_detail
            )
    
    # Sync wrapper methods for backward compatibility
    def create_report_sync(self, payslip_text: str = None, contract_text: str = None, 
                          attendance_text: str = None, type: str = "report", 
                          context: str = None) -> Dict:
        """Synchronous wrapper for create_report method with improved event loop handling"""
        # Handle backward compatibility parameter mapping
        try:
            # Try to get existing event loop
            loop = asyncio.get_event_loop()
            if loop.is_running():
                # If loop is already running, create a new thread
                import concurrent.futures
                import threading
                
                def run_in_thread():
                    # Create new event loop for this thread
                    new_loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(new_loop)
                    try:
                        return new_loop.run_until_complete(self.create_report(
                            payslip_text=payslip_text,
                            contract_text=contract_text, 
                            attendance_text=attendance_text,
                            analysis_type=type,  # Map 'type' to 'analysis_type'
                            context=context
                        ))
                    finally:
                        new_loop.close()
                
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(run_in_thread)
                    result = future.result()
            else:
                # No running loop, use asyncio.run
                result = asyncio.run(self.create_report(
                    payslip_text=payslip_text,
                    contract_text=contract_text, 
                    attendance_text=attendance_text,
                    analysis_type=type,  # Map 'type' to 'analysis_type'
                    context=context
                ))
        except RuntimeError as e:
            if "no current event loop" in str(e).lower() or "event loop is closed" in str(e).lower():
                # Create new event loop
                result = asyncio.run(self.create_report(
                    payslip_text=payslip_text,
                    contract_text=contract_text, 
                    attendance_text=attendance_text,
                    analysis_type=type,  # Map 'type' to 'analysis_type'
                    context=context
                ))
            else:
                raise e
        
        # Convert response to dict for backward compatibility
        return {
            "legal_analysis": result.legal_analysis,
            "status": result.status,
            # "relevant_laws": result.relevant_laws,
            # "relevant_judgements": result.relevant_judgements
        }
    
    def summarise_sync(self, ai_content_text: str) -> str:
        """Synchronous wrapper for summarise method with improved event loop handling"""
        try:
            # Try to get existing event loop
            loop = asyncio.get_event_loop()
            if loop.is_running():
                # If loop is already running, create a new thread
                import concurrent.futures
                
                def run_in_thread():
                    # Create new event loop for this thread
                    new_loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(new_loop)
                    try:
                        return new_loop.run_until_complete(self.summarise(ai_content_text))
                    finally:
                        new_loop.close()
                
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(run_in_thread)
                    return future.result()
            else:
                # No running loop, use asyncio.run
                return asyncio.run(self.summarise(ai_content_text))
        except RuntimeError as e:
            if "no current event loop" in str(e).lower() or "event loop is closed" in str(e).lower():
                # Create new event loop
                return asyncio.run(self.summarise(ai_content_text))
            else:
                raise e
